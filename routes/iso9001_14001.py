from fastapi import APIRouter
from fastapi import FastAPI, Response, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from typing import List, Optional
from docxtpl import DocxTemplate
from docx import Document
import io
import json
import math
import random
import httpx
import re
import traceback
from fastapi.responses import JSONResponse
from fastapi.responses import StreamingResponse
import ast
import zipfile
import asyncio
from datetime import datetime, timedelta

router = APIRouter()
stage1_minor_nc_store = []
stage1_observation_store = []

# =================== MODELS ===========================
class ISO9001_14001Stage1Audit(BaseModel):
    organizationName: str
    address: str
    siteAddress: Optional[str] = ""
    numberOfEmployees: int
    emailId: EmailStr
    contactPerson: str
    telephoneFax: str
    scope: str
    riskCategory: str
    iafCode: str
    auditTeam: List[str]
    auditManDays: str
    startDateOfAuditStage1: str
    endDateOfAuditStage1: str
    startDateOfAuditStage2: Optional[str] = None
    endDateOfAuditStage2: Optional[str] = None
    auditMode: str
    quotedManDaysAdequate: str
    changeInEmployeeDetail: str
    changeInScope: str
    additionalInformation: str
    internalAuditFrequency: str
    dateOfLastInternalAudit: str
    managementReviewFrequency: str
    dateOfLastManagementReview: str
    recommendationForStage2: str
    reviewedBy: str
    na_clauses: Optional[List[str]] = []
    attendanceSheet: List[str]
    dateOfReview: str
    clientName: str
    designation: str
    auditorName: str

class ISO9001_14001Stage2Audit(BaseModel):
    organizationName: str
    address: str
    siteAddress: Optional[str] = ""
    numberOfEmployees: int
    emailId: EmailStr
    contactPerson: str
    telephoneFax: str
    scope: str
    riskCategory: str
    iafCode: str
    auditTeam: List[str]
    auditManDays: str
    startDateOfAuditStage1: str  # Required
    endDateOfAuditStage1: str  # Required
    startDateOfAuditStage2: str
    endDateOfAuditStage2: str
    auditMode: str
    quotedManDaysAdequate: str
    changeInEmployeeDetail: str
    changeInScope: str
    additionalInformation: str
    internalAuditFrequency: str
    dateOfLastInternalAudit: str
    managementReviewFrequency: str
    dateOfLastManagementReview: str
    recommendationForStage2: str
    reviewedBy: str
    na_clauses: Optional[List[str]] = []
    attendanceSheet: List[str]
    dateOfReview: str
    clientName: str
    designation: str
    auditorName: str

class CombinedAuditRequest(BaseModel):
    stage1_audit: ISO9001_14001Stage1Audit
    stage2_audit: ISO9001_14001Stage2Audit
# ==============================================

def generate_document_dates(clause_map, stage1_start_date_str):
    """
    Generate a fixed random date for each unique document in clause_map,
    between 7–10 months before the Stage 1 start date.
    """
    start_date = datetime.strptime(stage1_start_date_str, "%Y-%m-%d")
    date_map = {}

    for clause, docs in clause_map.items():
        for doc in docs:
            key = f"{doc['Document Name']}|{doc['Document Number']}"
            if key not in date_map:
                # Offset between 7–10 months
                months_offset = random.randint(7, 10)
                days_offset = random.randint(0, 29)
                doc_date = start_date - timedelta(days=(months_offset * 30) + days_offset)
                date_map[key] = doc_date.strftime("%d-%b-%Y")

    return date_map


def extract_observation_rows(rows):
    results = []
    for row in rows:
        # Find C/NC/O or status column
        status = None
        for key in row.keys():
            if key.strip().lower().replace(" ", "") in ("c/nc/o", "c/nc/o.", "status"):
                status = row[key]
                break
        if status is None:
            continue
        # Normalize
        status_norm = str(status).strip().upper()
        if status_norm != "O":
            continue
        # Skip NA evidences
        evidence_val = None
        for key in row.keys():
            if "verification" in key.lower() or "conformity" in key.lower():
                evidence_val = row[key]
                break
        if evidence_val and str(evidence_val).strip().upper() == "NA":
            continue
        results.append(row)
    print(f"[DEBUG] extract_observation_rows: found {len(results)} O rows out of {len(rows)}")
    return results

def build_observation_summary_prompt(obs_rows):
    return f"""
You are summarizing ISO45001 audit Observations for an IMS audit report.

For each input item, write a short, factual summary (max 2–3 lines) describing the observation noted.

**Start each summary exactly like this**: Clause :

Output rules:
Output rules:
- One observation summary per clause
- Plain text only, separated by blank lines
- No JSON, bullets, code fences, or extra commentary
- The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols from markdown (- or * as formatting), tables, headings, or any other non-standard formatting.
- Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).

Input data:
{json.dumps(obs_rows, indent=2, ensure_ascii=False)}

Now return only the plain sentence summaries, one per clause, separated by blank lines.
"""

def clean_observation_summaries(summary_text):
    text = summary_text.strip().strip("`")
    text = re.sub(r"^``````$", "", text, flags=re.MULTILINE).strip()
    if text.startswith("[") or text.startswith("{"):
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                return [str(item.get("summary") if isinstance(item, dict) else str(item)).strip() for item in parsed if item]
            elif isinstance(parsed, dict) and "summary" in parsed:
                return [parsed["summary"].strip()]
        except:
            pass
    return [line.strip(" -•\t") for line in text.splitlines() if line.strip()]

def patch_observations_table(docx_buffer, summaries):
    from docx import Document
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    for table in doc.tables:
        if any("Observations raised:" in cell.text for cell in table.row_cells(0)):
            while len(table.rows) > 1:
                tbl = table._tbl
                tbl.remove(tbl.tr_lst[-1])
            for summary in summaries:
                row = table.add_row()
                row.cells[0].text = summary
            break
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

async def transfer_stage1_observations_to_stage2_doc(patched_buffer, audit, mistral_api_url, headers):
    print(f"[DEBUG] Transferring {len(stage1_observation_store)} Stage-1 Observations to Stage-2")
    if stage1_observation_store:
        attendance_text = ", ".join(audit.attendanceSheet)
        scope_text = audit.scope
        combined_entries = []
        for obs in stage1_observation_store:
            combined_entries.append(
                f"Clause {obs.get('Cl. No', '')}\nObservation: {obs.get('summary', '')}"
            )
    else:
        combined_entries = ["No observations from Stage 1."]

    from docx import Document
    patched_buffer.seek(0)
    doc = Document(patched_buffer)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Observations raised in Stage-1:" in cell.text:
                    cell.text = "Observations raised in Stage-1:\n\n" + "\n\n".join(combined_entries)
                    break
            else:
                continue
            break
        else:
            continue
        break
    patched_buffer.seek(0)
    patched_buffer.truncate(0)
    doc.save(patched_buffer)
    patched_buffer.seek(0)
    return patched_buffer

def remove_markdown_styling(text: str) -> str:
    """
    Removes markdown bold/italic markers (*, **) without altering other text.
    This avoids accidentally removing asterisks used for math or other purposes.
    """
    if not isinstance(text, str):
        return text
    # Remove **bold** or *italic* markers
    cleaned = re.sub(r'(\*\*|\*)', '', text)
    return cleaned

async def generate_completed_corrective_actions(stage1_minor_nc_store, scope_text, attendance_text, mistral_api_url, headers):
    """
    Given Stage 1 NCs, scope, and attendance, ask LLM for a past-tense corrective action for each NC.
    Returns a list of strings aligned with the store order.
    """
    actions = []
    for nc in stage1_minor_nc_store:
        clause_no = nc.get("Cl. No", "")
        summary = nc.get("summary", "")
        prompt = (
            f"For ISO clause {clause_no}, scope: {scope_text}, "
            f"attendance: {attendance_text}, "
            f"and the following nonconformity: \"{summary}\", "
            "generate a random, realistic corrective action that HAS ALREADY BEEN IMPLEMENTED "
            "and is now completed. "
            "Write the action in the past tense, e.g., 'The XYZ procedure was revised and all staff were trained accordingly.' "
            "Output only a concise, plain-English paragraph of 2-4 sentences describing what was done, "
            "including any implementation and verification. "
            "Do NOT return JSON, bullet points, lists, or code fences — just the text description."
        )
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(mistral_api_url, json={"prompt": prompt}, headers=headers)
            resp.raise_for_status()
            action = resp.json().get("response", "") if resp.headers.get("content-type", "").startswith("application/json") else resp.text
            actions.append(action.strip())
    return actions

def clean_corrective_action_text(raw_text: str) -> str:
    """
    Cleans corrective action text from LLM.
    If JSON is detected, extracts and flattens meaningful fields into a sentence.
    Removes markdown code fences.
    """
    if not raw_text:
        return ""

    # Remove markdown/json code fences
    cleaned = re.sub(r"^``````$", "", raw_text.strip(), flags=re.MULTILINE).strip()

    # Try parsing as JSON
    try:
        obj = json.loads(cleaned)
        # If top-level dict and has corrective_action
        if isinstance(obj, dict) and "corrective_action" in obj:
            ca_obj = obj["corrective_action"]
            if isinstance(ca_obj, dict):
                parts = []
                for key in ["action_taken", "implementation", "verification"]:
                    if ca_obj.get(key):
                        parts.append(str(ca_obj[key]).strip())
                return " ".join(parts)
        # If obj itself is a string, return it
        if isinstance(obj, str):
            return obj
    except Exception:
        pass  # if not JSON, just keep original text

    return cleaned

#  ================ STAGE-2 FUNCTIONS =======================
def extract_audit_table_iso9001_14001_stage2(docx_path_or_stream):
    """
    Extracts the audit clause verification table for ISO 9001/14001:2015 Stage-2 from a docx document.
    Returns a list of dicts, one per clause row, ONLY IF 'Cl. No' is not empty.
    """
    doc = Document(docx_path_or_stream)
    data = []

    # Possible column names (case-insensitive)
    cl_no_cols = {"cl. no", "clause"}
    desc_cols = {"description"}
    status_cols = {"c/nc/o", "status", "c / nc / o", "conformity"}
    doc_ver_cols = {"document verification detail with statement of conformity"}

    for table in doc.tables:
        # Read header row (first row)
        header = [cell.text.strip().lower() for cell in table.rows[0].cells]
        # Ensure all relevant columns present
        try:
            cl_no_idx = next(i for i, txt in enumerate(header) if txt in cl_no_cols)
            desc_idx = next(i for i, txt in enumerate(header) if txt in desc_cols)
            status_idx = next(i for i, txt in enumerate(header) if txt in status_cols)
            doc_ver_idx = next(i for i, txt in enumerate(header) if txt in doc_ver_cols)
        except StopIteration:
            continue

        # Table with expected headers found, extract rows
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) > max(cl_no_idx, desc_idx, status_idx, doc_ver_idx):
                cl_no = cells[cl_no_idx].text.strip()
                # Only add row if 'Cl. No' is NOT empty
                if cl_no:
                    entry = {
                        "Cl. No": cl_no,
                        "Description": cells[desc_idx].text.strip(),
                        "C/NC/O": cells[status_idx].text.strip(),
                        "Document Verification detail with statement of Conformity": cells[doc_ver_idx].text.strip()
                    }
                    data.append(entry)
        break  # Remove break if multiple such tables might exist

    return data

def update_cnc_placeholders_stage2(rows):
    """
    For ISO 9001/14001 Stage 2 tables: fill blank or '{{clause}}' 'C/NC/O' fields as follows:
      - 80% 'C'
      - 10% 'O'
      - 10% 'NC'
    Skip rows where 'Document Verification detail with statement of Conformity' is 'NA'.
    """
    def is_fillable(val):
        return not val or val.strip().lower() == "{{clause}}"

    evidence_field = "Document Verification detail with statement of Conformity"
    c_nc_o_field = "C/NC/O"

    indices_to_fill = [
        idx for idx, row in enumerate(rows)
        if is_fillable(row.get(c_nc_o_field, "")) and row.get(evidence_field, "").strip().upper() != "NA"
    ]

    total = len(indices_to_fill)
    if total == 0:
        return rows

    # Calculate fill counts for each status; sum to total
    nc_count = max(1, round(0.10 * total))
    o_count = max(1, round(0.10 * total))
    c_count = total - nc_count - o_count

    while nc_count + o_count + c_count > total:
        c_count = max(0, c_count - 1)
    while nc_count + o_count + c_count < total:
        c_count += 1

    replacements = (["C"] * c_count) + (["NC"] * nc_count) + (["O"] * o_count)
    random.shuffle(replacements)

    for i, idx in enumerate(indices_to_fill):
        rows[idx][c_nc_o_field] = replacements[i]

    # Explicitly clear C/NC/O for rows with evidence 'NA'
    for row in rows:
        if row.get(evidence_field, "").strip().upper() == "NA":
            row[c_nc_o_field] = ""

    return rows

def generate_prompt_for_iso9001_14001_stage2(batch, audit, clause_map, prompt_table_md, pattern_desc):
    # Build clause-specific prompts if you use them (customize per your IMS clause_map)
    stage2_prompts = []
    for clause, docs in clause_map.items():
        for doc in docs:
            if "Stage 2 Prompt" in doc and doc["Stage 2 Prompt"]:
                stage2_prompts.append(f"Clause {clause}: {doc['Stage 2 Prompt']}")
    stage2_prompt_text = "\n".join(stage2_prompts)

    attendance_list_text = "\n".join([f"- {member}" for member in audit.attendanceSheet])

    return f"""
You are an ISO 9001:2015 & ISO 14001:2015 Stage 2 Integrated Management System (IMS) audit reporting assistant.

Use the following document numbering format throughout all evidence:  
**Pattern**: {pattern_desc}  
When referencing documents, use ONLY the document NAME and NUMBER as provided in the input table for each clause.  
- If a document number has a prefix like "XXX" or "BLPL" (e.g., "XXX-EMS-F-01"), you MUST replace the prefix with the initials of the organization's name when writing the report.Only do this when document pattern starts with XXX . Dont do this if its just F-X or P-X.
- Dont modify the document number or details randomly.

{prompt_table_md}

--- 
### ABSOLUTE FORMATTING RULES (PLAIN TEXT ONLY – STRICT):
- The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols from markdown (- or * as formatting), tables, headings, or any other non-standard formatting.
- Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).
- Write all content in normal sentences using only letters, numerals, and standard punctuation.
- Document names and numbers must be written exactly as provided, without surrounding symbols or formatting.
- For spacing, only use actual line breaks; no markdown or decorative spacing.
- Even if the input contains markdown or symbols, remove them in the output — ensure the output is fully cleaned.
- Any output that contains forbidden formatting is invalid.

**STRICT and REDUNDANT RULES (do NOT break them):**
- For each clause, ONLY list as evidence the exact documents and their numbers provided for that clause in the input.
- If a clause has NO listed documents, do NOT mention, imply, or invent *any* document in the answer for that clause.
- NEVER create, paraphrase, infer, or generate document names or numbers based on the pattern or clause context.
- Do NOT combine, split, or otherwise modify listed document names/numbers.
- **In summary:** ONLY reference documents exactly as listed. DO NOT reference documents for a clause if none are provided.

Here are detailed prompts for each clause to guide your evidence generation:
{stage2_prompt_text}

---

### Audit Details:
- Organization: {audit.organizationName}
- IMS Scope: {audit.scope}
- Address: {audit.address}
- Stage 2 Audit Dates: {audit.startDateOfAuditStage2} to {audit.endDateOfAuditStage2}
- Stage 1 Audit Dates: {audit.startDateOfAuditStage1} to {audit.endDateOfAuditStage1}

### Attendance Sheet:
Below is the list of personnel present during the audit. Use these names accurately when drafting evidence. Assign relevant titles/roles (e.g., CEO, QA Manager, EMS Manager, Compliance Officer, etc.) from this list.

{attendance_list_text}

---

### Instructions for IMS Stage 2 Report Writing:
- Only update the 'evidence' field of each input item.
- Do NOT alter or remove any other fields (e.g., 'Cl. No', 'Description', 'C/NC/O').
- For 'C' (Conformity): Rephrase the "evidence" as a factual, positive confirmation that requirements of ISO 9001:2015 and/or ISO 14001:2015 for that clause are met, referencing only the clause(s) (e.g., 4.3, 7.2) and any relevant listed document(s).
- For 'NC' (Nonconformity): Clearly state what does not conform to ISO 9001:2015 and/or ISO 14001:2015, referencing only listed clause(s) and document(s).
- For 'O' (Observation): Reword the evidence as a neutral, factual observation, referencing only the listed clause(s) and document(s).
- If the 'Description' field includes multiple questions, write a comprehensive, structured response that clearly addresses **each question in order**.
- STRICTLY follow the order of batch items; do NOT change structure or order—modify 'evidence' only.
- If a clause has any document(s) referenced, use ONLY their names/numbers as provided, and date your evidence with a random date 7-10 months prior to the Stage 1 audit.
- Do NOT add, merge, or invent document references under any circumstances. Omit document references if none are listed.
- Use specific names and roles from the attendance list in your responses as appropriate.
- Responses must align with both ISO 9001:2015 and ISO 14001:2015 Stage 2 audit standards wherever the clause applies to one or both.
- Ensure every answer is separated by a blank line (two newlines) for clarity.
- Output must be only the list of dictionaries, updated as per these rules.
- Use document dates as specified in the prompt table. Dont generate dates randomly.
---

### Input:
Here is the list of clauses and requirements. Do NOT change structure—edit only the 'evidence' field.

{json.dumps(batch, indent=2, ensure_ascii=False)}

---

### Output:
Respond with ONLY the list of dictionaries, with revised 'evidence' fields.  
Do NOT add markdown, comments, or extra text.  
Separate each answer by a single blank line (\\n\\n) for readability.
"""

def choose_document_pattern_stage2(forced_pattern_name=None, date_map=None):
    """
    Randomly select one document-numbering pattern for ISO audit document references.
    Returns:
        pattern_name: 'ims_org', 'ims_only', 'qhse', or 'minimal'
        pattern_description: human-readable summary
        clause_map: dict mapping clause -> list of dicts with 'Document Name', 'Document Number'
        prompt_table: Markdown table as string for prompt
    """
    # --- Pattern definitions (expand for all clauses as needed) ---
    # Pattern 1: Org initials + IMS, e.g. XXX-IMS-F-01
    pattern_1 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-IMS-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe in detail about company activities, top management name"
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "XXX-IMS-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "XXX-IMS-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "XXX-IMS-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information according to the scope and address of the company."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "XXX-IMS-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the documented scope and also create a evidence(like purchase order, invoice or contract) to justify the scope."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "XXX-IMS-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the name of person in top management using the attendance sheet."
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored through websites, emails, display at reception."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "XXX-IMS-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also, mention a employee from the attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "XXX-IMS-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in IMS activities. Mention a name from the attendance sheet."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "XXX-IMS-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Mention 4-5 risks according to the scope and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "XXX-IMS-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "XXX-IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "XXX-IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "XXX-IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "XXX-IMS-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Give an example of recent new/changed requirement tracked and acted on. Also mention the legal requirement based on the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "XXX-IMS-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "XXX-IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "XXX-IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "XXX-IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "XXX-IMS-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Write at least 4 quality, environment, health and safety objectives and write the achieved results so far."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "XXX-IMS-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "XXX-IMS-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "XXX-IMS-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "XXX-IMS-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "XXX-IMS-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year.Mention the name of employee that competency verified from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "XXX-IMS-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions.Mention 2 to 3 training topics to be delivered in year based on the scope."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "XXX-IMS-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "XXX-IMS-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "XXX-IMS-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected.Mention a name to justify this from the attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-IMS-MAN-01",
                "Guidance/Description": "Manual describing the organization's IMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-IMS-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "XXX-IMS-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters). Also, list the topics what to communicate, when to communicate, when to communicate."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "XXX-IMS-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "XXX-IMS-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-IMS-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 2 Prompt": (
                    "List all core operational procedures maintained under the IMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "XXX-IMS-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatment",
                "Document Number": "XXX-IMS-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "XXX-IMS-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "XXX-IMS-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Create a sample of any product or service delivered to any client as per the scope and address of the company."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "XXX-IMS-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "XXX-IMS-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project. Create a sample of product based on the scope of the company."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "XXX-IMS-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "XXX-IMS-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "XXX-IMS-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "XXX-IMS-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "XXX-IMS-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "XXX-IMS-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "XXX-IMS-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome.Mention a HSE inspection record verified based on the scope."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "XXX-IMS-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "XXX-IMS-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-IMS-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "XXX-IMS-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "XXX-IMS-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "XXX-IMS-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-IMS-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "XXX-IMS-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "XXX-IMS-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification based on the scope of the company."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "XXX-IMS-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "XXX-IMS-P-17",
                "Guidance/Description": "Defines how IMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "XXX-IMS-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "XXX-IMS-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a recent update or review record for legal requirements based on the address and scope of the company, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "XXX-IMS-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "XXX-IMS-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "XXX-IMS-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "XXX-IMS-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "XXX-IMS-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "XXX-IMS-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "XXX-IMS-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "XXX-IMS-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "XXX-IMS-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "XXX-IMS-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "XXX-IMS-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "XXX-IMS-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process. Mention a product or sample."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    # Pattern 2: IMS only (IMS-...)
    pattern_2 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "IMS-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe in detail about company activities, top management name"
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "IMS-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "IMS-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "IMS-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information according to the scope and address of the company."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "IMS-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the documented scope and also create a evidence(like purchase order, invoice or contract) to justify the scope."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "IMS-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the name of person in top management using the attendance sheet."
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored through websites, emails, display at reception."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "IMS-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also, mention a employee from the attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "IMS-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in IMS activities. Mention a name from the attendance sheet."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "IMS-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Mention 4-5 risks according to the scope and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "IMS-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "IMS-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Give an example of recent new/changed requirement tracked and acted on. Also mention the legal requirement based on the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "IMS-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "IMS-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Write at least 4 quality, environment, health and safety objectives and write the achieved results so far."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "IMS-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "IMS-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "IMS-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "IMS-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "IMS-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year.Mention the name of employee that competency verified from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "IMS-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions.Mention 2 to 3 training topics to be delivered in year based on the scope."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "IMS-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "IMS-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "IMS-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected.Mention a name to justify this from the attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "IMS-MAN-01",
                "Guidance/Description": "Manual describing the organization's IMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "IMS-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "IMS-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters). Also, list the topics what to communicate, when to communicate, when to communicate."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "IMS-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "IMS-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "IMS-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 2 Prompt": (
                    "List all core operational procedures maintained under the IMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "IMS-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatment",
                "Document Number": "IMS-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "IMS-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "IMS-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Create a sample of any product or service delivered to any client as per the scope and address of the company."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "IMS-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "IMS-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project. Create a sample of product based on the scope of the company."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "IMS-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "IMS-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "IMS-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "IMS-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "IMS-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "IMS-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "IMS-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome.Mention a HSE inspection record verified based on the scope."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "IMS-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "IMS-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "IMS-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "IMS-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "IMS-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "IMS-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "IMS-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "IMS-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "IMS-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification based on the scope of the company."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "IMS-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "IMS-P-17",
                "Guidance/Description": "Defines how IMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "IMS-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "IMS-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a recent update or review record for legal requirements based on the address and scope of the company, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "IMS-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "IMS-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "IMS-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "IMS-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "IMS-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "IMS-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "IMS-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "IMS-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "IMS-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "IMS-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "IMS-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "IMS-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process. Mention a product or sample."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    pattern_3 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe in detail about company activities, top management name"
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "QHSE-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "QHSE-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "QHSE-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information according to the scope and address of the company."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "QHSE-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the documented scope and also create a evidence(like purchase order, invoice or contract) to justify the scope."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "QHSE-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the name of person in top management using the attendance sheet."
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored through websites, emails, display at reception."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "QHSE-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also, mention a employee from the attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "QHSE-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in IMS activities. Mention a name from the attendance sheet."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "QHSE-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Mention 4-5 risks according to the scope and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "QHSE-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "QHSE-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "QHSE-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "QHSE-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "QHSE-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Give an example of recent new/changed requirement tracked and acted on. Also mention the legal requirement based on the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "QHSE-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "QHSE-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "QHSE-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "QHSE-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "QHSE-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Write at least 4 quality, environment, health and safety objectives and write the achieved results so far."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "QHSE-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "QHSE-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "QHSE-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "QHSE-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "QHSE-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year.Mention the name of employee that competency verified from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "QHSE-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions.Mention 2 to 3 training topics to be delivered in year based on the scope."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "QHSE-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "QHSE-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "QHSE-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected.Mention a name to justify this from the attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Manual describing the organization's IMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "QHSE-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters). Also, list the topics what to communicate, when to communicate, when to communicate."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "QHSE-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "QHSE-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 2 Prompt": (
                    "List all core operational procedures maintained under the IMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatment",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "QHSE-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "QHSE-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Create a sample of any product or service delivered to any client as per the scope and address of the company."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "QHSE-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "QHSE-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project. Create a sample of product based on the scope of the company."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "QHSE-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "QHSE-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "QHSE-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "QHSE-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "QHSE-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "QHSE-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "QHSE-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome.Mention a HSE inspection record verified based on the scope."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "QHSE-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "QHSE-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "QHSE-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "QHSE-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "QHSE-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "QHSE-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "QHSE-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification based on the scope of the company."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "QHSE-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "QHSE-P-17",
                "Guidance/Description": "Defines how IMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "QHSE-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "QHSE-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a recent update or review record for legal requirements based on the address and scope of the company, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "QHSE-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "QHSE-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "QHSE-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "QHSE-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "QHSE-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "QHSE-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "QHSE-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "QHSE-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "QHSE-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "QHSE-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "QHSE-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "QHSE-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process. Mention a product or sample."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    pattern_4 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe in detail about company activities, top management name"
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information according to the scope and address of the company."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the documented scope and also create a evidence(like purchase order, invoice or contract) to justify the scope."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the name of person in top management using the attendance sheet."
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored through websites, emails, display at reception."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 2 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also, mention a employee from the attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in IMS activities. Mention a name from the attendance sheet."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Mention 4-5 risks according to the scope and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Give an example of recent new/changed requirement tracked and acted on. Also mention the legal requirement based on the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Mention any 4 enviornmental aspect based on the scope, and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Mention any 4 hazards based on the scope and elaborate the assessment methodology."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 2 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Write at least 4 quality, environment, health and safety objectives and write the achieved results so far."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 2 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year.Mention the name of employee that competency verified from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions.Mention 2 to 3 training topics to be delivered in year based on the scope."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 2 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected.Mention a name to justify this from the attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Manual describing the organization's IMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters). Also, list the topics what to communicate, when to communicate, when to communicate."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 2 Prompt": (
                    "List all core operational procedures maintained under the IMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatment",
                "Document Number": "F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Create a sample of any product or service delivered to any client as per the scope and address of the company."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 2 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project. Create a sample of product based on the scope of the company."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome.Mention a HSE inspection record verified based on the scope."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 2 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification based on the scope of the company."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 2 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "P-17",
                "Guidance/Description": "Defines how IMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a recent update or review record for legal requirements based on the address and scope of the company, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 2 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 2 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process. Mention a product or sample."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }


    patterns = [
        ("IMS_org",   "Org initials + IMS (XXX-IMS-...)",           pattern_1),
        ("IMS_only",  "IMS only (IMS-...)",                         pattern_2),
        ("qhse",      "QHSE system (QHSE-...)",                     pattern_3),
        ("minimal",   "Minimal prefix (MAN-01, P-01, etc.)",        pattern_4),
    ]
    if forced_pattern_name is not None:
        for pn, pd, cm in patterns:
            if pn == forced_pattern_name:
                pattern_name, pattern_desc, clause_map = pn, pd, cm
                break
        else:
            pattern_name, pattern_desc, clause_map = patterns[0]
    else:
        pattern_name, pattern_desc, clause_map = random.choice(patterns)

    # Generate full markdown table including all columns
    lines = [
        "| Clause | Document Name | Document Number | Document Date | Guidance/Description | Document Owner | Approved By | Stage 2 Prompt |",
        "|--------|---------------|----------------|--------------|----------------------|---------------|-------------|----------------|"
    ]
    for clause, docs in clause_map.items():
        for doc in docs:
            key = f"{doc['Document Name']}|{doc['Document Number']}"
            fixed_date = date_map.get(key, "") if date_map else ""
            lines.append(
                f"| {clause} | {doc['Document Name']} | {doc['Document Number']} | {fixed_date} | "
                f"{doc.get('Guidance/Description', '')} | {doc.get('Document Owner', '')} | "
                f"{doc.get('Approved By', '')} | {doc.get('Stage 2 Prompt', '')} |"
            )
    prompt_table = "\n".join(lines)

    return pattern_name, pattern_desc, clause_map, prompt_table

def patch_docx_by_row_index_iso9001_14001_stage2(docx_buffer, audit_rows, table_idx=None, data_start_idx=1):
    """
    Patch the ISO 9001/14001:2015 Stage 2 findings table in a Word docx buffer.
    - Finds the findings table if not given.
    - Patches ONLY rows with a non-empty 'Cl. No'.
    - Expects audit_rows as a list of dicts with keys:
        'Cl. No', 'Description', 'C/NC/O', 'Document Verification detail with statement of Conformity'
    - Leaves any client-info or blank 'Cl. No' rows UNTOUCHED.
    """
    docx_buffer.seek(0)
    doc = Document(docx_buffer)

    # Locate table by headers if needed
    if table_idx is None:
        expected_headers = [
            {"cl. no", "clause"},              # col 0
            {"description"},                   # col 1
            {"c/nc/o"},                        # col 2
            {"document verification detail with statement of conformity"}, # col 3
        ]
        for idx, table in enumerate(doc.tables):
            if not table.rows or len(table.rows[0].cells) < 4:
                continue
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if all(any(h in headers[i] for h in colset) for i, colset in enumerate(expected_headers)):
                table_idx = idx
                break
        else:
            raise ValueError("Could not locate the IMS Stage 2 findings table in DOCX.")

    table = doc.tables[table_idx]
    col_keys = ["Cl. No", "Description", "C/NC/O", "Document Verification detail with statement of Conformity"]

    audit_idx = 0
    for trow in table.rows[data_start_idx:]:
        if len(trow.cells) < 4:
            continue
        cl_no_cell = trow.cells[0].text.strip()
        if cl_no_cell:  # Only patch rows with non-empty 'Cl. No'
            if audit_idx >= len(audit_rows):
                break
            arow = audit_rows[audit_idx]
            for col, key in enumerate(col_keys):
                trow.cells[col].text = str(arow.get(key, ""))
            audit_idx += 1
        # Otherwise, leave the row untouched

    # Optionally: warn/log discrepancies
    if audit_idx < len(audit_rows):
        print(f"⚠️ Warning: Not all audit_rows were patched ({audit_idx} of {len(audit_rows)})")
    elif audit_idx > len(audit_rows):
        print(f"⚠️ Warning: More table lines than data rows ({audit_idx} > {len(audit_rows)})")

    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

async def transfer_stage1_ncs_to_stage2_doc(patched_buffer, audit, mistral_api_url, headers):
    print(f"[DEBUG] Transferring {len(stage1_minor_nc_store)} Stage-1 NCs to Stage-2")
    if stage1_minor_nc_store:
        attendance_text = ", ".join(audit.attendanceSheet)
        scope_text = audit.scope
        actions = await generate_completed_corrective_actions(
            stage1_minor_nc_store,
            scope_text,
            attendance_text,
            mistral_api_url,
            headers
        )
        combined_entries = []
        for nc, action in zip(stage1_minor_nc_store, actions):
            cleaned_action = clean_corrective_action_text(action)
            combined_entries.append(
                f"Clause {nc.get('Cl. No', '')}\n"
                f"Nonconformity: {nc.get('summary', '')}\n"
                f"Corrective action taken: {cleaned_action}"
            )
    else:
        combined_entries = ["No nonconformities from Stage 1."]

    from docx import Document
    patched_buffer.seek(0)
    doc = Document(patched_buffer)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Non conformities raised in Stage-1:" in cell.text:
                    cell.text = "Non conformities raised in Stage-1:\n\n" + "\n\n".join(combined_entries)
                    break
            else:
                continue
            break
        else:
            continue
        break

    patched_buffer.seek(0)
    patched_buffer.truncate(0)
    doc.save(patched_buffer)
    patched_buffer.seek(0)
    return patched_buffer


# =================== STAGE-1 FUNCTIONS ====================

def clean_llm_law_list(text: str) -> str:
    """Cleans up code block/JSON output for readable Word insertion."""
    # Remove markdown code block
    cleaned = re.sub(r"^``````$", "", text.strip(), flags=re.MULTILINE).strip()
    # Try JSON parse for ["1. Law...", ...]
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, list):
            return "\n".join(obj)
    except Exception:
        pass
    # Remove array brackets if present
    if cleaned.startswith("[") and cleaned.endswith("]"):
        # Remove possible starting/ending quotes/commas
        inner = cleaned[1:-1]
        lines = [x.strip().strip('",') for x in inner.split("\n") if x.strip()]
        lines = [x for x in lines if x]
        return "\n".join(lines)
    # If each line is a quoted string, remove
    lines = []
    for line in cleaned.splitlines():
        stripped = line.strip().strip('"').strip(",")
        if stripped:
            lines.append(stripped)
    return "\n".join(lines)

async def add_legal_requirements_to_docx_iso9001_14001_mistral(
    docx_buffer,
    address,
    scope,
    mistral_url="https://mistral-api-v2.onrender.com/api/mistral"
):
    """
    Calls Mistral LLM for Legal, Statutory & Regulatory Requirements relevant to IMS (ISO 9001 & 14001),
    and inserts the up-to-date law list into the correct cell in the docx buffer.
    """
    # 1. Prompt (for IMS: quality + environment)
    legal_prompt = f"""
You are an ISO 9001 (quality management) and ISO 14001 (environmental management) integrated management system (IMS) audit assistant.

Based only on the provided company address (use it to determine the country) and the organization’s scope, generate a well-organized, numbered list of all important, up-to-date legal, statutory, and regulatory requirements (laws, acts, rules, or major regulations) relevant to BOTH quality and environmental management for that country and scope.

Instructions:
- Do not invent non-existent statutes; use the country and sector as clues.
- Include environmental AND quality/legal requirements: e.g. product quality standards, manufacturing/industry-specific regulations, environmental/EHS/waste/air/water/pollution/fire/chemical rules, and any sector-specific EMS/QMS rules.
- Maximum 20 requirements. List only the names of legal instruments (law name ± national code/year), one per line; NO explanations, NO preface, NO commentary, NO markdown or JSON.
- Output ONLY the numbered law list.

Company Address: {address}
IMS Scope: {scope}
    """

    # 2. Call Mistral API
    async with httpx.AsyncClient(timeout=60.0) as client:
        api_response = await client.post(
            mistral_url,
            json={"prompt": legal_prompt},
            headers={"Content-Type": "application/json"}
        )
        api_response.raise_for_status()
        if api_response.headers.get("content-type", "").startswith("application/json"):
            law_list = api_response.json().get("response", "") or api_response.text
        else:
            law_list = api_response.text
    law_list = law_list.strip()
    law_list = clean_llm_law_list(law_list)  # Reuse your clean-up function for formatting

    # 3. Replace cell in docx buffer
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    written = False
    for table in doc.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                # Adjust keyword(s) if your IMS template uses a distinct heading
                if "Legal, Statutory" in cell.text:
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = law_list
                        written = True
                    else:
                        cell.text = "Legal, Statutory & Regulatory Requirements:\n\n" + law_list
                        written = True
    if not written:
        print("⚠️ Could not find 'Legal, Statutory & Regulatory Requirements' cell in the DOCX.")
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

async def add_org_brief_to_docx_iso9001_14001(
    docx_buffer,
    company_name,
    scope,
    mistral_url="https://mistral-api-v2.onrender.com/api/mistral"
):
    """
    Calls Mistral for an organization brief and inserts it into the IMS (ISO 9001+14001) DOCX stage 1 cell.
    Only the actual business/activities brief will appear (no JSON, no standards mention).
    """
    # 1. Strict, ISO-agnostic prompt
    brief_prompt = f"""
You are writing the opening organization summary for an IMS (ISO 9001/14001) Stage 1 audit report.

Based ONLY on the following company name and scope, write a concise, professional 2–3 sentence overview of this organization's main activities, products/services, and business focus.

DO NOT mention ISO, certifications, standards, quality/environmental compliance, or audit processes in any form.

- Company Name: {company_name}
- IMS Scope: {scope}

Output ONLY the brief text itself — do NOT include code block formatting, preface, or output as JSON. Output only the brief.
    """

    # 2. Call Mistral API
    async with httpx.AsyncClient(timeout=60.0) as client:
        api_response = await client.post(
            mistral_url,
            json={"prompt": brief_prompt},
            headers={"Content-Type": "application/json"}
        )
        api_response.raise_for_status()
        if api_response.headers.get("content-type", "").startswith("application/json"):
            result = api_response.json()
            brief_string = result.get("response", "") or str(result)
        else:
            brief_string = api_response.text
    brief_string = brief_string.strip()

    # 3. Robust extraction from all weird result formats
    for key in ("brief", "overview"):
        try:
            obj = json.loads(brief_string)
            if isinstance(obj, dict) and key in obj:
                brief_string = obj[key]
                break
        except Exception:
            pass
    # Remove any lingering codeblock or non-text artifacts:
    if brief_string.startswith("``````"):
        brief_string = brief_string.strip("`").strip()
    # Defensive: remove any lines about ISO or "standard"
    import re
    brief_string = "\n".join([
        line for line in brief_string.splitlines()
        if not re.search(r'(ISO ?9|ISO ?1|14001|certifi|standard|compliance)', line, re.I)
    ]).strip()

    # 4. Insert into DOCX
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    written = False
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            for idx, cell in enumerate(row.cells):
                if ("Brief about the organization" in cell.text) or ("Organization Brief" in cell.text):
                    # Insert into next cell (second column)
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = brief_string
                        written = True
                    else:
                        cell.text = "Brief about the organization:\n\n" + brief_string
                        written = True
    if not written:
        print("⚠️ Could not find 'Brief about the organization' cell in the DOCX.")
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


def org_initials(org_name: str) -> str:
    # E.g., "Livpure Limited" → "LL"
    return "".join([w[0].upper() for w in org_name.split() if w and w[0].isalpha()])

def generate_prompt_for_stage1(batch, audit, clause_map, prompt_table_md, pattern_desc):
    # Collect clause-specific prompts, if you use them (ISO 9001/14001/IMS blend)
    stage1_prompts = []
    for clause, docs in clause_map.items():
        for doc in docs:
            if "Stage 1 Prompt" in doc and doc["Stage 1 Prompt"]:
                stage1_prompts.append(f"Clause {clause}: {doc['Stage 1 Prompt']}")
    stage1_prompt_text = "\n".join(stage1_prompts)

    attendance_list_text = "\n".join([f"- {member}" for member in audit.attendanceSheet])

    return f"""
You are an audit reporting assistant for an ISO 9001:2015 and ISO 14001:2015 **Stage 1 Integrated Management System (IMS)** assessment.

Use the following document numbering format throughout the report:  
**Pattern**: {pattern_desc}  
When mentioning any document as evidence, you **MUST** always use its name and number from the table below.  
- If a document number has a prefix like "XXX" or "BLPL" (e.g., "XXX-EMS-F-01"), you MUST replace the prefix with the initials of the organization's name when writing the report.Only do this when document pattern starts with XXX . Dont do this if its just F-X or P-X.
- Dont modify the document number or details randomly.
- This rule is strict and must never be skipped. Under no circumstances should `XXX-` or `BLPL-` remain in any document number in your output.
- The correct initials to use for this organization are: **{org_initials(audit.organizationName)}**.

{prompt_table_md}

---

Here are detailed prompts for each clause to guide your evidence generation:
{stage1_prompt_text}

---

### Audit Details:
- Organization: {audit.organizationName}
- IMS Scope: {audit.scope}
- Address: {audit.address}
- Audit Dates: {audit.startDateOfAuditStage1} to {audit.endDateOfAuditStage1}

### Attendance Sheet:
List of personnel in attendance. Use these names accurately while writing evidence, assigning realistic roles relevant to IMS (e.g., CEO, IMS Manager, QMS/EMS/Compliance Officer, etc.):
{attendance_list_text}

---
### ABSOLUTE FORMATTING RULES (PLAIN TEXT ONLY – STRICT):
- The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols from markdown (- or * as formatting), tables, headings, or any other non-standard formatting.
- Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).
- Write all content in normal sentences using only letters, numerals, and standard punctuation.
- Document names and numbers must be written exactly as provided, without surrounding symbols or formatting.
- For spacing, only use actual line breaks; no markdown or decorative spacing.
- Even if the input contains markdown or symbols, remove them in the output — ensure the output is fully cleaned.
- Any output that contains forbidden formatting is invalid.

**STRICT and REDUNDANT RULES (do NOT break them):**
- For each clause, ONLY mention as evidence the exact documents and document numbers provided in the input for that clause.
- If a clause/question has NO documents given in the input, DO NOT invent, imply, or introduce ANY document forms, names, or numbers—leave out any document mention in your answer for that clause.
- Under NO circumstances should you add, paraphrase, or generate document names/numbers beyond what is provided for that clause.
- **Do NOT attempt to complete or create document numbers based on the pattern. ONLY use the exact document number provided. If a number is not listed, do not use one.**
- If you see a general description with NO specific documents, simply generate evidence without referring to any document at all.
- In short: **Never make up or combine document titles, forms, or numbers. Reference every document listed in the input for the clause, and nothing else.**
- Use document dates as specified in the prompt table. Dont generate dates randomly.

### Instructions for Report Writing:
- You are to ONLY update the 'Document Verification detail with statement of Conformity' field of each item in the input list.
- DO NOT change or remove any keys like 'Cl. NO', 'Description', or 'C/NC/O'.
- If the "C/NC/O" value is "C", rephrase the "Document Verification..." as a professional, positive confirmation that IMS/QMS/EMS requirements for this clause are met, referencing relevant ISO 9001/14001:2015 clauses and appropriate document(s).
- If "C/NC/O" is "NC", document a specific nonconformity—clearly stating what is not conforming, referencing the relevant clause and document(s).
- If "C/NC/O" is "O", rephrase neutrally as an observation, referencing the clause and relevant documents.
- Maintain the input order; do not reformat or change any field except 'Document Verification detail with statement of Conformity'.
- Every item is a dictionary, keep it exactly as-is, only updating the 'Document Verification...' field.
- For entries where the 'Description' field includes multiple questions, provide a detailed, structured answer addressing **each question in order**.
- Insert a blank line between each clause's answer for clarity (i.e. two newlines).
- Responses should align with best practices for ISO 9001/14001 Stage 1 IMS audits, referencing roles and documents naturally.
- If you reference a document, use a plausible date 7–10 months before the audit.
- If 'C/NC/O' or 'Document Verification...' is 'NA', do not fill in or modify the field.

---

### Input:
Here is the list of clause findings. Again, do NOT change the structure—only generate appropriate 'Document Verification detail with statement of Conformity' content.

{json.dumps(batch, indent=2, ensure_ascii=False)}

---

### Output:
Respond ONLY with the list of dictionaries, with updated 'Document Verification detail with statement of Conformity' fields.  
Do not add markdown, commentary, or explanations.  
Ensure each clause's answer is separated by one line (\\n\\n) for clarity.
"""


def choose_document_pattern_stage1(forced_pattern_name=None, date_map=None):
    """
    Randomly select one document-numbering pattern for ISO audit document references.
    Returns:
        pattern_name: 'ims_org', 'ims_only', 'qhse', or 'minimal'
        pattern_description: human-readable summary
        clause_map: dict mapping clause -> list of dicts with 'Document Name', 'Document Number'
        prompt_table: Markdown table as string for prompt
    """
    # --- Pattern definitions (expand for all clauses as needed) ---
    # Pattern 1: Org initials + IMS, e.g. XXX-IMS-F-01
    pattern_1 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-IMS-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "XXX-IMS-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "XXX-IMS-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "XXX-IMS-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "XXX-IMS-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes. and also create an evidence like ( purchase order, invoice or contract to justify the scope"
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "XXX-IMS-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition according to the scope."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "XXX-IMS-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "XXX-IMS-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/IMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "XXX-IMS-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "XXX-IMS-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "XXX-IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "XXX-IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "XXX-IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "XXX-IMS-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "XXX-IMS-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "XXX-IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "XXX-IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "XXX-IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "XXX-IMS-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "XXX-IMS-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "XXX-IMS-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "XXX-IMS-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "XXX-IMS-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "XXX-IMS-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "XXX-IMS-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "XXX-IMS-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "XXX-IMS-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "XXX-IMS-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-IMS-MAN-01",
                "Guidance/Description": "Manual describing the organization's IMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-IMS-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "XXX-IMS-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "XXX-IMS-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "XXX-IMS-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-IMS-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the IMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "XXX-IMS-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-IMS-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to the scope of the company. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "XXX-IMS-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "XXX-IMS-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "XXX-IMS-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "XXX-IMS-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "XXX-IMS-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "XXX-IMS-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "XXX-IMS-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "XXX-IMS-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "XXX-IMS-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "XXX-IMS-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "XXX-IMS-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "XXX-IMS-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "XXX-IMS-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-IMS-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "XXX-IMS-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "XXX-IMS-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "XXX-IMS-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-IMS-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "XXX-IMS-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "XXX-IMS-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "XXX-IMS-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "XXX-IMS-P-17",
                "Guidance/Description": "Defines how IMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "XXX-IMS-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "XXX-IMS-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "XXX-IMS-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "XXX-IMS-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "XXX-IMS-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "XXX-IMS-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "XXX-IMS-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "XXX-IMS-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "XXX-IMS-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "XXX-IMS-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "XXX-IMS-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "XXX-IMS-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "XXX-IMS-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "XXX-IMS-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    # Pattern 2: IMS only (IMS-...)
    pattern_2 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "IMS-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "IMS-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "IMS-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "IMS-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "IMS-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "IMS-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "IMS-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "IMS-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/IMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "IMS-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "IMS-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "IMS-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "IMS-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "IMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "IMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "IMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "IMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "IMS-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "IMS-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "IMS-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "IMS-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "IMS-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "IMS-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "IMS-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "IMS-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "IMS-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "IMS-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "IMS-MAN-01",
                "Guidance/Description": "Manual describing the organization's IMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "IMS-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "IMS-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "IMS-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "IMS-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "IMS-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the IMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "IMS-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "IMS-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "IMS-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "IMS-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "IMS-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "IMS-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "IMS-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "IMS-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "IMS-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "IMS-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "IMS-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "IMS-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "IMS-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "IMS-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "IMS-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-IMS-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "IMS-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "IMS-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "IMS-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "IMS-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "IMS-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "IMS-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "IMS-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "IMS-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "IMS-P-17",
                "Guidance/Description": "Defines how IMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "IMS-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "IMS-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "IMS-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "IMS-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "IMS-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "IMS-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "IMS-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "IMS-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "IMS-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "IMS-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "IMS-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "IMS-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "IMS-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "IMS-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "IMS-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    pattern_3 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "QHSE-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "QHSE-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "QHSE-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "QHSE-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "QHSE-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "QHSE-POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "QHSE-POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "QHSE-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "QHSE-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/IMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "QHSE-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "QHSE-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "QHSE-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "QHSE-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "QHSE-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "QHSE-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "QHSE-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "QHSE-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "QHSE-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "QHSE-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "QHSE-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "QHSE-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "QHSE-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "QHSE-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "QHSE-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "QHSE-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "QHSE-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "QHSE-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "QHSE-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "QHSE-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Manual describing the organization's QHSE.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "QHSE-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "QHSE-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "QHSE-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the QHSE, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "QHSE-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "QHSE-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "QHSE-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "QHSE-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "QHSE-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "QHSE-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "QHSE-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "QHSE-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "QHSE-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "QHSE-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "QHSE-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "QHSE-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "QHSE-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-QHSE-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "QHSE-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "QHSE-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "QHSE-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "QHSE-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "QHSE-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "QHSE-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "QHSE-P-17",
                "Guidance/Description": "Defines how QHSE performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "QHSE-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "QHSE-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "QHSE-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "QHSE-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "QHSE-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "QHSE-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "QHSE-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "QHSE-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "QHSE-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "QHSE-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "QHSE-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "QHSE-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "QHSE-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "QHSE-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    pattern_4 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/IMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Manual describing the organization's IMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the IMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "P-17",
                "Guidance/Description": "Defines how IMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    patterns = [
        ("IMS_org", "Org initials + IMS (XXX-IMS-...)", pattern_1),
        ("IMS_only", "IMS only (IMS-...)", pattern_2),
        ("qhse", "QHSE system (QHSE-...)", pattern_3),
        ("minimal", "Minimal prefix (MAN-01, P-01, etc.)", pattern_4),
    ]
    if forced_pattern_name is not None:
        for pn, pd, cm in patterns:
            if pn == forced_pattern_name:
                pattern_name, pattern_desc, clause_map = pn, pd, cm
                break
        else:
            pattern_name, pattern_desc, clause_map = patterns[0]
    else:
        pattern_name, pattern_desc, clause_map = random.choice(patterns)


    # Generate full markdown table including all columns
    lines = [
        "| Clause | Document Name | Document Number | Document Date | Guidance/Description | Document Owner | Approved By | Stage 1 Prompt |",
        "|--------|---------------|----------------|--------------|----------------------|---------------|-------------|----------------|"
    ]
    for clause, docs in clause_map.items():
        for doc in docs:
            key = f"{doc['Document Name']}|{doc['Document Number']}"
            fixed_date = date_map.get(key, "") if date_map else ""
            lines.append(
                f"| {clause} | {doc['Document Name']} | {doc['Document Number']} | {fixed_date} | "
                f"{doc.get('Guidance/Description', '')} | {doc.get('Document Owner', '')} | "
                f"{doc.get('Approved By', '')} | {doc.get('Stage 1 Prompt', '')} |"
            )
    prompt_table = "\n".join(lines)

    return pattern_name, pattern_desc, clause_map, prompt_table

def ensure_list_of_dicts(text: str) -> list[dict]:
    """
    Attempt to extract a JSON list of dictionaries from the input text.
    Cleans up common LLM artifacts like code fences, markdown, and invalid characters.
    """

    # Step 1: Remove markdown-style code fences like ```json ... ```
    cleaned_text = text.strip()
    cleaned_text = re.sub(r"^```(?:json)?\s*|```$", "", cleaned_text, flags=re.MULTILINE).strip()

    # Step 2: Extract JSON array from within larger text (if exists)
    json_array_match = re.search(r"(\[.*?\])", cleaned_text, re.DOTALL)
    if json_array_match:
        cleaned_text = json_array_match.group(1)

    # Step 3: Remove dangerous control characters (nulls, tabs, etc)
    cleaned_text = re.sub(r"[\x00-\x1F\x7F]", "", cleaned_text)

    # Step 4: Try to parse JSON
    try:
        data = json.loads(cleaned_text)
    except json.JSONDecodeError as e:
        raise ValueError(f"❌ Failed to parse JSON from Mistral response: {e}\nResponse text was:\n{cleaned_text[:500]}")

    # Step 5: Validate structure
    if not isinstance(data, list) or not all(isinstance(item, dict) for item in data):
        raise ValueError("❌ Parsed content is not a list of dictionaries")

    return data

def split_into_batches(data, batch_size=5):
    return [data[i:i + batch_size] for i in range(0, len(data), batch_size)]

def extract_audit_table_iso9001_14001_stage1_ordered(docx_path_or_stream):
    """
    Extracts the integrated ISO 9001+14001 Stage 1 main audit table from the given DOCX stream or file path.
    Returns a list of dicts with keys:
      - 'Cl. NO'
      - 'Description'
      - 'C/NC/O'
      - 'Document Verification detail with statement of Conformity'

    Skips section/group headers and blank or malformed rows.
    """
    doc = Document(docx_path_or_stream)
    data = []

    # Most flexible header search possible
    must_have = [
        ["cl. no", "cl no", "clause"],  # Clause number
        ["description"],  # Description
        ["c/nc/o", "status", "conformity"],  # Status
        ["document verification detail with statement of conformity",
         "statement of conformity", "document verification"]  # Doc verification
    ]

    # Iterate through all tables
    for table in doc.tables:
        header_idx = None
        header_map = {}

        # Locate header: scan each row for the needed columns (order, capitalization can vary)
        for i, row in enumerate(table.rows):
            header_cells = [cell.text.strip().replace('\n', ' ').lower() for cell in row.cells]
            col_indices = []
            for need in must_have:
                found = next((j for j, txt in enumerate(header_cells)
                              if any(h in txt for h in need)), None)
                col_indices.append(found)
            if all(idx is not None for idx in col_indices):
                header_idx = i
                header_map = {
                    "Cl. NO": col_indices[0],
                    "Description": col_indices[1],
                    "C/NC/O": col_indices[2],
                    "Document Verification detail with statement of Conformity": col_indices[3]
                }
                break  # Found the header row!

        # If the table doesn't have all required columns, skip it
        if header_idx is None or not header_map:
            continue

        # Now extract data rows, starting after header
        for row in table.rows[header_idx + 1:]:
            cells = row.cells
            # Defensive: Must have enough cells for all mapped indices
            if any(idx is None or idx >= len(cells) for idx in header_map.values()):
                continue
            vals = {key: cells[header_map[key]].text.strip() for key in header_map}

            # Skip SECTION HEADER or group header rows:
            # Heuristic: if 'Cl. NO' is empty, or all cells are identical & non-blank, it's likely a group/section heading
            if not vals["Cl. NO"]:
                continue
            if all(v == vals["Cl. NO"] for v in vals.values()) and vals["Cl. NO"]:
                continue
            if not any(vals.values()):
                continue

            data.append(vals)
        # For this template, only extract first matching audit requirements table
        if data:
            break

    return data

def mark_na_clauses(extracted_data, na_clauses):
    """
    For each row in extracted_data whose 'Cl. No' matches any clause number in na_clauses,
    set 'C/NC/O' and 'Document Verification detail with statement of Conformity' to 'NA'.

    - extracted_data: list of dicts (with 'Cl. No', etc.)
    - na_clauses: list of clause-strings (e.g. ["5.2 - ...", "8.4.2 & 8.4.3 - ..."])
    """
    if not na_clauses:
        return extracted_data

    # Helper to normalize and extract all clause numbers from a string
    def extract_clauses(clause_str):
        main = clause_str.split(" - ")[0].strip()
        # Handles multiple clauses like "8.4.2 & 8.4.3"
        nums = [p.strip() for p in main.replace("&", ",").split(",")]
        clauses = []
        for n in nums:
            for part in n.split("and"):
                c = part.strip()
                if c:
                    clauses.append(c)
        return clauses

    # Compile a set of all clause numbers that should be set to NA
    na_clauses_set = set()
    for nc in na_clauses:
        na_clauses_set.update(extract_clauses(nc))

    for row in extracted_data:
        cl_no = row.get("Cl. No", "").strip()
        if cl_no in na_clauses_set:
            row["C/NC/O"] = "NA"
            row["Document Verification detail with statement of Conformity"] = "NA"

    return extracted_data

def update_cnc_placeholders_stage1(rows):
    """
    For each row in 'rows', fill 'C/NC/O' if it is blank or set as a placeholder ('{{clause}}', case-insensitive),
    but ONLY if 'Document Verification detail with statement of Conformity' is not 'NA'.
    Distributes:
      - 2 rows as 'NC' (Non-Conformity),
      - 10% as 'O' (Observation, min 1 if there are any 'O's at all),
      - the rest as 'C' (Conformity).
    If 'evidence' field is 'NA' (case-insensitive), 'C/NC/O' is left blank.
    """
    def is_placeholder(val):
        return not val or val.strip().lower() == "{{clause}}"

    key_cnc = "C/NC/O"
    key_evidence = "Document Verification detail with statement of Conformity"

    # Indices of rows to fill
    indices_to_fill = [
        idx for idx, row in enumerate(rows)
        if is_placeholder(row.get(key_cnc))
        and row.get(key_evidence, "").strip().upper() != "NA"
    ]

    total = len(indices_to_fill)
    if total == 0:
        return rows

    # Calculate distribution
    nc_count = min(2, total)
    remaining = total - nc_count
    o_count = max(1, math.ceil(0.1 * total)) if remaining > 0 else 0
    o_count = min(o_count, remaining)
    c_count = remaining - o_count if remaining > 0 else 0

    # Build randomized distribution
    replacements = (["NC"] * nc_count) + (["C"] * c_count) + (["O"] * o_count)
    random.shuffle(replacements)

    # Assign the values
    for i, idx in enumerate(indices_to_fill):
        rows[idx][key_cnc] = replacements[i]

    # Ensure that any row where evidence is 'NA' gets an empty C/NC/O
    for row in rows:
        if row.get(key_evidence, "").strip().upper() == "NA":
            row[key_cnc] = ""

    return rows

def patch_docx_by_row_index_stage1(docx_buffer, audit_rows, table_idx=None, data_start_idx=1, return_table_idx=False):
    """
    Robustly patches an ISO 9001/14001/45001 Stage 1 audit clause table in a DOCX file (4-column).
    """
    from docx import Document
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    col_map, header_row_idx, table_found = {}, None, None

    # Flexible header detection keywords
    col_keywords = {
        "Cl. No": ["cl. no", "cl no", "cl.no", "clause no", "clause"],
        "Description": ["description"],
        "C/NC/O": ["c/nc/o", "status", "conformity"],
        "Document Verification detail with statement of Conformity": ["document verification", "statement of conformity"]
    }

    # 1. Detect header row
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows[:5]):
            header_cells = [cell.text.strip().lower().replace("\n", " ") for cell in row.cells]
            matched = {key: None for key in col_keywords}
            for idx, text in enumerate(header_cells):
                for col_key, keywords in col_keywords.items():
                    if any(kw in text for kw in keywords):
                        matched[col_key] = idx
            if all(v is not None for v in matched.values()):
                col_map = matched
                header_row_idx = r_idx
                table_found = table
                table_idx = t_idx
                break
        if header_row_idx is not None:
            break

    if not table_found:
        raise ValueError("Could not find Stage-1 clause table in DOCX.")

    # 2. Normalise audit_rows keys
    norm_rows = []
    for r in audit_rows:
        nr = {}
        for col_key in col_keywords.keys():
            for k, v in r.items():
                if k.strip().lower() == col_key.strip().lower():
                    nr[col_key] = v
                    break
            nr.setdefault(col_key, "")  # Ensure all keys present
        norm_rows.append(nr)

    # 3. Patch rows in order
    audit_idx = 0
    for trow in table_found.rows[header_row_idx + 1:]:
        if audit_idx >= len(norm_rows):
            break
        if all(col_map[key] < len(trow.cells) for key in col_map):
            cl_no_text = trow.cells[col_map["Cl. No"]].text.strip()
            if cl_no_text or any(trow.cells[col_map[k]].text.strip() for k in col_map):
                for key in col_map:
                    val = str(norm_rows[audit_idx].get(key, "")).replace("{{clause}}", "")
                    trow.cells[col_map[key]].text = val
                audit_idx += 1

    # 4. Cleanup stray placeholders
    for row in table_found.rows:
        for cell in row.cells:
            if "{{clause}}" in cell.text:
                cell.text = cell.text.replace("{{clause}}", "")

    # 5. Warnings
    if audit_idx < len(norm_rows):
        print(f"⚠️ Warning: Not all audit_rows were patched ({audit_idx} of {len(norm_rows)})")
    elif audit_idx > len(norm_rows):
        print(f"⚠️ Warning: More table lines than data rows ({audit_idx} > {len(norm_rows)})")

    # 6. Save back
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return (docx_buffer, table_idx) if return_table_idx else docx_buffer

def extract_minor_nc_rows(rows):
    results = []
    for row in rows:
        # Find C/NC/O or status column - normalize spaces & case
        status = None
        for key in row.keys():
            if key.strip().lower().replace(" ", "") in ("c/nc/o", "c/nc/o.", "status"):
                status = row[key]
                break
        if status is None:
            continue

        # Normalize status value
        status_norm = str(status).strip().upper().replace(" ", "")
        # Accept NC & variants
        if status_norm not in ("NC", "MINORNC") and not status_norm.endswith("NC"):
            continue

        # Find evidence/verification field
        evidence_val = None
        for key in row.keys():
            if "verification" in key.lower() or "conformity" in key.lower():
                evidence_val = row[key]
                break

        if evidence_val and str(evidence_val).strip().upper() == "NA":
            continue

        results.append(row)

    print(f"[DEBUG] extract_minor_nc_rows: found {len(results)} NC rows out of {len(rows)}")
    return results

def build_minor_nc_summary_prompt(nc_rows):
    return f"""
You are summarizing ISO 9001/14001/45001 minor nonconformities for an IMS audit report.

For each input item, write a short, factual summary (maximum 2–3 lines) describing the nonconformity.

**Start each summary with exactly this format**: Clause <clause-number>: <summary text>

Output rules:
- One observation summary per clause
- Plain text only, separated by blank lines
- No JSON, bullets, code fences, or extra commentary
- The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols from markdown (- or * as formatting), tables, headings, or any other non-standard formatting.
- Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).

Input data:
{json.dumps(nc_rows, indent=2, ensure_ascii=False)}

Now return only the plain text summaries, one per clause, separated by blank lines.
"""

def clean_minor_nc_summaries(summary_text):
    text = summary_text.strip().strip("`")  # strip code fences
    text = re.sub(r"^``````$", "", text, flags=re.MULTILINE).strip()

    # Try parsing JSON
    if text.startswith("[") or text.startswith("{"):
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                out = []
                for item in parsed:
                    if isinstance(item, dict) and item.get("summary"):
                        out.append(str(item["summary"]).strip())
                    elif isinstance(item, str):
                        out.append(item.strip())
                if out:
                    return out
            elif isinstance(parsed, dict) and "summary" in parsed:
                return [parsed["summary"].strip()]
        except Exception:
            pass

    # Fallback to plain text lines
    lines = [line.strip(" -•\t") for line in text.splitlines() if line.strip()]
    return lines

def patch_minor_ncs_table(docx_buffer, summaries):
    from docx import Document
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    for table in doc.tables:
        if any("Minor NCs raised:" in cell.text for cell in table.row_cells(0)):
            # Delete all rows except header
            while len(table.rows) > 1:
                tbl = table._tbl
                tbl.remove(tbl.tr_lst[-1])
            # Add summaries
            for summary in summaries:
                row = table.add_row()
                row.cells[0].text = summary
            break
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

async def add_work_process_to_docx_iso9001_14001_mistral(
    docx_buffer,
    company_name,
    scope,
    mistral_url="https://mistral-api-v2.onrender.com/api/mistral"
):
    """
    Calls Mistral for a work process flow based on the company's scope and inserts it
    into the 'Work process' cell in the IMS (ISO 9001+14001) DOCX.

    Output will be ONLY the process chain in the format:
    Step1 --> Step2 --> Step3 ...
    No extra wording, no markdown, no JSON.
    """
    # 1. Prompt for Work process
    work_process_prompt = f"""
You are writing the 'Work process' section of an IMS (ISO 9001/14001) Stage 1 audit report.

Based ONLY on the following company name and scope, list the main operational process steps of this organization in a logical left-to-right sequence.

FORMAT:
Output ONLY the process chain in the form:
Step1 --> Step2 --> Step3 --> ...
Do not include any explanatory text, numbering, bullet points, JSON, markdown, or other formatting.

- Company Name: {company_name}
- IMS Scope: {scope}
    """

    # 2. Call Mistral API
    async with httpx.AsyncClient(timeout=60.0) as client:
        api_response = await client.post(
            mistral_url,
            json={"prompt": work_process_prompt},
            headers={"Content-Type": "application/json"}
        )
        api_response.raise_for_status()
        if api_response.headers.get("content-type", "").startswith("application/json"):
            result = api_response.json()
            brief_string = result.get("response", "") or str(result)
        else:
            brief_string = api_response.text
    brief_string = brief_string.strip()

    # 3. Robust extraction from all weird result formats
    for key in ("process", "work_process", "steps", "chain", "summary"):
        try:
            obj = json.loads(brief_string)
            if isinstance(obj, dict) and key in obj:
                brief_string = obj[key]
                break
        except Exception:
            pass

    # Clean unwanted formatting
    if brief_string.startswith("```"):
        brief_string = brief_string.strip("`").strip()

    # Remove lines that may accidentally mention ISO/certification
    brief_string = "\n".join([
        line for line in brief_string.splitlines()
        if not re.search(r'(ISO ?9|ISO ?1|14001|certifi|standard|compliance)', line, re.I)
    ]).strip()

    # 4. Insert into DOCX in the 'Work process' cell
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    written = False
    for table in doc.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if ("Work process" in cell.text):
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = brief_string
                        written = True
                    else:
                        cell.text = "Work process:\n\n" + brief_string
                        written = True
    if not written:
        print("⚠️ Could not find 'Work process' cell in the DOCX.")

    # 5. Save back to buffer
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

async def add_materials_handled_to_docx_iso9001_14001(
    docx_buffer,
    company_name,
    scope,
    mistral_url="https://mistral-api-v2.onrender.com/api/mistral"
):
    """
    Calls Mistral for a summary of materials handled and consumed, and inserts it into the IMS (ISO 9001+14001) DOCX in the correct cell.
    The output is strictly based on company scope; no mention of ISO, certifications, standards, compliance, JSON formatting, or code blocks.
    """
    # 1. Strict, ISO-agnostic prompt for materials
    materials_prompt = f"""
You are writing a summary for the 'Materials handled and consumed' section of an IMS (ISO 9001/14001) Stage 1 audit report.

Based ONLY on the following company name and scope, write a concise, specific 1–2 sentence summary describing the main materials, substances, or product types this organization commonly handles, uses, or consumes as part of its operations.

DO NOT mention ISO, certifications, standards, quality/environmental compliance, or audit processes in any form.

- Company Name: {company_name}
- IMS Scope: {scope}

Output ONLY the brief text itself — do NOT include code block formatting, preface, or output as JSON. Output only the summary.
    """

    # 2. Call Mistral API
    async with httpx.AsyncClient(timeout=60.0) as client:
        api_response = await client.post(
            mistral_url,
            json={"prompt": materials_prompt},
            headers={"Content-Type": "application/json"}
        )
        api_response.raise_for_status()
        if api_response.headers.get("content-type", "").startswith("application/json"):
            result = api_response.json()
            brief_string = result.get("response", "") or str(result)
        else:
            brief_string = api_response.text
    brief_string = brief_string.strip()

    # 3. Robust extraction from all weird result formats
    for key in ("materials", "summary", "overview"):
        try:
            obj = json.loads(brief_string)
            if isinstance(obj, dict) and key in obj:
                brief_string = obj[key]
                break
        except Exception:
            pass
    # Remove any lingering codeblock or non-text artifacts:
    if brief_string.startswith("``````"):
        brief_string = brief_string.strip("`").strip()
    # Defensive: remove any lines about ISO or "standard"
    import re
    brief_string = "\n".join([
        line for line in brief_string.splitlines()
        if not re.search(r'(ISO ?9|ISO ?1|14001|certifi|standard|compliance)', line, re.I)
    ]).strip()

    # 4. Insert into DOCX
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    written = False
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            for idx, cell in enumerate(row.cells):
                if ("Materials handled and consumed" in cell.text):
                    # Insert into next cell (second column)
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = brief_string
                        written = True
                    else:
                        cell.text = "Materials handled and consumed:\n\n" + brief_string
                        written = True
    if not written:
        print("⚠️ Could not find 'Materials handled and consumed' cell in the DOCX.")
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

async def add_major_equipment_to_docx_iso9001_14001(
    docx_buffer,
    company_name,
    scope,
    mistral_url="https://mistral-api-v2.onrender.com/api/mistral"
):
    """
    Calls Mistral for a summary of major equipment used, and inserts it into the IMS (ISO 9001+14001) DOCX in the correct cell.
    The output is strictly based on company scope; no mention of ISO, certifications, standards, compliance, JSON formatting, or code blocks.
    """
    # 1. Strict, ISO-agnostic prompt for major equipment
    equipment_prompt = f"""
You are writing a summary for the 'Major Equipment used' section of an IMS (ISO 9001/14001) Stage 1 audit report.

Based ONLY on the following company name and scope, write a concise, specific 1–2 sentence summary describing the main equipment, machinery, or types of tools this organization commonly uses as part of its operations.

DO NOT mention ISO, certifications, standards, quality/environmental compliance, or audit processes in any form.

- Company Name: {company_name}
- IMS Scope: {scope}

Output ONLY the brief text itself — do NOT include code block formatting, preface, or output as JSON. Output only the summary.
    """

    # 2. Call Mistral API
    async with httpx.AsyncClient(timeout=60.0) as client:
        api_response = await client.post(
            mistral_url,
            json={"prompt": equipment_prompt},
            headers={"Content-Type": "application/json"}
        )
        api_response.raise_for_status()
        if api_response.headers.get("content-type", "").startswith("application/json"):
            result = api_response.json()
            brief_string = result.get("response", "") or str(result)
        else:
            brief_string = api_response.text
    brief_string = brief_string.strip()

    # 3. Robust extraction from all weird result formats
    for key in ("equipment", "summary", "overview"):
        try:
            obj = json.loads(brief_string)
            if isinstance(obj, dict) and key in obj:
                brief_string = obj[key]
                break
        except Exception:
            pass
    # Remove any lingering codeblock or non-text artifacts:
    if brief_string.startswith("``````"):
        brief_string = brief_string.strip("`").strip()
    # Defensive: remove any lines about ISO or "standard"
    import re
    brief_string = "\n".join([
        line for line in brief_string.splitlines()
        if not re.search(r'(ISO ?9|ISO ?1|14001|certifi|standard|compliance)', line, re.I)
    ]).strip()

    # 4. Insert into DOCX
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    written = False
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            for idx, cell in enumerate(row.cells):
                if ("Major Equipment used" in cell.text):
                    # Insert into next cell (second column)
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = brief_string
                        written = True
                    else:
                        cell.text = "Major Equipment used:\n\n" + brief_string
                        written = True
    if not written:
        print("⚠️ Could not find 'Major Equipment used' cell in the DOCX.")
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

#  ====================== STAGE-2 FUNCTIONS =============================






@router.post("/stage1/submit")
async def submit_iso9001_stage1(audit: ISO9001_14001Stage1Audit, forced_pattern_name=None, date_map=None):
    doc = DocxTemplate("templates/iso9001_14001_stage1.docx")
    context = {
        "organizationName": audit.organizationName,
        "address": audit.address,
        "siteAddress": audit.siteAddress,
        "numberOfEmployees": audit.numberOfEmployees,
        "emailId": audit.emailId,
        "contactPerson": audit.contactPerson,
        "telephoneFax": audit.telephoneFax,
        "scope": audit.scope,
        "riskCategory": audit.riskCategory,
        "iafCode": audit.iafCode,
        "auditTeam": "\n".join(audit.auditTeam),
        "auditManDays": audit.auditManDays,
        "startDateOfAudit": audit.startDateOfAuditStage1,
        "endDateOfAudit": audit.endDateOfAuditStage1,
        "auditMode": audit.auditMode,
        "quotedManDaysAdequate": audit.quotedManDaysAdequate,
        "changeInEmployeeDetail": audit.changeInEmployeeDetail,
        "changeInScope": audit.changeInScope,
        "additionalInformation": audit.additionalInformation,
        "internalAuditFrequency": audit.internalAuditFrequency,
        "dateOfLastInternalAudit": audit.dateOfLastInternalAudit,
        "managementReviewFrequency": audit.managementReviewFrequency,
        "dateOfLastManagementReview": audit.dateOfLastManagementReview,
        "reviewedBy": audit.reviewedBy,
        "dateOfReview": audit.dateOfReview,
        "na_clauses": audit.na_clauses,
        "clientName": audit.clientName,
        "designation": audit.designation,
        "auditorName": audit.auditorName,
        "clause": "{{clause}}",
    }

    doc.render(context)
    extract_buffer = io.BytesIO()
    doc.save(extract_buffer)
    extract_buffer.seek(0)

    extract_buffer = await add_org_brief_to_docx_iso9001_14001(
        extract_buffer,
        company_name=audit.organizationName,
        scope=audit.scope
    )

    extract_buffer = await add_legal_requirements_to_docx_iso9001_14001_mistral(
        extract_buffer,
        address=audit.address,
        scope=audit.scope
    )

    print("na clauses:", audit.na_clauses)
    rows = extract_audit_table_iso9001_14001_stage1_ordered(extract_buffer)
    rows = mark_na_clauses(rows, audit.na_clauses)
    rows = update_cnc_placeholders_stage1(rows)

    # Ensure date_map is always generated if not supplied
    if date_map is None:
        _, _, clause_map, _ = choose_document_pattern_stage1(forced_pattern_name=forced_pattern_name)
        date_map = generate_document_dates(clause_map, audit.startDateOfAuditStage1)

    # Select the correct document pattern
    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage1(
        forced_pattern_name=forced_pattern_name, date_map=date_map
    )

    print(prompt_table)

    batches = split_into_batches(rows, batch_size=5)
    updated_rows = []
    mistral_api_url = "https://mistral-api-v2.onrender.com/api/mistral"
    headers = {"Content-Type": "application/json"}
    MAX_RETRIES = 3

    # Step 4: Send batches to LLM for evidence rephrasing
    for i, batch in enumerate(batches):
        print(f"🔄 Sending batch {i + 1}/{len(batches)}")
        prompt = generate_prompt_for_stage1(
            batch, audit, clause_map, prompt_table, pattern_desc,
        )
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=300.0) as client:
                    response = await client.post(
                        mistral_api_url, json={"prompt": prompt}, headers=headers
                    )
                    response.raise_for_status()
                    rephrased_text = (
                        response.json().get("response", "")
                        if response.headers.get("content-type", "").startswith("application/json")
                        else response.text
                    )
                    batch_result = ensure_list_of_dicts(rephrased_text)
                    # Strip markdown styling
                    for row in batch_result:
                        for key in row:
                            if isinstance(row[key], str):
                                row[key] = remove_markdown_styling(row[key])
                    updated_rows.extend(batch_result)
                    print(f"✅ Batch {i + 1} succeeded on attempt {attempt}")
                    break
            except Exception as e:
                print(f"⚠️ Batch {i + 1}, attempt {attempt} failed: {e}")
                if attempt == MAX_RETRIES:
                    error_msg = f"Max batch retry reached. Batch {i + 1} failed."
                    print(f"❌ {error_msg}")
                    return {"error": error_msg}

    print("✅ All batches completed. Total rows:", len(updated_rows))

    # Step 5: Patch the rephrased results back onto the rendered docx table
    # print(updated_rows)
    patched_buffer = patch_docx_by_row_index_stage1(extract_buffer, updated_rows)

    # ---- MINOR NC Extraction, Summarization, and Table Patch -------
    minor_nc_rows = extract_minor_nc_rows(updated_rows)

    if minor_nc_rows:
        summary_prompt = build_minor_nc_summary_prompt(minor_nc_rows)
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(mistral_api_url, json={"prompt": summary_prompt}, headers=headers)
            resp.raise_for_status()
            summary_text = (
                resp.json().get("response", "")
                if resp.headers.get("content-type", "").startswith("application/json")
                else resp.text
            )
        minor_nc_summaries = clean_minor_nc_summaries(summary_text)
        patched_buffer = patch_minor_ncs_table(patched_buffer, minor_nc_summaries)

    # ---------------------------------------------------------------
    minor_nc_for_stage2 = []
    for summary in minor_nc_summaries:
        m = re.match(r"Clause\s*([\d\.]+)\s*:\s*(.+)", summary)
        if m:
            minor_nc_for_stage2.append({"Cl. No": m.group(1), "summary": m.group(2)})

    stage1_minor_nc_store.clear()  # Remove any existing from previous run
    stage1_minor_nc_store.extend(minor_nc_for_stage2)  # Save new NCs for Stage 2
    print("[DEBUG][Stage1] stage1_minor_nc_store after saving:", stage1_minor_nc_store)
    print("[DEBUG][Stage1] stage1_minor_nc_store length:", len(stage1_minor_nc_store))

    # ---- OBSERVATION Extraction, Summarization, and Table Patch -------
    obs_rows = extract_observation_rows(updated_rows)
    if obs_rows:
        summary_prompt_obs = build_observation_summary_prompt(obs_rows)
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(mistral_api_url, json={"prompt": summary_prompt_obs}, headers=headers)
            resp.raise_for_status()
            summary_text_obs = resp.json().get("response", "") if resp.headers.get("content-type", "").startswith(
                "application/json") else resp.text
        obs_summaries = clean_observation_summaries(summary_text_obs)
        patched_buffer = patch_observations_table(patched_buffer, obs_summaries)
        # Save for Stage 2 transfer
        obs_for_stage2 = []
        for summary in obs_summaries:
            m = re.match(r"(?:Clause\s*)?([\d\.]+)\s*[:\-–]?\s*(.+)", summary, re.I)
            if m:
                obs_for_stage2.append({"Cl. No": m.group(1), "summary": m.group(2)})
            else:
                # Keep even if parsing fails
                obs_for_stage2.append({"Cl. No": "", "summary": summary})
        stage1_observation_store.clear()
        stage1_observation_store.extend(obs_for_stage2)
        print("[DEBUG][Stage1] stage1_observation_store after saving:", stage1_observation_store)
    # ---------------------------------------------------------------

    final_doc_bytes = patched_buffer.getvalue()

    headers = {
        "Content-Disposition": f"attachment; filename={audit.organizationName}_iso9001_14001_stage1_report.docx"
    }

    return Response(
        content=final_doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

@router.post("/stage2/submit")
async def submit_iso9001_stage2(audit: ISO9001_14001Stage2Audit, forced_pattern_name=None, date_map=None):
    doc = DocxTemplate("templates/iso9001_14001_stage2.docx")
    context = {
        "organizationName": audit.organizationName,
        "address": audit.address,
        "siteAddress": audit.siteAddress,
        "numberOfEmployees": audit.numberOfEmployees,
        "emailId": audit.emailId,
        "contactPerson": audit.contactPerson,
        "telephoneFax": audit.telephoneFax,
        "scope": audit.scope,
        "iafCode": audit.iafCode,
        "riskCategory": audit.riskCategory,
        "auditTeam": "\n".join(audit.auditTeam),
        "auditManDays": audit.auditManDays,
        "startDateOfAudit": audit.startDateOfAuditStage2,
        "endDateOfAudit": audit.endDateOfAuditStage2,
        "internalAuditFrequency": audit.internalAuditFrequency,
        "dateOfLastInternalAudit": audit.dateOfLastInternalAudit,
        "managementReviewFrequency": audit.managementReviewFrequency,
        "dateOfLastManagementReview": audit.dateOfLastManagementReview,
        "clientName": audit.clientName,
        "designation": audit.designation,
        "auditorName": audit.auditorName,
        "clause": "{{clause}}",
    }

    doc.render(context)
    extract_buffer = io.BytesIO()
    doc.save(extract_buffer)
    extract_buffer.seek(0)

    extract_buffer = await add_org_brief_to_docx_iso9001_14001(
        extract_buffer,
        company_name=audit.organizationName,
        scope=audit.scope
    )
    print("✅ Brief added success")
    extract_buffer = await add_legal_requirements_to_docx_iso9001_14001_mistral(
        extract_buffer,
        address=audit.address,
        scope=audit.scope
    )
    print("✅ Legal requirements added success")

    extract_buffer = await add_work_process_to_docx_iso9001_14001_mistral(
        extract_buffer,
        company_name=audit.organizationName,
        scope=audit.scope
    )
    print("✅ Work process added successfully")

    extract_buffer = await add_materials_handled_to_docx_iso9001_14001(
        extract_buffer,
        company_name=audit.organizationName,  # Adjust property name to match your data structure
        scope=audit.scope
    )
    print("✅ Materials added success")

    extract_buffer = await add_major_equipment_to_docx_iso9001_14001(
        extract_buffer,
        company_name=audit.organizationName,  # Use the correct attribute from your audit object
        scope=audit.scope
    )
    print("✅ Major equipment added successfully")

    extracted_rows = extract_audit_table_iso9001_14001_stage2(extract_buffer)
    extracted_rows = mark_na_clauses(extracted_rows, audit.na_clauses)
    extracted_rows = update_cnc_placeholders_stage2(extracted_rows)
    print(extracted_rows)

    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage2(forced_pattern_name=forced_pattern_name, date_map=date_map)

    batches = split_into_batches(extracted_rows, batch_size=5)
    updated_rows = []
    mistral_api_url = "https://mistral-api-v2.onrender.com/api/mistral"
    headers = {"Content-Type": "application/json"}
    MAX_RETRIES = 3
    for i, batch in enumerate(batches):
        print(f"🔄 Sending batch {i + 1}/{len(batches)}")
        # Custom prompt for ISO 14001 Stage 2 (EMS audit): Ensure grammar & context fit EMS
        prompt = generate_prompt_for_iso9001_14001_stage2(
            batch, audit, clause_map, prompt_table, pattern_desc,
        )

        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=300.0) as client:
                    response = await client.post(mistral_api_url, json={"prompt": prompt}, headers=headers)
                    response.raise_for_status()

                    rephrased_text = (
                        response.json().get("response", "")
                        if response.headers.get("content-type", "").startswith("application/json")
                        else response.text
                    )
                    batch_result = ensure_list_of_dicts(rephrased_text)
                    # Strip markdown styling
                    for row in batch_result:
                        for key in row:
                            if isinstance(row[key], str):
                                row[key] = remove_markdown_styling(row[key])
                    updated_rows.extend(batch_result)
                    print(f"✅ Batch {i + 1} succeeded on attempt {attempt}")
                    break
            except Exception as e:
                print(f"⚠️ Batch {i + 1}, attempt {attempt} failed: {e}")
                if attempt == MAX_RETRIES:
                    error_msg = f"Max batch retry reached. Batch {i + 1} failed."
                    print(f"❌ {error_msg}")
                    return {"error": error_msg}

    print("✅ All batches completed. Total rows:", len(updated_rows))
    print("LLM Response")
    print(updated_rows)
    patched_buffer = patch_docx_by_row_index_iso9001_14001_stage2(
        extract_buffer,
        updated_rows
    )

    # --- Insert Stage‑1 NC carry-forward & corrective actions ---
    patched_buffer = await transfer_stage1_ncs_to_stage2_doc(
        patched_buffer, audit, mistral_api_url, headers
    )

    patched_buffer = await transfer_stage1_observations_to_stage2_doc(
        patched_buffer, audit, mistral_api_url, headers
    )

    # ---- MINOR NC Extraction, Summarization, and Table Patch -------
    minor_nc_rows = extract_minor_nc_rows(updated_rows)
    if minor_nc_rows:
        summary_prompt = build_minor_nc_summary_prompt(minor_nc_rows)
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(mistral_api_url, json={"prompt": summary_prompt}, headers=headers)
            resp.raise_for_status()
            summary_text = (
                resp.json().get("response", "")
                if resp.headers.get("content-type", "").startswith("application/json")
                else resp.text
            )
        minor_nc_summaries = clean_minor_nc_summaries(summary_text)
        patched_buffer = patch_minor_ncs_table(patched_buffer, minor_nc_summaries)

    # ---------------------------------------------------------------

    # ---- OBSERVATION Extraction, Summarization, and Table Patch -------
    obs_rows = extract_observation_rows(updated_rows)
    if obs_rows:
        summary_prompt_obs = build_observation_summary_prompt(obs_rows)
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(mistral_api_url, json={"prompt": summary_prompt_obs}, headers=headers)
            resp.raise_for_status()
            summary_text_obs = resp.json().get("response", "") if resp.headers.get("content-type", "").startswith(
                "application/json") else resp.text
        obs_summaries = clean_observation_summaries(summary_text_obs)
        patched_buffer = patch_observations_table(patched_buffer, obs_summaries)

    final_doc_bytes = patched_buffer.getvalue()
    headers = {
        "Content-Disposition": f"attachment; filename={audit.organizationName}_iso14001_stage2_report.docx"
    }

    return Response(
        content=final_doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

@router.post("/download-both")
async def download_stage1_and_stage2_reports(payload: CombinedAuditRequest):
    stage1_audit = payload.stage1_audit
    stage2_audit = payload.stage2_audit

    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage2()

    date_map = generate_document_dates(clause_map, stage1_audit.startDateOfAuditStage1)
    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage2(
        forced_pattern_name=pattern_name,
        date_map=date_map
    )

    print("Pattern chosen for both:", pattern_name)

    # Forward that pattern to both generation calls (force them to use the same numbering)
    stage1_response = await submit_iso9001_stage1(stage1_audit, forced_pattern_name=pattern_name, date_map=date_map)
    stage2_response = await submit_iso9001_stage2(stage2_audit, forced_pattern_name=pattern_name, date_map=date_map)

    # Both responses are FastAPI Response objects; get .body!
    stage1_bytes = stage1_response.body
    stage2_bytes = stage2_response.body
    stage1_filename = f"{stage1_audit.organizationName}_iso9001_14001_stage1.docx"
    stage2_filename = f"{stage2_audit.organizationName}_iso9001_14001_stage2.docx"

    # Make and return ZIP
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        zipf.writestr(stage1_filename, stage1_bytes)
        zipf.writestr(stage2_filename, stage2_bytes)

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={
            "Content-Disposition": "attachment"
        },
    )