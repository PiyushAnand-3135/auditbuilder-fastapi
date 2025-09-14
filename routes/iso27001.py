from fastapi import APIRouter
from fastapi import FastAPI, Response, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from typing import List, Optional
from docxtpl import DocxTemplate
from docx import Document
import io
import json
import math
import random
import httpx
import re
import traceback
from fastapi.responses import JSONResponse
from fastapi.responses import StreamingResponse
import ast
import zipfile
import asyncio
from datetime import datetime, timedelta

router = APIRouter()
stage1_minor_nc_store = []
stage1_observation_store = []

# ==================== MODELS ==========================

class ISO27001Stage1Audit(BaseModel):
    organizationName: str
    address: str
    siteAddress: str
    virtualSiteDetails: str
    numberOfEmployees: int
    numberOfShifts: int
    numberOfUsers: int
    numberOfServers: int
    numberOfWorkStations: int
    numberOfDevStaff: int
    numberOfEmployeesOnSite: int
    emailId: EmailStr
    contactPerson: str
    telephoneFax: str
    scope: str
    auditTeam: List[str]
    auditManDays: str
    businessSector: str
    businessSectorRisk: str
    auditMode: str
    ictArrangement: str
    effectivenessIfRemote: str
    anyDeviationFromAuditPlan: str
    anySignificantIssues: str
    identificationOfAuditTeam: str
    anySignificantChange: str
    startDateOfAuditStage1: str
    endDateOfAuditStage1: str
    startDateOfAuditStage2: Optional[str] = None
    endDateOfAuditStage2: Optional[str] = None
    quotedManDaysAdequate: str  # "Yes" / "No"
    changeInEmployeeDetail: str  # "Yes" / "No"
    changeInScope: str  # "Yes" / "No"
    additionalInformation: str
    attendanceSheet: List[str]
    clientName: str
    designation: str
    auditorName: str
    reviewerName: str
    qualityManagerName: str


class ISO27001Stage2Audit(BaseModel):
    organizationName: str
    address: str
    siteAddress: str
    virtualSiteDetails: str
    numberOfEmployees: int
    numberOfShifts: int
    numberOfUsers: int
    numberOfServers: int
    numberOfWorkStations: int
    numberOfDevStaff: int
    numberOfEmployeesOnSite: int
    emailId: EmailStr
    contactPerson: str
    telephoneFax: str
    scope: str
    businessSector: str
    businessSectorRisk: str
    auditMode: str
    ictArrangement: str
    effectivenessIfRemote: str
    anyDeviationFromAuditPlan: str
    anySignificantIssues: str
    identificationOfAuditTeam: str
    anySignificantChange: str
    startDateOfAuditStage1: str   # Stage 1 audit date (needed for continuity)
    endDateOfAuditStage1: str     # Stage 1 audit date (needed for continuity)
    startDateOfAuditStage2: str
    endDateOfAuditStage2: str
    quotedManDaysAdequate: str
    changeInEmployeeDetail: str
    auditTeam: List[str]
    auditManDays: str
    changeInScope: str
    additionalInformation: str
    attendanceSheet: List[str]
    clientName: str
    designation: str
    auditorName: str
    reviewerName: str
    qualityManagerName: str


class CombinedISO27001AuditRequest(BaseModel):
    stage1_audit: ISO27001Stage1Audit
    stage2_audit: ISO27001Stage2Audit


# ===================== FUNCTIONS ========================

async def add_org_brief_to_docx_iso9001_14001(
    docx_buffer,
    company_name,
    scope,
    mistral_url="https://nodeapi.accuratereport.org/api/mistral/",
    max_retries=3,
    backoff_factor=2,
):
    """
    Calls Mistral for an organization brief and inserts it into the ISMS (ISO 9001+14001) DOCX stage 1 cell.
    Only the actual business/activities brief will appear (no JSON, no standards mention).
    """
    # 1. Strict, ISO-agnostic prompt
    brief_prompt = f"""
You are writing the opening organization summary for an ISMS (ISO 9001/14001) Stage 1 audit report.

Based ONLY on the following company name and scope, write a concise, professional 2–3 sentence overview of this organization's main activities, products/services, and business focus.

DO NOT mention ISO, certifications, standards, quality/environmental compliance, or audit processes in any form.

- Company Name: {company_name}
- ISMS Scope: {scope}

Output ONLY the brief text itself — do NOT include code block formatting, preface, or output as JSON. Output only the brief.
    """

    # 2. Call Mistral API with retry logic
    brief_string = ""
    for attempt in range(1, max_retries + 1):
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                api_response = await client.post(
                    mistral_url,
                    json={"prompt": brief_prompt},
                    headers={"Content-Type": "application/json"}
                )
                api_response.raise_for_status()
                if api_response.headers.get("content-type", "").startswith("application/json"):
                    result = api_response.json()
                    brief_string = result.get("response", "") or str(result)
                else:
                    brief_string = api_response.text
            print(f"✅ Mistral API succeeded on attempt {attempt}")
            break  # success → exit retry loop
        except Exception as e:
            if attempt == max_retries:
                raise  # re-raise last exception
            wait_time = backoff_factor ** (attempt - 1)
            print(f"⚠️ Attempt {attempt} failed: {e}. Retrying in {wait_time}s...")
            await asyncio.sleep(wait_time)

    brief_string = brief_string.strip()

    # 3. Robust extraction from all weird result formats
    for key in ("brief", "overview"):
        try:
            obj = json.loads(brief_string)
            if isinstance(obj, dict) and key in obj:
                brief_string = obj[key]
                break
        except Exception:
            pass
    # Remove any lingering codeblock or non-text artifacts:
    if brief_string.startswith("``````"):
        brief_string = brief_string.strip("`").strip()
    # Defensive: remove any lines about ISO or "standard"
    import re
    brief_string = "\n".join([
        line for line in brief_string.splitlines()
        if not re.search(r'(ISO ?9|ISO ?1|14001|certifi|standard|compliance)', line, re.I)
    ]).strip()

    # 4. Insert into DOCX
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    written = False
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            for idx, cell in enumerate(row.cells):
                if ("Brief about the organization" in cell.text) or ("Organization Brief" in cell.text):
                    # Insert into next cell (second column)
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = brief_string
                        written = True
                    else:
                        cell.text = "Brief about the organization:\n\n" + brief_string
                        written = True
    if not written:
        print("⚠️ Could not find 'Brief about the organization' cell in the DOCX.")
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def org_initials(org_name: str) -> str:
    # E.g., "Livpure Limited" → "LL"
    return "".join([w[0].upper() for w in org_name.split() if w and w[0].isalpha()])

def choose_document_pattern_stage1(forced_pattern_name=None, date_map=None):
    """
    Randomly select one document-numbering pattern for ISO audit document references.
    Returns:
        pattern_name: 'ISMS_org', 'ISMS_only', 'qhse', or 'minimal'
        pattern_description: human-readable summary
        clause_map: dict mapping clause -> list of dicts with 'Document Name', 'Document Number'
        prompt_table: Markdown table as string for prompt
    """
    # --- Pattern definitions (expand for all clauses as needed) ---
    # Pattern 1: Org initials + ISMS, e.g. XXX-ISMS-F-01
    pattern_1 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-ISMS-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "XXX-ISMS-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "XXX-ISMS-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "XXX-ISMS-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "XXX-ISMS-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes. and also create an evidence like ( purchase order, invoice or contract to justify the scope"
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "XXX-ISMS-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "XXX-ISMS-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition according to the scope."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "XXX-POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "XXX-ISMS-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "XXX-ISMS-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/ISMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "XXX-ISMS-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "XXX-ISMS-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "XXX-ISMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "XXX-ISMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "XXX-ISMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-ISMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "XXX-ISMS-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "XXX-ISMS-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "XXX-ISMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "XXX-ISMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "XXX-ISMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-ISMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "XXX-ISMS-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "XXX-ISMS-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "XXX-ISMS-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "XXX-ISMS-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "XXX-ISMS-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "XXX-ISMS-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "XXX-ISMS-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "XXX-ISMS-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "XXX-ISMS-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "XXX-ISMS-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-ISMS-MAN-01",
                "Guidance/Description": "Manual describing the organization's ISMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "XXX-ISMS-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "XXX-ISMS-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "XXX-ISMS-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-ISMS-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "XXX-ISMS-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-ISMS-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "XXX-ISMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the ISMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "XXX-ISMS-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "XXX-ISMS-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "XXX-ISMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to the scope of the company. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "XXX-ISMS-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "XXX-ISMS-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "XXX-ISMS-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "XXX-ISMS-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "XXX-ISMS-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "XXX-ISMS-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "XXX-ISMS-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "XXX-ISMS-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "XXX-ISMS-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "XXX-ISMS-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "XXX-ISMS-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "XXX-ISMS-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "XXX-ISMS-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-ISMS-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-ISMS-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "XXX-ISMS-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "XXX-ISMS-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "XXX-ISMS-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "XXX-ISMS-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "XXX-ISMS-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "XXX-ISMS-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "XXX-ISMS-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "XXX-ISMS-P-17",
                "Guidance/Description": "Defines how ISMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "XXX-ISMS-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "XXX-ISMS-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "XXX-ISMS-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "XXX-ISMS-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "XXX-ISMS-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "XXX-ISMS-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "XXX-ISMS-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "XXX-ISMS-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "XXX-ISMS-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "XXX-ISMS-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "XXX-ISMS-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "XXX-ISMS-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "XXX-ISMS-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "XXX-ISMS-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "XXX-ISMS-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    # Pattern 2: ISMS only (ISMS-...)
    pattern_2 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "ISMS-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "ISMS-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "ISMS-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "ISMS-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "ISMS-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "ISMS-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "ISMS-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "ISMS-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "ISMS-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/ISMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "ISMS-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "ISMS-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "ISMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "ISMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "ISMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "ISMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "ISMS-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "ISMS-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "ISMS-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "ISMS-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "ISMS-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "ISMS-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "ISMS-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "ISMS-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "ISMS-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "ISMS-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "ISMS-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "ISMS-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "ISMS-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "ISMS-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "ISMS-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "ISMS-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "ISMS-MAN-01",
                "Guidance/Description": "Manual describing the organization's ISMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "ISMS-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "ISMS-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "ISMS-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "ISMS-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "ISMS-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "ISMS-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "ISMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the ISMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "ISMS-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "ISMS-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "ISMS-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "ISMS-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "ISMS-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "ISMS-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "ISMS-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "ISMS-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "ISMS-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "ISMS-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "ISMS-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "ISMS-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "ISMS-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "ISMS-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "ISMS-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "ISMS-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-ISMS-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "ISMS-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "ISMS-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "ISMS-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "ISMS-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "ISMS-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "ISMS-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "ISMS-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "ISMS-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "ISMS-P-17",
                "Guidance/Description": "Defines how ISMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "ISMS-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "ISMS-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "ISMS-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "ISMS-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "ISMS-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "ISMS-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "ISMS-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "ISMS-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "ISMS-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "ISMS-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "ISMS-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "ISMS-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "ISMS-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "ISMS-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "ISMS-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    pattern_3 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "QHSE-F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "QHSE-F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "QHSE-P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "QHSE-F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "QHSE-PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "QHSE-POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "QHSE-POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "QHSE-P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "QHSE-P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/ISMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "QHSE-P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "QHSE-F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "QHSE-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "QHSE-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "QHSE-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "QHSE-P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "QHSE-F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "QHSE-P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "QHSE-P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "QHSE-F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "QHSE-OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "QHSE-F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "QHSE-F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "QHSE-F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "QHSE-P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "QHSE-F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "QHSE-F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "QHSE-F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "QHSE-F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "QHSE-F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Manual describing the organization's QHSE.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "QHSE-MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "QHSE-P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "QHSE-P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "QHSE-F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the QHSE, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "QHSE-F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "QHSE-P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "QHSE-P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "QHSE-F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "QHSE-F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "QHSE-P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "QHSE-P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "QHSE-F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "QHSE-F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "QHSE-F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "QHSE-P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "QHSE-F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "QHSE-P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "QHSE-F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "XXX-QHSE-F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "QHSE-F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "QHSE-F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "QHSE-P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "QHSE-F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "QHSE-F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "QHSE-P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "QHSE-F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "QHSE-P-17",
                "Guidance/Description": "Defines how QHSE performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "QHSE-P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "QHSE-P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "QHSE-F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "QHSE-F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "QHSE-F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "QHSE-F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "QHSE-F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "QHSE-F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "QHSE-P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "QHSE-F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "QHSE-P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "QHSE-F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "QHSE-F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "QHSE-F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "QHSE-F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    pattern_4 = {
        "4.1": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Describes the organization's integrated management system and its context.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe in detail the organization’s name, nature of business, core activities, industry sector, and the names/designations of its top management from the attendance sheet, in a formal ISO audit style."
                )
            },
            {
                "Document Name": "SWOT Analysis",
                "Document Number": "F-01",
                "Guidance/Description": "Identifies strengths, weaknesses, opportunities, and threats.",
                "Document Owner": "Process Owner",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Provide a completed SWOT analysis form and describe how results influence actions, with a concrete example of a weakness or opportunity addressed."
                )
            },
            {
                "Document Name": "Context of Organization",
                "Document Number": "F-02",
                "Guidance/Description": "Defines external and internal issues relevant to organizational purpose and QMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the documented issues and explain, with examples, where any of these have prompted operational or policy changes."
                )
            },
        ],
        "4.2": [
            {
                "Document Name": "Procedure for Determining Context and Interested Parties",
                "Document Number": "P-01",
                "Guidance/Description": "Process for identifying interested parties and their relevant needs and expectations.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Describe and show evidence that the organization has identified all relevant interested parties and determined their needs and expectations. List at least 4-5 actual interested parties (e.g., customers, regulators, staff, suppliers, contractors) and specific expectations for each. Show how this information is integrated into your management system."
                )
            },
            {
                "Document Name": "List of Interested Parties",
                "Document Number": "F-03",
                "Guidance/Description": "Lists internal and external interested parties with their needs.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current interested parties list, and for 4-5 entries, show evidence that their expectations are tracked and acted on, such as communications, meeting minutes, or actions taken."
                )
            }
        ],
        "4.3": [
            {
                "Document Name": "Scope of the Quality management system",
                "Document Number": "General Description",
                "Guidance/Description": "Defines the boundaries and applicability of the management system.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the documented scope and give real examples showing what is included and excluded; e.g., reference specific departments, locations, or processes."
                )
            }
        ],
        "4.4": [
            {
                "Document Name": "Process Interaction Chart",
                "Document Number": "PIC-01",
                "Guidance/Description": "A diagram showing process interactions and interfaces.",
                "Document Owner": "Process Owner",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the process map/chart and explain with evidence (e.g., training records, cross-functional meetings) how these interactions are communicated and implemented."
                )
            },
            {
                "Document Name": "List of All procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Comprehensive inventory of all active management system procedures.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current version of the procedures list and a tracked change showing a recent update or addition."
                )
            }
        ],
        "5.1": [
            {
                "Document Name": "Leadership-general",
                "Document Number": "General Description",
                "Guidance/Description": "Describes top management’s leadership approach in the QMS.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Give examples of leadership in action (e.g., signed policies, management review participation, communication records) showing how top management demonstrates commitment and also mention the names of person in top management using the attendance sheet that he/she is commited for conformance of the management system"
                )
            },
            {
                "Document Name": "Customer Focus",
                "Document Number": "POL-02",
                "Guidance/Description": "Policy for customer focus, satisfaction, and requirements.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence (customer feedback reviews, complaint logs, actions taken) that customer focus is implemented and monitored throught the policy."
                )
            },
        ],
        "5.2": [
            {
                "Document Name": "Quality, Environment, Health & Safety Policy",
                "Document Number": "POL-02",
                "Guidance/Description": "Signed and communicated QHSE policy document.",
                "Document Owner": "Managing Director",
                "Approved By": "Board",
                "Stage 1 Prompt": (
                    "Provide the current signed QHSE policy; show evidence of how it is communicated and understood at relevant functions and levels."
                )
            }
        ],
        "5.3": [
            {
                "Document Name": "Procedure for Roles, Responsibilities & Authorities",
                "Document Number": "P-02",
                "Guidance/Description": "Defines functional roles, responsibilities, authorities.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide responsibility matrix or job descriptions and show evidence they are followed (e.g. signed documents, role-based process assignments). Also mention a name to justify this using a name from attendance sheet."
                )
            }
        ],
        "5.4": [
            {
                "Document Name": "Procedure for Consultation and participation of Workers",
                "Document Number": "P-03",
                "Guidance/Description": "Process for involving employees in decisions affecting QHSE.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Present meeting minutes, suggestion system outputs, or other records showing worker involvement in QMS/EMS/ISMS activities."
                )
            }
        ],
        "6.1.1": [
            {
                "Document Name": "Procedure for Addressing Risk and Opportunity",
                "Document Number": "P-04",
                "Guidance/Description": "Documents risk and opportunity assessment and handling.",
                "Document Owner": "Risk Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show risk register(s) and evidence that risks/opportunities are regularly identified, assessed, and acted on (e.g., risk reviews, controls implemented). Also, mention 4-5 risks according to the scope of the company and their mitigation plan."
                )
            },
            {
                "Document Name": "Registry of Key Risks & opportunities",
                "Document Number": "F-08",
                "Guidance/Description": "Record of identified risks and opportunities.",
                "Document Owner": "Risk Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current registry and examples of actions taken on identified risks/opportunities."
                )
            },
        ],
        "6.1.2": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.1.3": [
            {
                "Document Name": "Procedure for identification for legal requirements",
                "Document Number": "P-07",
                "Guidance/Description": "Process to identify, access and comply with legal/other requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show the process for legal requirement identification, and current legal register. Mention legal requirement according to the country and scope of the company."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "F-11",
                "Guidance/Description": "Register of legal/other compliance requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide an up-to-date legal register and show evidence of ongoing review/updates."
                )
            }
        ],
        "6.1.4": [
            {
                "Document Name": "Procedure for Environmental Impact Assessment",
                "Document Number": "P-05",
                "Guidance/Description": "How to identify and manage environmental aspects and impacts.",
                "Document Owner": "EHS Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a completed impact assessment, and show how significant impacts are managed in your daily/operational practices. Also, mention two environmental aspect according to the scope."
                )
            },
            {
                "Document Name": "Procedure for Hazard Identification",
                "Document Number": "P-06",
                "Guidance/Description": "Process for hazard identification and risk evaluation.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Show a record of identified hazards and evidence of implemented controls or corrective actions. Also, write any 2 hazards based on the scope of company."
                )
            },
            {
                "Document Name": "Record of Environmental Aspect and Impact Analysis",
                "Document Number": "F-09",
                "Guidance/Description": "Filled record of aspects and impact analyses.",
                "Document Owner": "EHS Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show the current filled record and evidence that it's regularly reviewed and updated."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Actual hazard analysis and risk treatment records.",
                "Document Owner": "Safety Officer",
                "Approved By": "EHS Manager",
                "Stage 1 Prompt": (
                    "Present completed forms with evidence of actions taken as a result (e.g. mitigation implemented)."
                )
            }
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "OBJ-01",
                "Guidance/Description": "Records QHSE objectives, plans, performance indicators.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide current QHSE objectives, records of performance vs objectives, and evidence of action when objectives are not achieved. Write atleast 4 objectives according to the scope."
                )
            },
            {
                "Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
                "Document Number": "F-12",
                "Guidance/Description": "Filled records of objective monitoring/action plans.",
                "Document Owner": "Quality Manager",
                "Approved By": "Top Management",
                "Stage 1 Prompt": (
                    "Show filled action plans and monitoring records, and describe a real corrective action triggered following a missed target."
                )
            }
        ],
        "7.1": [
            {
                "Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
                "Document Number": "F-13",
                "Guidance/Description": "Inventory of major assets and machinery.",
                "Document Owner": "Asset Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide the current asset list and evidence it's maintained and updated regularly; provide an example of how maintenance is scheduled using the list."
                )
            },
            {
                "Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
                "Document Number": "F-42",
                "Guidance/Description": "Schedules and records for maintenance/calibration.",
                "Document Owner": "Maintenance Supervisor",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show this year's plan and proof that maintenance and calibration are performed as scheduled (e.g., completed checklists, certificates)."
                )
            }
        ],
        "7.2": [
            {
                "Document Name": "Procedure for Training & Competenacy",
                "Document Number": "P-08",
                "Guidance/Description": "How to manage and verify employee competency.",
                "Document Owner": "HR Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show training and competence procedure and evidence (training records, competence evaluations) that personnel are competent for roles assigned."
                )
            },
            {
                "Document Name": "Competence Matrix",
                "Document Number": "F-14",
                "Guidance/Description": "Matrix of staff roles, competencies, qualification status.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the current competence matrix and evidence that gaps are tracked and closed with actual data from the past year. Also, mention a name with the role from the attendance sheet."
                )
            },
            {
                "Document Name": "Annual training Calendar",
                "Document Number": "F-15",
                "Guidance/Description": "Planned training events for the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide recent annual training plan and proof of completed sessions."
                )
            },
            {
                "Document Name": "Effecetiveness of Training Provided",
                "Document Number": "F-16",
                "Guidance/Description": "Evaluation of training effectiveness.",
                "Document Owner": "Training Coordinator",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Present completed effectiveness evaluations and corrective actions taken if training outcomes were not met."
                )
            },
            {
                "Document Name": "Annual Training Records",
                "Document Number": "F-17",
                "Guidance/Description": "Records of all training carried out in the year.",
                "Document Owner": "HR Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show signed training attendance records and certificates for at least 4 different trainings."
                )
            },
            {
                "Document Name": "Competence Evaluation",
                "Document Number": "F-18",
                "Guidance/Description": "Evaluation records for individual competence.",
                "Document Owner": "Quality Manager",
                "Approved By": "HR Manager",
                "Stage 1 Prompt": (
                    "Submit actual filled-in evaluation forms and show how incompetence is identified and corrected. Also, mention a name with designation and what training was provided using the names from attendance sheet."
                )
            }
        ],
        "7.3": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Manual describing the organization's ISMS.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show physical/digital copies of the manual and evidence that staff have access and reference it in work."
                )
            }
        ],
        "7.4": [
            {
                "Document Name": "Integrated Management System Manual",
                "Document Number": "MAN-01",
                "Guidance/Description": "Manual includes communication procedures.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Explain how communication requirements from the manual are followed in practice; provide communications sent using the guidance."
                )
            },
            {
                "Document Name": "Procedure for Internal and External Communication",
                "Document Number": "P-09",
                "Guidance/Description": "How the organization manages its internal/external communications.",
                "Document Owner": "Communications Coordinator",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show an example of actual internal and external communication regarding QMS (e.g. safety alerts, customer letters)."
                )
            }
        ],
        "7.5": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "P-09",
                "Guidance/Description": "Document control process explained.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a controlled document with revision history, and evidence that obsolete versions are removed from use."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "F-04",
                "Guidance/Description": "List of all controlled documents.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current master list, mark controlled/uncontrolled copies, and show an example of a document recently added or revised."
                )
            },
            {
                "Document Name": "List of External Origin Documents",
                "Document Number": "F-19",
                "Guidance/Description": "Documents controlled that come from outside the organization.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide examples showing external documents tracked and updated—e.g., a regulation update tracked in the system."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-20",
                "Guidance/Description": "Form for requesting changes to documents.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Provide one completed change request form and show how requests are logged and tracked."
                )
            }
        ],
        "8.1": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                # Not directly specified; usually Operations Manager or Process Owner – fill as per your org chart
                "Approved By": "",
                "Stage 1 Prompt": (
                    "List all core operational procedures maintained under the ISMS, as per your organization's scope. Show how these procedures are controlled and regularly reviewed."
                )
            },
            {
                "Document Name": "Change management Form",
                "Document Number": "F-21",
                "Guidance/Description": "Change management documentation related to operational processes.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show completed change management forms reflecting changes in any operational procedure or process over the last year."
                )
            },
            {
                "Document Name": "Records of Hazard Analysis and Risk Treatement",
                "Document Number": "F-10",
                "Guidance/Description": "Write a prompt about the hazard analysis and risk treatment identified for each operational procedure.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide filled records of hazard analysis and risk treatment performed for major operational activities. Give an example illustrating how results from these records led to implemented controls."
                )
            }
        ],
        "8.2": [
            {
                "Document Name": "Master List of Operational Procedures",
                "Document Number": "F-04",
                "Guidance/Description": "Write a prompt to list of core operation procedures as per the scope.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show the master list of operational procedures with reference to customer requirements. Provide a sample showing the trace from customer requirements to documented procedures."
                )
            },
            {
                "Document Name": "Procedure for Emergency Preparedness",
                "Document Number": "P-10",
                "Guidance/Description": "Write a prompt that emergency evacuation plan verified and found evident.",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show your current emergency preparedness procedure. Provide evidence (e.g., evacuation drill records) that the plan is tested and known by staff."
                )
            }
        ],
        "8.3": [
            {
                "Document Name": "Procedure for Identification of Design Input & Output of the product and services",
                "Document Number": "P-11",
                "Guidance/Description": "Write a prompt that design & development prompt verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Show a sample of design & development documentation for a product or service as per your organization’s scope. Include input/output records and version control."
                )
            },
            {
                "Document Name": "Design Review and Approval Sheet",
                "Document Number": "F-22",
                "Guidance/Description": "Write a prompt that design review and approval verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Provide a completed design review and approval sheet for a recent project, showing signatures/approvals and identified issues. Also, mention a sample of product or service delivered to client as per the scope."
                )
            },
            {
                "Document Name": "Design Progess Sheet",
                "Document Number": "F-23",
                "Guidance/Description": "Write a prompt that design process flow verified (create a sample of any product or service delivered to any client as per the scope).",
                "Document Owner": "",
                "Approved By": "",
                "Stage 1 Prompt": (
                    "Present a filled-in design process flow/progress sheet documenting key stages, milestones, and change management actions for a selected design project.Also, mention a sample of product or service delivered to client as per the scope."
                )
            }
        ],
        "8.4": [
            {
                "Document Name": "Procedure for Selection & Evaluation of Vendors",
                "Document Number": "P-12",
                "Guidance/Description": "Describes selection, approval, and evaluation of suppliers.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide supplier evaluation records showing at least two vendors evaluated with outcomes. Include criteria used for assessment and ongoing monitoring actions."
                )
            },
            {
                "Document Name": "Procedure for Purchasing Management",
                "Document Number": "P-13",
                "Guidance/Description": "Defines the purchasing process and controls.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show sample purchase orders and evidence of implementation of purchasing procedures, including approval and verification steps."
                )
            },
            {
                "Document Name": "Vendor and Sub Contractor Registration Form",
                "Document Number": "F-24",
                "Guidance/Description": "Form used for registering new vendors/subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a completed registration form for a sample supplier, noting evaluation and approval process."
                )
            },
            {
                "Document Name": "List of Approved Vendors and Sub Contractors",
                "Document Number": "F-25",
                "Guidance/Description": "Current list of all approved suppliers and subcontractors.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the list with at least two example suppliers, including approval status and date of last evaluation."
                )
            },
            {
                "Document Name": "Vendor Registration Form",
                "Document Number": "F-26",
                "Guidance/Description": "Form evidencing vendor registration and approval.",
                "Document Owner": "Procurement Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a completed and signed registration form for a vendor from the current year."
                )
            }
        ],
        "8.5.1": [
            {
                "Document Name": "Procedure for Service/Production/Contract",
                "Document Number": "P-14",
                "Guidance/Description": "Describes service, production, and contract controls.",
                "Document Owner": "Operations Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide records (job cards, work instructions, service records) showing controls during product or service provision for a sample client."
                )
            },
            {
                "Document Name": "HSE work Instructions",
                "Document Number": "F-27",
                "Guidance/Description": "Work instructions addressing health, safety, and environment.",
                "Document Owner": "HSE Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Present completed HSE inspection checklists or records, including a recent random safety inspection outcome."
                )
            }
        ],
        "8.5.2": [
            {
                "Document Name": "Procedure for Document and Record Control",
                "Document Number": "P-14",
                "Guidance/Description": "Defines controls for document identification and traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show examples (logs, tags, digital tracking) of how documents or products are identified and traced throughout service or production."
                )
            },
            {
                "Document Name": "Change Management Form",
                "Document Number": "F-28",
                "Guidance/Description": "Form to log and authorize changes to production/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a completed form for a recent change in production or service, detailing the traceability process."
                )
            },
            {
                "Document Name": "Master List of Documents",
                "Document Number": "F-04",
                "Guidance/Description": "Master index of all documents for traceability.",
                "Document Owner": "Document Control Coordinator",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show how the master list supports document traceability, with an annotated example."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-21",
                "Guidance/Description": "Request form for document changes affecting traceability.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a sample of this form including traceability notes and resolution for one recent request."
                )
            }
        ],
        "8.5.3": [
            {
                "Document Name": "List of Item Received",
                "Document Number": "F-29",
                "Guidance/Description": "Log of customer or external provider property received.",
                "Document Owner": "Warehouse Supervisor",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show a filled form evidencing the receipt and safeguarding of property from a customer or supplier."
                )
            }
        ],
        "8.5.4": [
            {
                "Document Name": "Preservation",
                "Document Number": "General Description",
                "Guidance/Description": "Describes measures for preservation of product through production/service lifecycle.",
                "Document Owner": "Operations Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Explain methods for preserving product/service conformity (packaging, storage, labeling), and give an example from a recent job."
                )
            }
        ],
        "8.5.5": [
            {
                "Document Name": "Customer Feedback Analysis Report",
                "Document Number": "F-30",
                "Guidance/Description": "Reports on customer feedback and post-delivery activities.",
                "Document Owner": "Customer Service Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show a recent customer feedback report, delivery note, or post-delivery survey analysis with a completed action."
                )
            }
        ],
        "8.5.6": [
            {
                "Document Name": "Procedure for Change Management",
                "Document Number": "P-15",
                "Guidance/Description": "Describes the process for managing changes affecting product/service.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide a sample record of a product/service change from initial request to implementation for one project/client."
                )
            },
            {
                "Document Name": "Documents Change Request Form",
                "Document Number": "F-21",
                "Guidance/Description": "Form for logging changes as part of the change management process.",
                "Document Owner": "All Staff",
                "Approved By": "Document Control Coordinator",
                "Stage 1 Prompt": (
                    "Show a completed change request relating to a service/product delivered, including status and approvals."
                )
            }
        ],
        "8.6": [
            {
                "Document Name": "Final Inspection Report",
                "Document Number": "F-30",
                "Guidance/Description": "Final inspection record for product/service before release.",
                "Document Owner": "Quality Inspector",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show at least one signed final inspection report for a product or service delivered to a client relevant to your scope."
                )
            }
        ],
        "8.7": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "P-16",
                "Guidance/Description": "Procedure to identify, control, and correct nonconforming outputs.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Provide records/log for 2-3 nonconforming outputs, including actions taken and follow-up verification."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "F-31",
                "Guidance/Description": "Log/register of nonconformities, corrections, and status.",
                "Document Owner": "Quality Manager",
                "Approved By": "Operations Manager",
                "Stage 1 Prompt": (
                    "Show the registry for the last year with two sample entries, their closure or current status, including evidence for action."
                )
            }
        ],
        "9.1": [
            {
                "Document Name": "Procedure for Monitoring & Measurement",
                "Document Number": "P-17",
                "Guidance/Description": "Defines how ISMS performance is measured, analyzed, and evaluated.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show monitoring/measurement plan, filled records, and summaries/results for the last quarter."
                )
            }
        ],
        "9.1.1": [
            {
                "Document Name": "Procedure for Compliance Management",
                "Document Number": "P-18",
                "Guidance/Description": "Procedure for evaluation and management of compliance obligations.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide compliance monitoring records, audits, or status reports showing periodic review."
                )
            }
        ],
        "9.1.2": [
            {
                "Document Name": "Procedure for Identification for Legal Requirements",
                "Document Number": "P-19",
                "Guidance/Description": "Describes how legal and other requirements are identified and complied with.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a recent update or review record for legal requirements, with one example of a regulatory change tracked and addressed."
                )
            },
            {
                "Document Name": "List of all legal documents and legal requirements",
                "Document Number": "F-11",
                "Guidance/Description": "Up-to-date register of all relevant legal requirements.",
                "Document Owner": "Compliance Officer",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a copy of the legal register and highlight the learning/action taken on a new requirement in the last 6 months."
                )
            }
        ],
        "9.1.3": [
            {
                "Document Name": "Data Analysis Record",
                "Document Number": "F-03",
                "Guidance/Description": "Record and analysis/results of monitored data for continual improvement.",
                "Document Owner": "Quality Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show a sample analysis record and explain the actions decided based on this data analysis."
                )
            }
        ],
        "9.2": [
            {
                "Document Name": "Procedure for Internal Audit",
                "Document Number": "F-04",
                "Guidance/Description": "Describes how internal audits are planned, conducted, and followed up.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last two internal audit reports, including audit program and corrections for non-conformities identified."
                )
            },
            {
                "Document Name": "Internal Audit Program",
                "Document Number": "F-32",
                "Guidance/Description": "Schedule/calendar of planned internal audits.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide the current annual audit program including areas covered and assigned auditors."
                )
            },
            {
                "Document Name": "Internal Audit Schedule",
                "Document Number": "F-33",
                "Guidance/Description": "Detailed audit timetable and auditor assignments.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Show the detailed schedule and confirmation/audit notifications sent."
                )
            },
            {
                "Document Name": "Internal Audit Report",
                "Document Number": "F-34",
                "Guidance/Description": "Completed report with findings, recommendations, and corrective action.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Quality Manager",
                "Stage 1 Prompt": (
                    "Provide a recent audit report, and summarize 2-3 nonconformities found, including their closure status and responsible persons."
                )
            }
        ],
        "9.3": [
            {
                "Document Name": "Procedure for Management Review",
                "Document Number": "P-20",
                "Guidance/Description": "Defines the management review process and requirements.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show management review schedule, agenda, and minutes for the most recent meeting, including actions and persons responsible."
                )
            },
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "F-35",
                "Guidance/Description": "Signed minutes from management review meetings.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide the last approved minutes and highlight key outputs, decisions, and assigned actions."
                )
            }
        ],
        "10.2": [
            {
                "Document Name": "Procedure for Management of Non-Conformities and Corrective Actions",
                "Document Number": "P-21",
                "Guidance/Description": "Details how non-conformities are corrected and actions tracked.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide evidence of at least two corrective actions still in progress, along with their status, owner, and planned closure date."
                )
            },
            {
                "Document Name": "Registry and Status Nonconformities and Corrective Actions",
                "Document Number": "F-36",
                "Guidance/Description": "Register/log showing status of all non-conformities and corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show registry with status updates and details for at least two nonconformities (open and closed)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "F-37",
                "Guidance/Description": "Detailed report evidencing closure and verification for each nonconformity.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Provide a sample nonconformity closure report, including root cause, corrections, actions, and verification."
                )
            }
        ],
        "10.3": [
            {
                "Document Name": "Management Review Minutes",
                "Document Number": "F-35",
                "Guidance/Description": "Signed minutes, including continual improvement review and actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Show evidence that continual improvement is reviewed and driven through management review (e.g., improvement actions and tracking)."
                )
            },
            {
                "Document Name": "Non Conformity Report",
                "Document Number": "F-37",
                "Guidance/Description": "Evidence that continual improvement is achieved through corrective actions.",
                "Document Owner": "QHSE Manager",
                "Approved By": "Managing Director",
                "Stage 1 Prompt": (
                    "Submit an example where a nonconformity or suggestion led to a documented improvement of the management system or process."
                )
            }
        ]
        # ...expand for remaining clauses as needed...
    }

    patterns = [
        ("ISMS_org", "Org initials + ISMS (XXX-ISMS-...)", pattern_1),
        ("ISMS_only", "ISMS only (ISMS-...)", pattern_2),
        ("qhse", "QHSE system (QHSE-...)", pattern_3),
        ("minimal", "Minimal prefix (MAN-01, P-01, etc.)", pattern_4),
    ]
    if forced_pattern_name is not None:
        for pn, pd, cm in patterns:
            if pn == forced_pattern_name:
                pattern_name, pattern_desc, clause_map = pn, pd, cm
                break
        else:
            pattern_name, pattern_desc, clause_map = patterns[0]
    else:
        pattern_name, pattern_desc, clause_map = random.choice(patterns)


    # Generate full markdown table including all columns
    lines = [
        "| Clause | Document Name | Document Number | Document Date | Guidance/Description | Document Owner | Approved By | Stage 1 Prompt |",
        "|--------|---------------|----------------|--------------|----------------------|---------------|-------------|----------------|"
    ]
    for clause, docs in clause_map.items():
        for doc in docs:
            key = f"{doc['Document Name']}|{doc['Document Number']}"
            fixed_date = date_map.get(key, "") if date_map else ""
            lines.append(
                f"| {clause} | {doc['Document Name']} | {doc['Document Number']} | {fixed_date} | "
                f"{doc.get('Guidance/Description', '')} | {doc.get('Document Owner', '')} | "
                f"{doc.get('Approved By', '')} | {doc.get('Stage 1 Prompt', '')} |"
            )
    prompt_table = "\n".join(lines)

    return pattern_name, pattern_desc, clause_map, prompt_table

def generate_prompt_for_stage1_iso27001(batch, audit, clause_map, prompt_table_md, pattern_desc):
    # Collect clause-specific prompts if you use them (customize per your clause_map)
    stage1_prompts = []
    for clause, docs in clause_map.items():
        for doc in docs:
            if "Stage 1 Prompt" in doc and doc["Stage 1 Prompt"]:
                stage1_prompts.append(f"Clause {clause}: {doc['Stage 1 Prompt']}")
    stage1_prompt_text = "\n".join(stage1_prompts)

    attendance_list_text = "\n".join([f"- {member}" for member in audit.attendanceSheet])
    org_initials_text = org_initials(audit.organizationName)

    # Extra IMS_org instructions (only if pattern_desc matches IMS_org style)
    ims_org_instructions = ""
    if "Org initials + ISMS (XXX-ISMS-...)" in pattern_desc:
        ims_org_instructions = f"""
        Apply the below pattern only and only if the document number starts with XXX do the below thing
        - If a document number has a prefix like "XXX" or "BLPL" (e.g., "XXX-ISMS-F-01"), you MUST replace the prefix with **{org_initials_text}** (the initials of the organization's name).
        - Do not modify the document number or details randomly.
        - For example, if the organization's name is "{audit.organizationName}", then you must use "{org_initials_text}-ISMS-F-01" instead of "XXX-ISMS-F-01".
        - This rule is strict and must never be skipped. Under no circumstances should `XXX-` or `BLPL-` remain in any document number in your output.
        """

    return f"""
    You are an ISO 27001:2022 Information Security Management System (ISMS) Stage 1 audit reporting assistant.

    Use the following document numbering format throughout all evidence:
    **Pattern**: {pattern_desc}
    {ims_org_instructions}

    {prompt_table_md}

    ---

    ### ABSOLUTE FORMATTING RULES (PLAIN TEXT ONLY – STRICT):
    - The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols from markdown (- or * as formatting), tables, headings, or any other non-standard formatting.
    - Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).
    - Write all content in normal sentences using only letters, numerals, and standard punctuation.
    - Document names and numbers must be written exactly as provided, without surrounding symbols or formatting.
    - For spacing, only use actual line breaks; no markdown or decorative spacing.
    - Even if the input contains markdown or symbols, remove them in the output — ensure the output is fully cleaned.
    - Any output that contains forbidden formatting is invalid.

    **STRICT and REDUNDANT RULES (do NOT break them):**
    - Use the document date for each document as mentioned in the prompt table. Do not generate them randomly.
    - For each clause, ONLY list as evidence the exact documents and their numbers provided for that clause in the input.
    - If a clause has NO listed documents, do NOT mention, imply, or invent any document in the answer for that clause.
    - NEVER create, paraphrase, infer, or generate document names or numbers based on the pattern or clause context.
    - Do NOT combine, split, or otherwise modify listed document names/numbers.
    - **In summary:** ONLY reference documents exactly as listed. DO NOT reference documents for a clause if none are provided.
    - Format every answer strictly as paragraphs, one paragraph for each distinct answer or point. If the content contains any markdown or tables, rewrite them into plain paragraph form while preserving all details and meaning. No table, markdown, or code formatting are allowed in the output.

    ### Audit Details:
    - Organization: {audit.organizationName}
    - ISMS Scope: {audit.scope}
    - Address: {audit.address}
    - Stage 1 Audit Dates: {audit.startDateOfAuditStage1} to {audit.endDateOfAuditStage1}

    ### Attendance Sheet:
    Below is the list of personnel present during the audit. Use these names accurately when drafting evidence. Assign relevant titles/roles (e.g., CEO, IT Manager, ISMS Coordinator, Information Security Officer, etc.) from this list.

    {attendance_list_text}

    ---

    ### Instructions for ISO 27001:2022 Stage 1 ISMS Report Writing:
    - For each clause, the answer must be concise and limited to approximately 80 to 100 words, including only the necessary information relevant to the clause and documents.
    - Only update the 'Document Verification detail with statement of Conformity' field of each input item.
    - Do NOT alter or remove any other fields (e.g., 'Clause Number', 'C/NC/O').
    - For 'C' (Conformity): Rephrase the evidence as a factual, positive confirmation that ISO 27001:2022 requirements for that clause are met, referencing only the clause(s) and any relevant listed document(s).
    - For 'NC' (Nonconformity): Clearly state what does not conform to ISO 27001:2022, referencing only listed clause(s) and document(s).
    - For 'O' (Observation): Reword the evidence as a neutral, factual observation, referencing only the listed clause(s) and document(s).
    - If the 'Clause Number' field includes multiple items, write a structured response that clearly addresses each in order.
    - STRICTLY follow the order of batch items; do NOT change structure or order — modify only the evidence field.
    - Do NOT add, merge, or invent document references under any circumstances. Omit document references if none are listed.
    - Use specific names and roles from the attendance list in your responses as appropriate.
    - Responses must align with ISO 27001:2022 Stage 1 ISMS audit standards wherever the clause applies.
    - Ensure every answer is separated by a blank line (two newlines) for clarity.
    - Output must be only the list of dictionaries, updated as per these rules.
    ---

    ### Input:
    Here is the list of clauses and requirements. Do NOT change structure — edit only the 'Document Verification detail with statement of Conformity' field.

    {json.dumps(batch, indent=2, ensure_ascii=False)}

    ---

    ### Output:
    Respond with ONLY the list of dictionaries, with revised 'Document Verification detail with statement of Conformity' fields.
    Do NOT add markdown, comments, or extra text.
    Separate each answer by a single blank line (\\n\\n) for readability.
    """


def extract_audit_table_iso27001_stage1_ordered(docx_path_or_stream):
    from docx import Document

    doc = Document(docx_path_or_stream)
    data = []
    expected_headers = [
        "Clause Number",
        "C/NC/O",
        "Document Verification detail with statement of Conformity"
    ]

    for table in doc.tables:
        header_idx = None
        # Find header row
        for i, row in enumerate(table.rows):
            header_cells = [cell.text.strip().replace('\n', ' ').replace('\r', ' ') for cell in row.cells]
            found = 0
            for h in expected_headers:
                for cell in header_cells:
                    if h.lower() in cell.lower():
                        found += 1
                        break
            if found == 3:
                header_idx = i
                break

        if header_idx is None:
            continue

        # Extract rows after header
        for row in table.rows[header_idx + 1:]:
            cells = row.cells
            if len(cells) < 3:
                continue
            vals = [cells[i].text.strip() if i < len(cells) else '' for i in range(3)]
            # Skip repeated heading rows
            if vals[0] and vals[0] == vals[1] == vals[2]:
                continue
            if not any(vals):
                continue
            data.append({
                "Clause Number": vals[0],
                "C/NC/O": vals[1],
                "Document Verification detail with statement of Conformity": vals[2],
            })
        break
    return data

def mark_na_clauses_stage1_iso27001(extracted_data, na_clauses):
    if not na_clauses:
        return extracted_data

    # Prepare normalized sets for numbers and descriptions
    na_clause_numbers = set()
    na_full_descriptions = set()
    for c in na_clauses:
        tokens = c.split(" - ", 1)
        if len(tokens) == 2:
            na_clause_numbers.add(tokens[0].strip())
            na_full_descriptions.add(tokens[1].strip())
        else:
            na_full_descriptions.add(c.strip())

    for row in extracted_data:
        clause_number = row.get("Clause Number", "").strip()
        desc = row.get("Document Verification detail with statement of Conformity", "").strip()

        # Try splitting clause_number into number and text (if combined)
        tokens = clause_number.split(" ", 1)
        number_at_start = tokens[0] if tokens and tokens[0][0].isdigit() else None
        description_rest = tokens[1].strip() if len(tokens) > 1 else clause_number

        # Mark NA if either clause number or description matches
        if (number_at_start and number_at_start in na_clause_numbers) or \
           (clause_number in na_full_descriptions) or \
           (desc in na_full_descriptions) or \
           (description_rest in na_full_descriptions):
            row["C/NC/O"] = "NA"
            row["Document Verification detail with statement of Conformity"] = "NA"

    return extracted_data

def update_cnc_placeholders_stage1_iso27001(rows):
    """
    Fill 'C/NC/O' fields that are blank or contain '{{clause}}' (case-insensitive)
    with:
        - 2 NC,
        - 10% O (at least 1 if possible),
        - remainder C.

    Skips rows where 'Document Verification detail with statement of Conformity' is 'NA'.
    """

    def is_placeholder(val: str) -> bool:
        return not val or val.strip().lower() == "{{clause}}"

    key_cnc = "C/NC/O"
    key_evidence = "Document Verification detail with statement of Conformity"

    indices_to_fill = [
        idx for idx, row in enumerate(rows)
        if is_placeholder(row.get(key_cnc, ""))
        and row.get(key_evidence, "").strip().upper() != "NA"
    ]

    total = len(indices_to_fill)
    if total == 0:
        return rows

    nc_count = min(2, total)
    remaining = total - nc_count
    o_count = max(1, math.ceil(0.1 * total)) if remaining > 0 else 0
    o_count = min(o_count, remaining)
    c_count = remaining - o_count if remaining > 0 else 0

    replacements = (["NC"] * nc_count) + (["C"] * c_count) + (["O"] * o_count)
    random.shuffle(replacements)

    for i, idx in enumerate(indices_to_fill):
        rows[idx][key_cnc] = replacements[i]

    # Ensure C/NC/O is cleared if evidence is NA
    for row in rows:
        if row.get(key_evidence, "").strip().upper() == "NA":
            row[key_cnc] = ""

    return rows

def ensure_list_of_dicts(text: str) -> list[dict]:
    """
    Attempt to extract a JSON list of dictionaries from the input text.
    Cleans up common LLM artifacts like code fences, markdown, and invalid characters.
    """

    # Step 1: Remove markdown-style code fences like ```json ... ```
    cleaned_text = text.strip()
    cleaned_text = re.sub(r"^```(?:json)?\s*|```$", "", cleaned_text, flags=re.MULTILINE).strip()

    # Step 2: Extract JSON array from within larger text (if exists)
    json_array_match = re.search(r"(\[.*?\])", cleaned_text, re.DOTALL)
    if json_array_match:
        cleaned_text = json_array_match.group(1)

    # Step 3: Remove dangerous control characters (nulls, tabs, etc)
    cleaned_text = re.sub(r"[\x00-\x1F\x7F]", "", cleaned_text)

    # Step 4: Try to parse JSON
    try:
        data = json.loads(cleaned_text)
    except json.JSONDecodeError as e:
        raise ValueError(f"❌ Failed to parse JSON from Mistral response: {e}\nResponse text was:\n{cleaned_text[:500]}")

    # Step 5: Validate structure
    if not isinstance(data, list) or not all(isinstance(item, dict) for item in data):
        raise ValueError("❌ Parsed content is not a list of dictionaries")

    return data

def split_into_batches(data, batch_size=5):
    return [data[i:i + batch_size] for i in range(0, len(data), batch_size)]

def remove_markdown_styling(text: str) -> str:
    """
    Removes markdown bold/italic markers (*, **) without altering other text.
    This avoids accidentally removing asterisks used for math or other purposes.
    """
    if not isinstance(text, str):
        return text
    # Remove **bold** or *italic* markers
    cleaned = re.sub(r'(\*\*|\*)', '', text)
    return cleaned

def patch_docx_by_row_index_stage1_iso27001(docx_buffer, audit_rows):
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    target_table = None
    data_start_idx = None

    expected_headers = [
        "clause number",
        "c/nc/o",
        "document verification detail with statement of conformity"
    ]

    def normalize(text: str) -> str:
        # lowercase + strip + collapse multiple spaces/newlines
        return " ".join(text.lower().split())

    def is_section_heading(row):
        vals = [cell.text.strip() for cell in row.cells]
        return all(vals) and len(set(vals)) == 1

    # Find the correct table by checking if its header row contains all expected headers
    for table in doc.tables:
        for i, row in enumerate(table.rows[:3]):
            headers = [normalize(cell.text) for cell in row.cells]
            if all(any(exp == h for h in headers) for exp in expected_headers):
                target_table = table
                data_start_idx = i + 1
                break
        if target_table:
            break

    if target_table is None or data_start_idx is None:
        raise ValueError("Could not locate ISO 27001 Stage-1 table in the docx!")

    clause_row_idx = 0
    for row in target_table.rows[data_start_idx:]:
        if is_section_heading(row):
            continue
        if clause_row_idx >= len(audit_rows):
            break

        row.cells[0].text = str(audit_rows[clause_row_idx].get("Clause Number", ""))
        row.cells[1].text = str(audit_rows[clause_row_idx].get("C/NC/O", ""))
        row.cells[2].text = str(audit_rows[clause_row_idx].get("Document Verification detail with statement of Conformity", ""))
        clause_row_idx += 1

    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def extract_minor_nc_rows_iso27001(rows):
    results = []
    for row in rows:
        # Find the C/NC/O or status column
        status = row.get("C/NC/O", None)
        if status is None:
            continue

        # Normalize status value
        status_norm = str(status).strip().upper().replace(" ", "")
        # Accept NC & variants (e.g., "NC", "MINORNC")
        if status_norm not in ("NC", "MINORNC") and not status_norm.endswith("NC"):
            continue

        # Get evidence/verification field
        evidence_val = row.get("Document Verification detail with statement of Conformity", "")

        if evidence_val and str(evidence_val).strip().upper() == "NA":
            continue

        results.append(row)

    print(f"[DEBUG] extract_minor_nc_rows_iso27001: found {len(results)} NC rows out of {len(rows)}")
    return results

def build_minor_nc_summary_prompt_iso27001(nc_rows):
    return f"""
You are summarizing ISO 27001:2022 ISMS minor nonconformities for a Stage-1 audit report.

For each input item, write a short, factual summary (maximum 2–3 lines) describing the nonconformity.

**Start each summary with exactly this format**: Clause <clause-number>: <summary text>

Output rules:
- Write each summary on a new line or as separate short paragraphs.
- DO NOT return JSON, lists, bullet points, or any special formatting — just plain sentences/paragraphs.
- No code fences, no additional notes, no explanations.
- The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols from markdown (- or * as formatting), tables, headings, or any other non-standard formatting.
- Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).
- Write all content in normal sentences using only letters, numerals, and standard punctuation.

Input data:
{json.dumps(nc_rows, indent=2, ensure_ascii=False)}

Now return only the plain text summaries, one per clause, separated by blank lines.
"""

def clean_minor_nc_summaries(summary_text):
    text = summary_text.strip().strip("`")  # strip code fences
    text = re.sub(r"^``````$", "", text, flags=re.MULTILINE).strip()

    # Try parsing JSON
    if text.startswith("[") or text.startswith("{"):
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                out = []
                for item in parsed:
                    if isinstance(item, dict) and item.get("summary"):
                        out.append(str(item["summary"]).strip())
                    elif isinstance(item, str):
                        out.append(item.strip())
                if out:
                    return out
            elif isinstance(parsed, dict) and "summary" in parsed:
                return [parsed["summary"].strip()]
        except Exception:
            pass

    # Fallback to plain text lines
    lines = [line.strip(" -•\t") for line in text.splitlines() if line.strip()]
    return lines

def patch_minor_ncs_table(docx_buffer, summaries):
    from docx import Document
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    for table in doc.tables:
        if any("Minor NCs raised:" in cell.text for cell in table.row_cells(0)):
            # Delete all rows except header
            while len(table.rows) > 1:
                tbl = table._tbl
                tbl.remove(tbl.tr_lst[-1])
            # Add summaries
            for summary in summaries:
                row = table.add_row()
                row.cells[0].text = summary
            break
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def extract_observation_rows_iso27001(rows):
    results = []
    for row in rows:
        # Get status from the known key
        status = row.get("C/NC/O", None)
        if not status:
            continue

        # Normalize
        status_norm = str(status).strip().upper()
        if status_norm != "O":
            continue

        # Skip if evidence is NA
        evidence_val = row.get("Document Verification detail with statement of Conformity", "")
        if evidence_val and str(evidence_val).strip().upper() == "NA":
            continue

        results.append(row)

    print(f"[DEBUG] extract_observation_rows_iso27001: found {len(results)} O rows out of {len(rows)}")
    return results

def build_observation_summary_prompt_iso27001(obs_rows):
    return f"""
You are summarizing ISO 27001:2022 Stage 1 audit Observations for an ISMS audit report.

For each input item, write a short, factual summary (maximum 2–3 lines) describing the observation noted.

**Start each summary exactly like this**: Clause <clause-number>: <summary text>

Output rules:
- One observation summary per clause.
- Plain text only, separated by blank lines.
- No JSON, bullets, code fences, or extra commentary.
- The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols, tables, headings, or any other non-standard formatting.
- Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).
- Write all content in normal sentences using only letters, numerals, and standard punctuation.

Input data:
{json.dumps(obs_rows, indent=2, ensure_ascii=False)}

Now return only the plain text summaries, one per clause, separated by blank lines.
"""

def clean_observation_summaries(summary_text):
    text = summary_text.strip().strip("`")
    text = re.sub(r"^``````$", "", text, flags=re.MULTILINE).strip()
    if text.startswith("[") or text.startswith("{"):
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                return [str(item.get("summary") if isinstance(item, dict) else str(item)).strip() for item in parsed if item]
            elif isinstance(parsed, dict) and "summary" in parsed:
                return [parsed["summary"].strip()]
        except:
            pass
    return [line.strip(" -•\t") for line in text.splitlines() if line.strip()]

def patch_observations_table(docx_buffer, summaries):
    from docx import Document
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    for table in doc.tables:
        if any("Observations raised:" in cell.text for cell in table.row_cells(0)):
            while len(table.rows) > 1:
                tbl = table._tbl
                tbl.remove(tbl.tr_lst[-1])
            for summary in summaries:
                row = table.add_row()
                row.cells[0].text = summary
            break
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def generate_document_dates(clause_map, stage1_start_date_str):
    """
    Generate a fixed random date for each unique document in clause_map,
    between 7–10 months before the Stage 1 start date.
    Returned in YYYY-MM-DD format.
    """
    start_date = datetime.strptime(stage1_start_date_str, "%Y-%m-%d")
    date_map = {}

    for clause, docs in clause_map.items():
        for doc in docs:
            key = f"{doc['Document Name']}|{doc['Document Number']}"
            if key not in date_map:
                # Offset between 7–10 months
                months_offset = random.randint(7, 10)
                days_offset = random.randint(0, 29)
                doc_date = start_date - timedelta(days=(months_offset * 30) + days_offset)
                date_map[key] = doc_date.strftime("%Y-%m-%d")

    return date_map


# ==================== STAGE-2 FUNCTIONS ===============================

def extract_audit_table_iso27001_stage2_ordered(docx_path_or_stream):
    from docx import Document

    doc = Document(docx_path_or_stream)
    data = []
    expected_headers = [
        "Clause Number",
        "C/NC/O",
        "Document Verification detail with statement of Conformity"
    ]

    for table in doc.tables:
        header_idx = None
        # Find header row
        for i, row in enumerate(table.rows):
            header_cells = [cell.text.strip().replace('\n', ' ').replace('\r', ' ')
                            for cell in row.cells]
            found = sum(1 for h in expected_headers
                        if any(h.lower() in cell.lower() for cell in header_cells))
            if found == 3:
                header_idx = i
                break

        if header_idx is None:
            continue

        # ✅ Extract rows in clean format
        for row in table.rows[header_idx + 1:]:
            cells = row.cells
            if len(cells) < 1:
                continue

            clause = cells[0].text.strip()

            # Skip repeated header-like rows
            if clause.lower().startswith("clause number"):
                continue

            # 🚫 Skip merged section headers (e.g. INFORMATION SECURITY CONTROLS, 6.0 People controls, 7.0 Physical controls)
            if clause.isupper() or clause.endswith("controls"):
                continue

            if not clause:
                continue

            data.append({
                "Clause Number": clause,
                "C/NC/O": "{{clause}}",
                "Document Verification detail with statement of Conformity": "Answer according to the prompt",
            })

        break  # ✅ Only the first audit checklist table

    return data



def generate_prompt_for_stage2_iso27001(batch, audit, clause_map, prompt_table_md, pattern_desc):
    # Collect clause-specific prompts if you use them (customize per your clause_map)
    stage1_prompts = []
    for clause, docs in clause_map.items():
        for doc in docs:
            if "Stage 1 Prompt" in doc and doc["Stage 1 Prompt"]:
                stage1_prompts.append(f"Clause {clause}: {doc['Stage 1 Prompt']}")
    stage1_prompt_text = "\n".join(stage1_prompts)

    attendance_list_text = "\n".join([f"- {member}" for member in audit.attendanceSheet])
    org_initials_text = org_initials(audit.organizationName)

    # Extra IMS_org instructions (only if pattern_desc matches IMS_org style)
    ims_org_instructions = ""
    if "Org initials + ISMS (XXX-ISMS-...)" in pattern_desc:
        ims_org_instructions = f"""
        Apply the below pattern only and only if the document number starts with XXX do the below thing
        - If a document number has a prefix like "XXX" or "BLPL" (e.g., "XXX-ISMS-F-01"), you MUST replace the prefix with **{org_initials_text}** (the initials of the organization's name).
        - Do not modify the document number or details randomly.
        - For example, if the organization's name is "{audit.organizationName}", then you must use "{org_initials_text}-ISMS-F-01" instead of "XXX-ISMS-F-01".
        - This rule is strict and must never be skipped. Under no circumstances should `XXX-` or `BLPL-` remain in any document number in your output.
        """

    return f"""
    You are an ISO 27001:2022 Information Security Management System (ISMS) Stage 1 audit reporting assistant.

    Use the following document numbering format throughout all evidence:
    **Pattern**: {pattern_desc}
    {ims_org_instructions}

    {prompt_table_md}

    ---

    ### ABSOLUTE FORMATTING RULES (PLAIN TEXT ONLY – STRICT):
    - The output must be in strict plain text — no markdown, no bold (**), italics (*), underscores (_), bullet symbols from markdown (- or * as formatting), tables, headings, or any other non-standard formatting.
    - Do not generate any special characters used for styling in markdown (such as *, _, `, >, |, ~, #, [], ()).
    - Write all content in normal sentences using only letters, numerals, and standard punctuation.
    - Document names and numbers must be written exactly as provided, without surrounding symbols or formatting.
    - For spacing, only use actual line breaks; no markdown or decorative spacing.
    - Even if the input contains markdown or symbols, remove them in the output — ensure the output is fully cleaned.
    - Any output that contains forbidden formatting is invalid.

    **STRICT and REDUNDANT RULES (do NOT break them):**
    - Use the document date for each document as mentioned in the prompt table. Do not generate them randomly.
    - For each clause, ONLY list as evidence the exact documents and their numbers provided for that clause in the input.
    - If a clause has NO listed documents, do NOT mention, imply, or invent any document in the answer for that clause.
    - NEVER create, paraphrase, infer, or generate document names or numbers based on the pattern or clause context.
    - Do NOT combine, split, or otherwise modify listed document names/numbers.
    - **In summary:** ONLY reference documents exactly as listed. DO NOT reference documents for a clause if none are provided.
    - Format every answer strictly as paragraphs, one paragraph for each distinct answer or point. If the content contains any markdown or tables, rewrite them into plain paragraph form while preserving all details and meaning. No table, markdown, or code formatting are allowed in the output.

    ### Audit Details:
    - Organization: {audit.organizationName}
    - ISMS Scope: {audit.scope}
    - Address: {audit.address}
    - Stage 2 Audit Dates: {audit.startDateOfAuditStage2} to {audit.endDateOfAuditStage1}

    ### Attendance Sheet:
    Below is the list of personnel present during the audit. Use these names accurately when drafting evidence. Assign relevant titles/roles (e.g., CEO, IT Manager, ISMS Coordinator, Information Security Officer, etc.) from this list.

    {attendance_list_text}

    ---

    ### Instructions for ISO 27001:2022 Stage 1 ISMS Report Writing:
    - For each clause, the answer must be concise and limited to approximately 80 to 100 words, including only the necessary information relevant to the clause and documents.
    - Only update the 'Document Verification detail with statement of Conformity' field of each input item.
    - Do NOT alter or remove any other fields (e.g., 'Clause Number', 'C/NC/O').
    - For 'C' (Conformity): Rephrase the evidence as a factual, positive confirmation that ISO 27001:2022 requirements for that clause are met, referencing only the clause(s) and any relevant listed document(s).
    - For 'NC' (Nonconformity): Clearly state what does not conform to ISO 27001:2022, referencing only listed clause(s) and document(s).
    - For 'O' (Observation): Reword the evidence as a neutral, factual observation, referencing only the listed clause(s) and document(s).
    - If the 'Clause Number' field includes multiple items, write a structured response that clearly addresses each in order.
    - STRICTLY follow the order of batch items; do NOT change structure or order — modify only the evidence field.
    - Do NOT add, merge, or invent document references under any circumstances. Omit document references if none are listed.
    - Use specific names and roles from the attendance list in your responses as appropriate.
    - Responses must align with ISO 27001:2022 Stage 1 ISMS audit standards wherever the clause applies.
    - Ensure every answer is separated by a blank line (two newlines) for clarity.
    - Output must be only the list of dictionaries, updated as per these rules.
    ---

    ### Input:
    Here is the list of clauses and requirements. Do NOT change structure — edit only the 'Document Verification detail with statement of Conformity' field.

    {json.dumps(batch, indent=2, ensure_ascii=False)}

    ---

    ### Output:
    Respond with ONLY the list of dictionaries, with revised 'Document Verification detail with statement of Conformity' fields.
    Do NOT add markdown, comments, or extra text.
    Separate each answer by a single blank line (\\n\\n) for readability.
    """

def patch_docx_by_row_index_stage2_iso27001(docx_buffer, audit_rows):
    from docx import Document

    docx_buffer.seek(0)
    doc = Document(docx_buffer)

    header_candidates = {
        "clause number": "Clause Number",
        "c/nc/o": "C/NC/O",
        "document verification detail with statement of conformity": "Document Verification detail with statement of Conformity"
    }

    def normalize(txt):
        return " ".join(txt.lower().replace('\n', ' ').replace('\r', ' ').replace('\t', ' ').split())

    # Find table & header mapping
    target_table, header_row_idx, col_map = None, None, {}
    for table in doc.tables:
        for i, row in enumerate(table.rows[:5]):
            headers = [normalize(cell.text) for cell in row.cells]
            matches = {}
            for key in header_candidates:
                for colidx, hdrtxt in enumerate(headers):
                    if key in hdrtxt:
                        matches[key] = colidx
                        break
            if len(matches) == len(header_candidates):
                target_table, header_row_idx, col_map = table, i, matches
                break
        if target_table:
            break

    if not target_table:
        raise ValueError("Could not locate ISO 27001 table with correct header.")

    data_start_idx = header_row_idx + 1
    clause_row_idx = 0

    for row in target_table.rows[data_start_idx:]:
        if clause_row_idx >= len(audit_rows):
            break

        vals = [cell.text.strip() for cell in row.cells]
        if not any(vals):
            continue

        clause_text = vals[col_map["clause number"]] if col_map["clause number"] < len(vals) else ""

        # 🚫 Skip merged/section rows (e.g. INFORMATION SECURITY CONTROLS, "6.0 People controls")
        if not clause_text or clause_text.isupper() or clause_text.endswith("controls"):
            continue

        audit_row = audit_rows[clause_row_idx]

        # ✅ Only update non-clause columns
        for key, audit_key in header_candidates.items():
            if key == "clause number":
                continue  # don't overwrite clause numbers
            idx = col_map[key]
            if idx < len(row.cells):
                row.cells[idx].text = str(audit_row.get(audit_key, row.cells[idx].text))

        clause_row_idx += 1

    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

async def generate_completed_corrective_actions(stage1_minor_nc_store, scope_text, attendance_text, mistral_api_url, headers):
    """
    Given Stage 1 NCs, scope, and attendance, ask LLM for a past-tense corrective action for each NC.
    Returns a list of strings aligned with the store order.
    """
    actions = []
    max_retries = 3
    base_delay = 2

    for nc in stage1_minor_nc_store:
        clause_no = nc.get("Cl. No", "")
        summary = nc.get("summary", "")
        prompt = (
            f"For ISO clause {clause_no}, scope: {scope_text}, "
            f"attendance: {attendance_text}, "
            f"and the following nonconformity: \"{summary}\", "
            "generate a random, realistic corrective action that HAS ALREADY BEEN IMPLEMENTED "
            "and is now completed. "
            "Write the action in the past tense, e.g., 'The XYZ procedure was revised and all staff were trained accordingly.' "
            "Output only a concise, plain-English paragraph of 2-4 sentences describing what was done, "
            "including any implementation and verification. "
            "Do NOT return JSON, bullet points, lists, or code fences — just the text description."
        )

        for attempt in range(1, max_retries + 1):
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    resp = await client.post(mistral_api_url, json={"prompt": prompt}, headers=headers)
                    resp.raise_for_status()
                    if resp.headers.get("content-type", "").startswith("application/json"):
                        action = resp.json().get("response", "")
                    else:
                        action = resp.text
                    actions.append(action.strip())
                    break  # success, exit retry loop
            except (httpx.RequestError, httpx.HTTPStatusError) as e:
                if attempt == max_retries:
                    actions.append(f"Failed to generate corrective action for NC (Clause {clause_no}): {e}")
                else:
                    delay = base_delay * (2 ** (attempt - 1)) + random.uniform(0, 1)
                    await asyncio.sleep(delay)
    return actions

def clean_corrective_action_text(raw_text: str) -> str:
    """
    Cleans corrective action text from LLM.
    If JSON is detected, extracts and flattens meaningful fields into a sentence.
    Removes markdown code fences.
    """
    if not raw_text:
        return ""

    # Remove markdown/json code fences
    cleaned = re.sub(r"^``````$", "", raw_text.strip(), flags=re.MULTILINE).strip()

    # Try parsing as JSON
    try:
        obj = json.loads(cleaned)
        # If top-level dict and has corrective_action
        if isinstance(obj, dict) and "corrective_action" in obj:
            ca_obj = obj["corrective_action"]
            if isinstance(ca_obj, dict):
                parts = []
                for key in ["action_taken", "implementation", "verification"]:
                    if ca_obj.get(key):
                        parts.append(str(ca_obj[key]).strip())
                return " ".join(parts)
        # If obj itself is a string, return it
        if isinstance(obj, str):
            return obj
    except Exception:
        pass  # if not JSON, just keep original text

    return cleaned

async def transfer_stage1_ncs_to_stage2_doc(patched_buffer, audit, mistral_api_url, headers):
    print(f"[DEBUG] Transferring {len(stage1_minor_nc_store)} Stage-1 NCs to Stage-2")
    if stage1_minor_nc_store:
        attendance_text = ", ".join(audit.attendanceSheet)
        scope_text = audit.scope
        actions = await generate_completed_corrective_actions(
            stage1_minor_nc_store,
            scope_text,
            attendance_text,
            mistral_api_url,
            headers
        )
        combined_entries = []
        for nc, action in zip(stage1_minor_nc_store, actions):
            cleaned_action = clean_corrective_action_text(action)
            combined_entries.append(
                f"Clause {nc.get('Cl. No', '')}\n"
                f"Nonconformity: {nc.get('summary', '')}\n"
                f"Corrective action taken: {cleaned_action}"
            )
    else:
        combined_entries = ["No nonconformities from Stage 1."]

    from docx import Document
    patched_buffer.seek(0)
    doc = Document(patched_buffer)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Non conformities raised in Stage-1:" in cell.text:
                    cell.text = "Non conformities raised in Stage-1:\n\n" + "\n\n".join(combined_entries)
                    break
            else:
                continue
            break
        else:
            continue
        break

    patched_buffer.seek(0)
    patched_buffer.truncate(0)
    doc.save(patched_buffer)
    patched_buffer.seek(0)
    return patched_buffer

async def transfer_stage1_observations_to_stage2_doc(patched_buffer, audit, mistral_api_url, headers):
    print(f"[DEBUG] Transferring {len(stage1_observation_store)} Stage-1 Observations to Stage-2")
    if stage1_observation_store:
        attendance_text = ", ".join(audit.attendanceSheet)
        scope_text = audit.scope
        combined_entries = []
        for obs in stage1_observation_store:
            combined_entries.append(
                f"Clause {obs.get('Cl. No', '')}\nObservation: {obs.get('summary', '')}"
            )
    else:
        combined_entries = ["No observations from Stage 1."]

    from docx import Document
    patched_buffer.seek(0)
    doc = Document(patched_buffer)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Observations raised in Stage-1:" in cell.text:
                    cell.text = "Observations raised in Stage-1:\n\n" + "\n\n".join(combined_entries)
                    break
            else:
                continue
            break
        else:
            continue
        break
    patched_buffer.seek(0)
    patched_buffer.truncate(0)
    doc.save(patched_buffer)
    patched_buffer.seek(0)
    return patched_buffer

# ======================= ROUTES ===================================

@router.post("/stage1/submit")
async def submit_iso27001_stage1(audit : ISO27001Stage1Audit, forced_pattern_name=None, date_map=None):
    doc = DocxTemplate("templates/iso27001_stage1.docx")
    context = {
        "organizationName": audit.organizationName,
        "address": audit.address,
        "siteAddress": audit.siteAddress,
        "virtualSiteDetails": audit.virtualSiteDetails,
        "numberOfEmployees": audit.numberOfEmployees,
        "numberOfShifts": audit.numberOfShifts,
        "numberOfUsers": audit.numberOfUsers,
        "numberOfServers": audit.numberOfServers,
        "numberOfWorkStations": audit.numberOfWorkStations,
        "numberOfDevStaff": audit.numberOfDevStaff,
        "numberOfEmployeesOnSite": audit.numberOfEmployeesOnSite,
        "emailId": audit.emailId,
        "contactPerson": audit.contactPerson,
        "telephoneFax": audit.telephoneFax,
        "scope": audit.scope,
        "businessSector": audit.businessSector,
        "businessSectorRisk": audit.businessSectorRisk,
        "auditMode": audit.auditMode,
        "ictArrangement": audit.ictArrangement,
        "effectivenessIfRemote": audit.effectivenessIfRemote,
        "startDateOfAudit": audit.startDateOfAuditStage1,
        "endDateOfAudit": audit.endDateOfAuditStage1,
        "auditTeam": "\n".join(audit.auditTeam),
        "auditManDays": audit.auditManDays,
        "quotedManDaysAdequate": audit.quotedManDaysAdequate,
        "changeInEmployeeDetail": audit.changeInEmployeeDetail,
        "changeInScope": audit.changeInScope,
        "additionalInformation": audit.additionalInformation,
        "attendanceSheet": "\n".join(audit.attendanceSheet),
        "clientName": audit.clientName,
        "designation": audit.designation,
        "auditorName": audit.auditorName,
        "reviewerName": audit.reviewerName,
        "qualityManagerName": audit.qualityManagerName,
        "clause": "{{clause}}",
    }

    doc.render(context)
    extract_buffer = io.BytesIO()
    doc.save(extract_buffer)
    extract_buffer.seek(0)


    extract_buffer = await add_org_brief_to_docx_iso9001_14001(
        extract_buffer,
        company_name=audit.organizationName,

        scope=audit.scope
    )
    print("✅ Brief added success")

    asyncio.sleep(2)

    rows = extract_audit_table_iso27001_stage1_ordered(extract_buffer)
    rows = mark_na_clauses_stage1_iso27001(rows, getattr(audit, "na_clauses", []))
    rows = update_cnc_placeholders_stage1_iso27001(rows)


    if not forced_pattern_name or not date_map:
        _pattern_name, _, _clause_map, _ = choose_document_pattern_stage1()
        _date_map = generate_document_dates(_clause_map, audit.startDateOfAuditStage1)
    else:
        _pattern_name, _date_map = forced_pattern_name, date_map

    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage1(
        forced_pattern_name=_pattern_name,
        date_map=_date_map
    )

    batches = split_into_batches(rows, batch_size=5)
    updated_rows = []
    mistral_api_url = "https://nodeapi.accuratereport.org/api/mistral/"
    headers = {"Content-Type": "application/json"}
    MAX_RETRIES = 5

    # Step 4: Send batches to LLM for evidence rephrasing
    for i, batch in enumerate(batches):
        print(f"🔄 Sending ISO 27001 Stage 1 batch {i + 1}/{len(batches)}")
        prompt = generate_prompt_for_stage1_iso27001(  # <-- ISO 27001 specific
            batch, audit, clause_map, prompt_table, pattern_desc,
        )
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=300.0) as client:
                    response = await client.post(
                        mistral_api_url, json={"prompt": prompt}, headers=headers
                    )

                    # Log details if server returns error code
                    if response.status_code >= 400:
                        print(f"⚠️ Server returned {response.status_code}")
                        print("Headers:", response.headers)
                        print("Body (truncated):", response.text[:1000])

                    response.raise_for_status()

                    rephrased_text = (
                        response.json().get("response", "")
                        if response.headers.get("content-type", "").startswith("application/json")
                        else response.text
                    )
                    batch_result = ensure_list_of_dicts(rephrased_text)

                    # Strip markdown styling
                    for row in batch_result:
                        for key in row:
                            if isinstance(row[key], str):
                                row[key] = remove_markdown_styling(row[key])

                    updated_rows.extend(batch_result)
                    print(f"✅ ISO 27001 Stage 1 batch {i + 1} succeeded on attempt {attempt}")
                    await asyncio.sleep(2)
                    break

            except httpx.HTTPStatusError as e:
                print(f"❌ ISO 27001 Stage 1 batch {i + 1}, attempt {attempt} failed with HTTP {e.response.status_code}")
                await asyncio.sleep(5)
                print("Response text (truncated):", e.response.text[:1000])
                if attempt == MAX_RETRIES:
                    error_msg = f"Max batch retry reached. ISO 27001 Stage 1 batch {i + 1} failed."
                    print(f"❌ {error_msg}")
                    return {"error": error_msg}

            except Exception as e:
                print(f"❌ ISO 27001 Stage 1 batch {i + 1}, attempt {attempt} failed with error: {e}")
                await asyncio.sleep(5)
                if attempt == MAX_RETRIES:
                    error_msg = f"Max batch retry reached. ISO 27001 Stage 1 batch {i + 1} failed."
                    print(f"❌ {error_msg}")
                    return {"error": error_msg}

    print("✅ All ISO 27001 Stage 1 batches completed. Total rows:", len(updated_rows))
    patched_buffer = patch_docx_by_row_index_stage1_iso27001(extract_buffer, updated_rows)

    # ---- ISO 27001 Stage-1 MINOR NC Extraction, Summarization, and Table Patch -------
    minor_nc_rows = extract_minor_nc_rows_iso27001(updated_rows)
    minor_nc_summaries = []
    MAX_RETRIES = 3  # Robust retry logic for LLM query

    if minor_nc_rows:
        summary_prompt = build_minor_nc_summary_prompt_iso27001(minor_nc_rows)
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    resp = await client.post(
                        mistral_api_url,
                        json={"prompt": summary_prompt},
                        headers=headers
                    )
                    resp.raise_for_status()
                    summary_text = (
                        resp.json().get("response", "")
                        if resp.headers.get("content-type", "").startswith("application/json")
                        else resp.text
                    )
                minor_nc_summaries = clean_minor_nc_summaries(summary_text)
                patched_buffer = patch_minor_ncs_table(patched_buffer, minor_nc_summaries)
                break  # Success!
            except Exception as e:
                print(f"[WARN] ISO 27001 Stage-1 Minor NC batch attempt {attempt} failed: {e}")
                await asyncio.sleep(2)
                if attempt == MAX_RETRIES:
                    print("❌ Max retries reached for ISO 27001 Stage-1 Minor NC summarization. Skipping patching.")
                    minor_nc_summaries = []
                    # Optionally leave patched_buffer unchanged
                else:
                    continue  # Retry
    # -----------------------------------------------------------------
    minor_nc_for_stage2 = []
    for summary in minor_nc_summaries:
        m = re.match(r"Clause\s*([\d\.]+)\s*:\s*(.+)", summary)
        if m:
            minor_nc_for_stage2.append({"Cl. No": m.group(1), "summary": m.group(2)})
    stage1_minor_nc_store.clear()  # Remove any existing from previous run
    stage1_minor_nc_store.extend(minor_nc_for_stage2)  # Save new NCs for Stage 2
    print("[DEBUG][Stage1] stage1_minor_nc_store after saving:", stage1_minor_nc_store)
    print("[DEBUG][Stage1] stage1_minor_nc_store length:", len(stage1_minor_nc_store))

    await asyncio.sleep(3)

    # ---- OBSERVATION Extraction, Summarization, and Table Patch (ISO 27001 Stage 1) -------
    obs_rows = extract_observation_rows_iso27001(updated_rows)
    if obs_rows:
        summary_prompt_obs = build_observation_summary_prompt_iso27001(obs_rows)
        summary_text_obs = None
        MAX_RETRIES = 3
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    resp = await client.post(
                        mistral_api_url, json={"prompt": summary_prompt_obs}, headers=headers
                    )
                    resp.raise_for_status()
                    summary_text_obs = (
                        resp.json().get("response", "")
                        if resp.headers.get("content-type", "").startswith("application/json")
                        else resp.text
                    )
                print(f"✅ ISO 27001 Observation batch succeeded on attempt {attempt}")
                break
            except Exception as e:
                print(f"⚠️ ISO 27001 Observation batch attempt {attempt} failed: {e}")
                if attempt == MAX_RETRIES:
                    print("❌ Max observation batch retry reached. Observation batch failed.")
                    summary_text_obs = ""
        if summary_text_obs:  # Only proceed if we got a response
            obs_summaries = clean_observation_summaries(summary_text_obs)
            patched_buffer = patch_observations_table(patched_buffer, obs_summaries)
            # Save for Stage 2 transfer
            obs_for_stage2 = []
            for summary in obs_summaries:
                m = re.match(r"(?:Clause\s*)?([\d\.]+)\s*[:\-–]?\s*(.+)", summary, re.I)
                if m:
                    obs_for_stage2.append({"Cl. No": m.group(1), "summary": m.group(2)})
                else:
                    # Keep even if parsing fails
                    obs_for_stage2.append({"Cl. No": "", "summary": summary})
            stage1_observation_store.clear()
            stage1_observation_store.extend(obs_for_stage2)
            print("[DEBUG][Stage1] stage1_observation_store after saving:", stage1_observation_store)
    # --------------------------------------------------------------------------------------

    final_doc_bytes = patched_buffer.getvalue()

    headers = {
        "Content-Disposition": f"attachment; filename={audit.organizationName}_iso27001_stage1_report.docx"
    }

    return Response(
        content=final_doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.post("/stage2/submit")
async def submit_iso27001_stage2(audit: ISO27001Stage2Audit, forced_pattern_name=None, date_map=None):
    doc = DocxTemplate("templates/iso27001_stage2.docx")
    context = {
        "organizationName": audit.organizationName,
        "address": audit.address,
        "siteAddress": audit.siteAddress,
        "virtualSiteDetails": audit.virtualSiteDetails,
        "numberOfEmployees": audit.numberOfEmployees,
        "numberOfShifts": audit.numberOfShifts,
        "numberOfUsers": audit.numberOfUsers,
        "numberOfServers": audit.numberOfServers,
        "numberOfWorkStations": audit.numberOfWorkStations,
        "numberOfDevStaff": audit.numberOfDevStaff,
        "numberOfEmployeesOnSite": audit.numberOfEmployeesOnSite,
        "emailId": audit.emailId,
        "contactPerson": audit.contactPerson,
        "telephoneFax": audit.telephoneFax,
        "scope": audit.scope,
        "businessSector": audit.businessSector,
        "businessSectorRisk": audit.businessSectorRisk,
        "auditMode": audit.auditMode,
        "ictArrangement": audit.ictArrangement,
        "effectivenessIfRemote": audit.effectivenessIfRemote,
        "anyDeviationFromAuditPlan": audit.anyDeviationFromAuditPlan,
        "anySignificantIssues": audit.anySignificantIssues,
        "identificationOfAuditTeam": audit.identificationOfAuditTeam,
        "anySignificantChange": audit.anySignificantChange,
        "startDateOfAudit": audit.startDateOfAuditStage2,
        "endDateOfAudit": audit.endDateOfAuditStage2,
        "auditTeam": "\n".join(audit.auditTeam),
        "auditManDays": audit.auditManDays,
        "quotedManDaysAdequate": audit.quotedManDaysAdequate,
        "changeInEmployeeDetail": audit.changeInEmployeeDetail,
        "changeInScope": audit.changeInScope,
        "additionalInformation": audit.additionalInformation,
        "attendanceSheet": "\n".join(audit.attendanceSheet),
        "clientName": audit.clientName,
        "designation": audit.designation,
        "auditorName": audit.auditorName,
        "reviewerName": audit.reviewerName,
        "qualityManagerName": audit.qualityManagerName,
        "clause": "{{clause}}",
    }

    doc.render(context)
    extract_buffer = io.BytesIO()
    doc.save(extract_buffer)
    extract_buffer.seek(0)



    extract_buffer = await add_org_brief_to_docx_iso9001_14001(
        extract_buffer,
        company_name=audit.organizationName,

        scope=audit.scope
    )
    print("✅ Brief added success")

    asyncio.sleep(2)

    rows = extract_audit_table_iso27001_stage2_ordered(extract_buffer)
    rows = mark_na_clauses_stage1_iso27001(rows, getattr(audit, "na_clauses", []))
    rows = update_cnc_placeholders_stage1_iso27001(rows)

    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage1(
        forced_pattern_name=forced_pattern_name,
        date_map=date_map
    )

    batches = split_into_batches(rows, batch_size=8)
    updated_rows = []
    mistral_api_url = "https://nodeapi.accuratereport.org/api/mistral/"
    headers = {"Content-Type": "application/json"}
    MAX_RETRIES = 5

    # Step 4: Send batches to LLM for evidence rephrasing
    for i, batch in enumerate(batches):
        print(f"🔄 Sending ISO 27001 Stage 1 batch {i + 1}/{len(batches)}")
        prompt = generate_prompt_for_stage2_iso27001(  # <-- ISO 27001 specific
            batch, audit, clause_map, prompt_table, pattern_desc,
        )
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=300.0) as client:
                    response = await client.post(
                        mistral_api_url, json={"prompt": prompt}, headers=headers
                    )

                    # Log details if server returns error code
                    if response.status_code >= 400:
                        print(f"⚠️ Server returned {response.status_code}")
                        print("Headers:", response.headers)
                        print("Body (truncated):", response.text[:1000])

                    response.raise_for_status()

                    rephrased_text = (
                        response.json().get("response", "")
                        if response.headers.get("content-type", "").startswith("application/json")
                        else response.text
                    )
                    batch_result = ensure_list_of_dicts(rephrased_text)

                    # Strip markdown styling
                    for row in batch_result:
                        for key in row:
                            if isinstance(row[key], str):
                                row[key] = remove_markdown_styling(row[key])

                    updated_rows.extend(batch_result)
                    print(f"✅ ISO 27001 Stage 2 batch {i + 1} succeeded on attempt {attempt}")
                    await asyncio.sleep(2)
                    break

            except httpx.HTTPStatusError as e:
                print(f"❌ ISO 27001 Stage 2 batch {i + 1}, attempt {attempt} failed with HTTP {e.response.status_code}")
                await asyncio.sleep(5)
                print("Response text (truncated):", e.response.text[:1000])
                if attempt == MAX_RETRIES:
                    error_msg = f"Max batch retry reached. ISO 27001 Stage 1 batch {i + 1} failed."
                    print(f"❌ {error_msg}")
                    return {"error": error_msg}

            except Exception as e:
                print(f"❌ ISO 27001 Stage 2 batch {i + 1}, attempt {attempt} failed with error: {e}")
                await asyncio.sleep(5)
                if attempt == MAX_RETRIES:
                    error_msg = f"Max batch retry reached. ISO 27001 Stage 1 batch {i + 1} failed."
                    print(f"❌ {error_msg}")
                    return {"error": error_msg}

    print("✅ All ISO 27001 Stage 1 batches completed. Total rows:", len(updated_rows))
    # print("Updated rows: ",updated_rows)
    patched_buffer = patch_docx_by_row_index_stage2_iso27001(extract_buffer, updated_rows)

    patched_buffer = await transfer_stage1_ncs_to_stage2_doc(
        patched_buffer, audit, mistral_api_url, headers
    )

    patched_buffer = await transfer_stage1_observations_to_stage2_doc(
        patched_buffer, audit, mistral_api_url, headers
    )

    # ---- ISO 27001 Stage-2 MINOR NC Extraction, Summarization, and Table Patch -------
    minor_nc_rows = extract_minor_nc_rows_iso27001(updated_rows)
    minor_nc_summaries = []
    MAX_RETRIES = 3  # Robust retry logic for LLM query

    if minor_nc_rows:
        summary_prompt = build_minor_nc_summary_prompt_iso27001(minor_nc_rows)
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    resp = await client.post(
                        mistral_api_url,
                        json={"prompt": summary_prompt},
                        headers=headers
                    )
                    resp.raise_for_status()
                    summary_text = (
                        resp.json().get("response", "")
                        if resp.headers.get("content-type", "").startswith("application/json")
                        else resp.text
                    )
                minor_nc_summaries = clean_minor_nc_summaries(summary_text)
                patched_buffer = patch_minor_ncs_table(patched_buffer, minor_nc_summaries)
                break  # Success!
            except Exception as e:
                print(f"[WARN] ISO 27001 Stage-1 Minor NC batch attempt {attempt} failed: {e}")
                await asyncio.sleep(2)
                if attempt == MAX_RETRIES:
                    print("❌ Max retries reached for ISO 27001 Stage-1 Minor NC summarization. Skipping patching.")
                    minor_nc_summaries = []
                    # Optionally leave patched_buffer unchanged
                else:
                    continue  # Retry
    # -----------------------------------------------------------------

    # ---- OBSERVATION Extraction, Summarization, and Table Patch (ISO 27001 Stage 1) -------
    obs_rows = extract_observation_rows_iso27001(updated_rows)
    if obs_rows:
        summary_prompt_obs = build_observation_summary_prompt_iso27001(obs_rows)
        summary_text_obs = None
        MAX_RETRIES = 3
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    resp = await client.post(
                        mistral_api_url, json={"prompt": summary_prompt_obs}, headers=headers
                    )
                    resp.raise_for_status()
                    summary_text_obs = (
                        resp.json().get("response", "")
                        if resp.headers.get("content-type", "").startswith("application/json")
                        else resp.text
                    )
                print(f"✅ ISO 27001 Observation batch succeeded on attempt {attempt}")
                break
            except Exception as e:
                print(f"⚠️ ISO 27001 Observation batch attempt {attempt} failed: {e}")
                if attempt == MAX_RETRIES:
                    print("❌ Max observation batch retry reached. Observation batch failed.")
                    summary_text_obs = ""
        if summary_text_obs:  # Only proceed if we got a response
            obs_summaries = clean_observation_summaries(summary_text_obs)
            patched_buffer = patch_observations_table(patched_buffer, obs_summaries)
    # --------------------------------------------------------------------------------------


    final_doc_bytes = patched_buffer.getvalue()

    headers = {
        "Content-Disposition": f"attachment; filename={audit.organizationName}_iso27001_stage2_report.docx"
    }

    return Response(
        content=final_doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

@router.post("/download-both")
async def download_stage1_and_stage2_reports(payload: CombinedISO27001AuditRequest):
    stage1_audit = payload.stage1_audit
    stage2_audit = payload.stage2_audit

    # Pick YOUR pattern ONCE, using stage2’s function: (It’s fine for IMS blending)
    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage1()

    date_map = generate_document_dates(clause_map, stage1_audit.startDateOfAuditStage1)
    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern_stage1(
        forced_pattern_name=pattern_name,
        date_map=date_map
    )

    print("Pattern chosen for both:", pattern_name)

    # Forward that pattern to both generation calls
    stage1_response = await submit_iso27001_stage1(stage1_audit, forced_pattern_name=pattern_name, date_map=date_map)
    await asyncio.sleep(2)
    stage2_response = await submit_iso27001_stage2(stage2_audit, forced_pattern_name=pattern_name, date_map=date_map)

    # --- Safe extraction of bytes regardless of return type ---
    def extract_bytes(resp):
        if resp is None:
            return b""
        if hasattr(resp, "body"):           # FastAPI Response
            return resp.body
        if hasattr(resp, "content"):        # httpx.Response
            return resp.content
        if isinstance(resp, dict):          # Our code returned dict
            # adjust key if different in your dict
            return resp.get("file") or resp.get("content") or b""
        return bytes(resp) if isinstance(resp, (bytes, bytearray)) else b""

    stage1_bytes = extract_bytes(stage1_response)
    stage2_bytes = extract_bytes(stage2_response)

    stage1_filename = f"{stage1_audit.organizationName}_iso27001_stage1.docx"
    stage2_filename = f"{stage2_audit.organizationName}_iso27001_stage2.docx"

    # Make and return ZIP
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        zipf.writestr(stage1_filename, stage1_bytes)
        zipf.writestr(stage2_filename, stage2_bytes)

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={
            "Content-Disposition": "attachment"
        },
    )



