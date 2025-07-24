from fastapi import FastAPI, Response, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from typing import List, Optional
from docxtpl import DocxTemplate
from docx import Document
import io
import json
import math
import random
import httpx
import re
import traceback
from fastapi.responses import JSONResponse


def extract_audit_table_iso9001_stage2(docx_path_or_stream):
    """
    Extracts the main audit clause table from ISO 9001 Stage 2 DOCX.
    Returns a dict like: "CLAUS E NO. - REQUIREMENTS": "EVIDENCE"
    """
    from docx import Document

    doc = Document(docx_path_or_stream)
    extracted = {}

    for table in doc.tables:
        # Get header row texts (first row after the table header, which may be some empty or merged rows)
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        # Match header: CLAUS E NO. | REQUIREMENTS | C/NC/O | EVIDENCE
        header_match = (
            "CLAUSE NO." in [c.upper() for c in first_row]
            and "REQUIREMENTS" in [c.upper() for c in first_row]
            and "EVIDENCE" in [c.upper() for c in first_row]
        )
        # Try a little fuzzy matching in case of extra spaces
        header_upper = [cell.upper().replace(" ", "") for cell in first_row]
        if not header_match and (
            "CLAUSENO." in header_upper and
            "REQUIREMENTS" in header_upper and
            "EVIDENCE" in header_upper
        ):
            header_match = True

        if header_match:
            # Determine real column indices
            try:
                cl_no_idx = next(i for i, txt in enumerate(first_row) if "CLAUS" in txt.upper())
                req_idx = next(i for i, txt in enumerate(first_row) if "REQUIRE" in txt.upper())
                evid_idx = next(i for i, txt in enumerate(first_row) if "EVIDENCE" in txt.upper())
            except StopIteration:
                continue  # Skip if not found

            # Now extract rows after the header
            for row in table.rows[1:]:
                cells = row.cells
                if len(cells) <= max(cl_no_idx, req_idx, evid_idx):
                    continue
                cl_no = cells[cl_no_idx].text.strip()
                req = cells[req_idx].text.strip()
                evid = cells[evid_idx].text.strip()
                # Only add if clause no and requirements present
                if cl_no or req:
                    key = f"{cl_no} - {req}" if cl_no and req else cl_no or req
                    extracted[key] = evid
            break  # Assume only one main table
    return extracted




# =========================== ISO:9001 STAGE 1 FUNCTIONS START HERE ===================================================================

async def add_org_brief_to_docx(
    docx_buffer,
    company_name,
    scope,
    mistral_url="https://mistral-api-5icm.onrender.com/api/mistral"
):
    """
    Calls Mistral for a company brief and inserts into the proper DOCX cell.
    """
    # 1. Prepare prompt
    brief_prompt = f"""
    You are an ISO 9001 audit assistant.

    Based only on the following company name and ISO 9001 scope, write a concise 2-3 sentence professional overview describing this company's main activities and business focus, suitable for use at the beginning of a stage 1 ISO 9001 audit report.

    Company Name: {company_name}
    Scope: {scope}

    Output ONLY the brief, no explanation, no code block, no labels.
    """

    # 2. Call Mistral API
    async with httpx.AsyncClient(timeout=60.0) as client:
        api_response = await client.post(
            mistral_url,
            json={"prompt": brief_prompt},
            headers={"Content-Type": "application/json"}
        )
        api_response.raise_for_status()
        if api_response.headers.get("content-type", "").startswith("application/json"):
            result = api_response.json()
            # Try to handle "response" field if present (sometimes mistral does this)
            brief_string = result.get("response", "") or str(result)
        else:
            brief_string = api_response.text
    brief_string = brief_string.strip()
    try:
        obj = json.loads(brief_string)
        if isinstance(obj, dict) and "overview" in obj:
            brief_string = obj["overview"]
    except Exception:
        pass

    # 3. Insert the brief into the DOCX
    from docx import Document
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    written = False
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            for idx, cell in enumerate(row.cells):
                if "Brief about the organization" in cell.text:
                    # Put brief in next cell if multi-column
                    if idx + 1 < len(row.cells):
                        row.cells[idx + 1].text = brief_string
                        written = True
                    else:
                        cell.text = "Brief about the organization:\n\n" + brief_string
                        written = True
    # Optionally warn if not found
    if not written:
        print("⚠️ Could not find 'Brief about the organization' cell in the DOCX.")
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

import random

def choose_document_pattern():
    """
    Randomly select one document-numbering pattern for ISO audit document references.
    Returns:
        pattern_name: 'ims_org', 'ims_only', 'qhse', or 'minimal'
        pattern_description: human-readable summary
        clause_map: dict mapping clause -> list of dicts with 'Document Name', 'Document Number'
        prompt_table: Markdown table as string for prompt
    """
    # --- Pattern definitions (expand for all clauses as needed) ---
    # Pattern 1: Org initials + IMS, e.g. XXX-IMS-F-01
    pattern_1 = {
        "4.1": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "XXX-IMS-MAN-01"},
            {"Document Name": "SWOT Analysis", "Document Number": "XXX-IMS-F-01"},
            {"Document Name": "Context of Organization", "Document Number": "XXX-IMS-F-02"},
        ],
        "4.2": [
            {"Document Name": "Procedure for Determining Context and Interested Parties",
             "Document Number": "XXX-IMS-P-01"},
            {"Document Name": "List of Interested Parties", "Document Number": "XXX-IMS-F-03"}
        ],
        "4.3": [
            {"Document Name": "Scope of the Quality management system", "Document Number": "General Description"}
        ],
        "4.4": [
            {"Document Name": "Process Interaction Chart", "Document Number": "XXX-IMS-PIC-01"},
            {"Document Name": "List of All procedures", "Document Number": "XXX-IMS-F-04"}
        ],
        "5.1": [
            {"Document Name": "Leadership-general", "Document Number": "General Description"},
            {"Document Name": "Customer Focus", "Document Number": "XXX-POL-02"}
        ],
        "5.2": [
            {"Document Name": "Quality, Environment, Health & Safety Policy", "Document Number": "XXX-POL-02"}
        ],
        "5.3": [
            {"Document Name": "Procedure for Roles, Responsibilities & Authorities", "Document Number": "XXX-IMS-P-02"}
        ],
        "5.4": [
            {"Document Name": "Procedure for Consultation and participation of Workers",
             "Document Number": "XXX-IMS-P-03"}
        ],
        "6.1.1": [
            {"Document Name": "Procedure for Addressing Risk and Opportunity", "Document Number": "XXX-IMS-P-04"},
            {"Document Name": "Registry of Key Risks & opportunities", "Document Number": "XXX-IMS-F-08"}
        ],
        "6.1.2": [
            {"Document Name": "Procedure for Environmental Impact Assessment", "Document Number": "XXX-IMS-P-05"},
            {"Document Name": "Procedure for Hazard Identification", "Document Number": "XXX-IMS-P-06"},
            {"Document Name": "Record of Environmental Aspect and Impact Analysis", "Document Number": "XXX-IMS-F-09"},
            {"Document Name": "Records of Hazard Analysis and Risk Treatement", "Document Number": "XXX-IMS-F-10"}
        ],
        "6.1.3": [
            {"Document Name": "Procedure for identification for legal requirements", "Document Number": "XXX-IMS-P-07"},
            {"Document Name": "List of all legal documents and legal requirements", "Document Number": "XXX-IMS-F-11"}
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "XXX-IMS-OBJ-01"},
            {"Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
             "Document Number": "XXX-IMS-F-12"}
        ],
        "7.1": [
            {"Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
             "Document Number": "XXX-IMS-F-13"},
            {"Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
             "Document Number": "XXX-IMS-F-42"}
        ],
        "7.2": [
            {"Document Name": "Procedure for Training & Competenacy", "Document Number": "XXX-IMS-P-08"},
            {"Document Name": "Competence Matrix", "Document Number": "XXX-IMS-F-14"},
            {"Document Name": "Annual training Calendar", "Document Number": "XXX-IMS-F-15"},
            {"Document Name": "Effecetiveness of Training Provided", "Document Number": "XXX-IMS-F-16"},
            {"Document Name": "Annual Training Records", "Document Number": "XXX-IMS-F-17"},
            {"Document Name": "Competence Evaluation", "Document Number": "XXX-IMS-F-18"}
        ],
        "7.3": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "XXX-IMS-MAN-01"}
        ],
        "7.4": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "XXX-IMS-MAN-01"},
            {"Document Name": "Procedure for Internal and External Communication", "Document Number": "XXX-IMS-P-09"}
        ],
        "7.5": [
            {"Document Name": "Procedure for Document and Record Control", "Document Number": "XXX-IMS-P-09"},
            {"Document Name": "Master List of Documents", "Document Number": "XXX-IMS-F-04"},
            {"Document Name": "List of External Origin Documents", "Document Number": "XXX-IMS-F-19"},
            {"Document Name": "Documents Change Request Form", "Document Number": "XXX-IMS-F-20"}
        ],

        # ...continue for all process, operation, audit, management review, nonconformance etc.
    }

    # Pattern 2: IMS only (IMS-...)
    pattern_2 = {
        "4.1": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "IMS-MAN-01"},
            {"Document Name": "SWOT Analysis", "Document Number": "IMS-F-01"},
            {"Document Name": "Context of Organization", "Document Number": "IMS-F-02"},
        ],
        "4.2": [
            {"Document Name": "Procedure for Determining Context and Interested Parties",
             "Document Number": "IMS-P-01"},
            {"Document Name": "List of Interested Parties", "Document Number": "IMS-F-03"}
        ],
        "4.3": [
            {"Document Name": "Scope of the Quality management system", "Document Number": "General Description"}
        ],
        "4.4": [
            {"Document Name": "Process Interaction Chart", "Document Number": "IMS-PIC-01"},
            {"Document Name": "List of All procedures", "Document Number": "IMS-F-04"}
        ],
        "5.1": [
            {"Document Name": "Leadership-general", "Document Number": "General Description"},
            {"Document Name": "Customer Focus", "Document Number": "IMS-POL-02"}
        ],
        "5.2": [
            {"Document Name": "Quality, Environment, Health & Safety Policy", "Document Number": "IMS-POL-02"}
        ],
        "5.3": [
            {"Document Name": "Procedure for Roles, Responsibilities & Authorities", "Document Number": "IMS-P-02"}
        ],
        "5.4": [
            {"Document Name": "Procedure for Consultation and participation of Workers", "Document Number": "IMS-P-03"}
        ],
        "6.1.1": [
            {"Document Name": "Procedure for Addressing Risk and Opportunity", "Document Number": "IMS-P-04"},
            {"Document Name": "Registry of Key Risks & opportunities", "Document Number": "IMS-F-08"}
        ],
        "6.1.2": [
            {"Document Name": "Procedure for Environmental Impact Assessment", "Document Number": "IMS-P-05"},
            {"Document Name": "Procedure for Hazard Identification", "Document Number": "IMS-P-06"},
            {"Document Name": "Record of Environmental Aspect and Impact Analysis", "Document Number": "IMS-F-09"},
            {"Document Name": "Records of Hazard Analysis and Risk Treatement", "Document Number": "IMS-F-10"}
        ],
        "6.1.3": [
            {"Document Name": "Procedure for identification for legal requirements", "Document Number": "IMS-P-07"},
            {"Document Name": "List of all legal documents and legal requirements", "Document Number": "IMS-F-11"}
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "IMS-OBJ-01"},
            {"Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
             "Document Number": "IMS-F-12"}
        ],
        "7.1": [
            {"Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
             "Document Number": "IMS-F-13"},
            {"Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
             "Document Number": "IMS-F-42"}
        ],
        "7.2": [
            {"Document Name": "Procedure for Training & Competenacy", "Document Number": "IMS-P-08"},
            {"Document Name": "Competence Matrix", "Document Number": "IMS-F-14"},
            {"Document Name": "Annual training Calendar", "Document Number": "IMS-F-15"},
            {"Document Name": "Effecetiveness of Training Provided", "Document Number": "IMS-F-16"},
            {"Document Name": "Annual Training Records", "Document Number": "IMS-F-17"},
            {"Document Name": "Competence Evaluation", "Document Number": "IMS-F-18"}
        ],
        "7.3": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "IMS-MAN-01"}
        ],
        "7.4": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "IMS-MAN-01"},
            {"Document Name": "Procedure for Internal and External Communication", "Document Number": "IMS-P-09"}
        ],
        "7.5": [
            {"Document Name": "Procedure for Document and Record Control", "Document Number": "IMS-P-09"},
            {"Document Name": "Master List of Documents", "Document Number": "IMS-F-04"},
            {"Document Name": "List of External Origin Documents", "Document Number": "IMS-F-19"},
            {"Document Name": "Documents Change Request Form", "Document Number": "IMS-F-20"}
        ],
        # ...continue for all process, operation, audit, management review, nonconformance etc.
    }

    # Pattern 3: QHSE (QHSE-...)
    pattern_3 = {
        "4.1": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "QHSE-MAN-01"},
            {"Document Name": "SWOT Analysis", "Document Number": "QHSE-F-01"},
            {"Document Name": "Context of Organization", "Document Number": "QHSE-F-02"}
        ],
        "4.2": [
            {"Document Name": "Procedure for Determining Context and Interested Parties",
             "Document Number": "QHSE-P-01"},
            {"Document Name": "List of Interested Parties", "Document Number": "QHSE-F-03"}
        ],
        "4.3": [
            {"Document Name": "Scope of the Quality management system", "Document Number": "General Description"}
        ],
        "4.4": [
            {"Document Name": "Process Interaction Chart", "Document Number": "QHSE-PIC-01"},
            {"Document Name": "List of All procedures", "Document Number": "QHSE-F-04"}
        ],
        "5.1": [
            {"Document Name": "Leadership-general", "Document Number": "General Description"},
            {"Document Name": "Customer Focus", "Document Number": "QHSE-POL-02"}
        ],
        "5.2": [
            {"Document Name": "Quality, Environment, Health & Safety Policy", "Document Number": "QHSE-POL-02"}
        ],
        "5.3": [
            {"Document Name": "Procedure for Roles, Responsibilities & Authorities", "Document Number": "QHSE-P-02"}
        ],
        "5.4": [
            {"Document Name": "Procedure for Consultation and participation of Workers", "Document Number": "QHSE-P-03"}
        ],
        "6.1.1": [
            {"Document Name": "Procedure for Addressing Risk and Opportunity", "Document Number": "QHSE-P-04"},
            {"Document Name": "Registry of Key Risks & opportunities", "Document Number": "QHSE-F-08"}
        ],
        "6.1.2": [
            {"Document Name": "Procedure for Environmental Impact Assessment", "Document Number": "QHSE-P-05"},
            {"Document Name": "Procedure for Hazard Identification", "Document Number": "QHSE-P-06"},
            {"Document Name": "Record of Environmental Aspect and Impact Analysis", "Document Number": "QHSE-F-09"},
            {"Document Name": "Records of Hazard Analysis and Risk Treatement", "Document Number": "QHSE-F-10"}
        ],
        "6.1.3": [
            {"Document Name": "Procedure for identification for legal requirements", "Document Number": "QHSE-P-07"},
            {"Document Name": "List of all legal documents and legal requirements", "Document Number": "QHSE-F-11"}
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "QHSE-OBJ-01"},
            {"Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
             "Document Number": "QHSE-F-12"}
        ],
        "7.1": [
            {"Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
             "Document Number": "QHSE-F-13"},
            {"Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
             "Document Number": "QHSE-F-42"}
        ],
        "7.2": [
            {"Document Name": "Procedure for Training & Competenacy", "Document Number": "QHSE-P-08"},
            {"Document Name": "Competence Matrix", "Document Number": "QHSE-F-14"},
            {"Document Name": "Annual training Calendar", "Document Number": "QHSE-F-15"},
            {"Document Name": "Effecetiveness of Training Provided", "Document Number": "QHSE-F-16"},
            {"Document Name": "Annual Training Records", "Document Number": "QHSE-F-17"},
            {"Document Name": "Competence Evaluation", "Document Number": "QHSE-F-18"}
        ],
        "7.3": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "QHSE-MAN-01"}
        ],
        "7.4": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "QHSE-MAN-01"},
            {"Document Name": "Procedure for Internal and External Communication", "Document Number": "QHSE-P-09"}
        ],
        "7.5": [
            {"Document Name": "Procedure for Document and Record Control", "Document Number": "QHSE-P-09"},
            {"Document Name": "Master List of Documents", "Document Number": "QHSE-F-04"},
            {"Document Name": "List of External Origin Documents", "Document Number": "QHSE-F-19"},
            {"Document Name": "Documents Change Request Form", "Document Number": "QHSE-F-20"}
        ],
        # ...continue for all process, operation, audit, management review, nonconformance etc.
    }

    # Pattern 4: Minimal (MAN-01, F-01, etc.)
    pattern_4 = {
        "4.1": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "MAN-01"},
            {"Document Name": "SWOT Analysis", "Document Number": "F-01"},
            {"Document Name": "Context of Organization", "Document Number": "F-02"}
        ],
        "4.2": [
            {"Document Name": "Procedure for Determining Context and Interested Parties", "Document Number": "P-01"},
            {"Document Name": "List of Interested Parties", "Document Number": "F-03"}
        ],
        "4.3": [
            {"Document Name": "Scope of the Quality management system", "Document Number": "General Description"}
        ],
        "4.4": [
            {"Document Name": "Process Interaction Chart", "Document Number": "PIC-01"},
            {"Document Name": "List of All procedures", "Document Number": "F-04"}
        ],
        "5.1": [
            {"Document Name": "Leadership-general", "Document Number": "General Description"},
            {"Document Name": "Customer Focus", "Document Number": "POL-02"}
        ],
        "5.2": [
            {"Document Name": "Quality, Environment, Health & Safety Policy", "Document Number": "POL-02"}
        ],
        "5.3": [
            {"Document Name": "Procedure for Roles, Responsibilities & Authorities", "Document Number": "P-02"}
        ],
        "5.4": [
            {"Document Name": "Procedure for Consultation and participation of Workers", "Document Number": "P-03"}
        ],
        "6.1.1": [
            {"Document Name": "Procedure for Addressing Risk and Opportunity", "Document Number": "P-04"},
            {"Document Name": "Registry of Key Risks & opportunities", "Document Number": "F-08"}
        ],
        "6.1.2": [
            {"Document Name": "Procedure for Environmental Impact Assessment", "Document Number": "P-05"},
            {"Document Name": "Procedure for Hazard Identification", "Document Number": "P-06"},
            {"Document Name": "Record of Environmental Aspect and Impact Analysis", "Document Number": "F-09"},
            {"Document Name": "Records of Hazard Analysis and Risk Treatement", "Document Number": "F-10"}
        ],
        "6.1.3": [
            {"Document Name": "Procedure for identification for legal requirements", "Document Number": "P-07"},
            {"Document Name": "List of all legal documents and legal requirements", "Document Number": "F-11"}
        ],
        "6.2": [
            {
                "Document Name": "Quality & HSE Objectives, Quality & HSE objective monitoring sheets, Results of the Quality Objectives",
                "Document Number": "OBJ-01"},
            {"Document Name": "Objective Monitoring Action Plan and Results of Monitored Data",
             "Document Number": "F-12"}
        ],
        "7.1": [
            {"Document Name": "List of Machinery, List of Computers, List of Assets, List of equipments",
             "Document Number": "F-13"},
            {"Document Name": "Annual maintainance plan and calibration plan for machines and equipments",
             "Document Number": "F-42"}
        ],
        "7.2": [
            {"Document Name": "Procedure for Training & Competenacy", "Document Number": "P-08"},
            {"Document Name": "Competence Matrix", "Document Number": "F-14"},
            {"Document Name": "Annual training Calendar", "Document Number": "F-15"},
            {"Document Name": "Effecetiveness of Training Provided", "Document Number": "F-16"},
            {"Document Name": "Annual Training Records", "Document Number": "F-17"},
            {"Document Name": "Competence Evaluation", "Document Number": "F-18"}
        ],
        "7.3": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "MAN-01"}
        ],
        "7.4": [
            {"Document Name": "Integrated Management System Manual", "Document Number": "MAN-01"},
            {"Document Name": "Procedure for Internal and External Communication", "Document Number": "P-09"}
        ],
        "7.5": [
            {"Document Name": "Procedure for Document and Record Control", "Document Number": "P-09"},
            {"Document Name": "Master List of Documents", "Document Number": "F-04"},
            {"Document Name": "List of External Origin Documents", "Document Number": "F-19"},
            {"Document Name": "Documents Change Request Form", "Document Number": "F-20"}
        ],
        # ...continue for all process, operation, audit, management review, nonconformance etc.
    }

    patterns = [
        ("ims_org",   "Org initials + IMS (XXX-IMS-...)",           pattern_1),
        ("ims_only",  "IMS only (IMS-...)",                         pattern_2),
        ("qhse",      "QHSE system (QHSE-...)",                     pattern_3),
        ("minimal",   "Minimal prefix (MAN-01, P-01, etc.)",        pattern_4),
    ]
    # Randomly select pattern
    pattern_name, pattern_desc, clause_map = random.choice(patterns)

    # Generate markdown table
    lines = [
        "| Clause | Document Name                           | Document Number |",
        "|--------|-----------------------------------------|-----------------|"
    ]
    for clause, docs in clause_map.items():
        for doc in docs:
            lines.append(f"| {clause}   | {doc['Document Name']} | {doc['Document Number']} |")
    prompt_table = "\n".join(lines)

    return pattern_name, pattern_desc, clause_map, prompt_table


def build_nc_section(updated_rows):
    nc_lines = []
    for row in updated_rows:
        if row.get("C/NC/O", "").strip().upper() == "NC":
            clause_no = row.get("Cl. NO", "").strip()
            detail = row.get("Document Verification detail with statement of Conformity", "").strip()
            nc_lines.append(f"Clause {clause_no}: {detail}\n")
    return "\n".join(nc_lines) if nc_lines else "No non-conformities found."

def build_observation_section(updated_rows):
    obs_lines = []
    for row in updated_rows:
        if row.get("C/NC/O", "").strip().upper() == "O":
            clause_no = row.get("Cl. NO", "").strip()
            detail = row.get("Document Verification detail with statement of Conformity", "").strip()
            obs_lines.append(f"Clause {clause_no}: {detail}\n")
    return "\n".join(obs_lines) if obs_lines else "No observations found."

def insert_audit_sections(patched_buffer, nc_text, obs_text):
    patched_buffer.seek(0)
    doc = Document(patched_buffer)
    # Non-Conformities
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Non-Conformities Raised" in cell.text:
                    cell.text = "Non-Conformities Raised:\n\n" + nc_text
    # Observations/SUMMARY
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "SUMMARY (including general observations/comments" in cell.text or "SUMMARY" in cell.text:
                    cell.text = "SUMMARY (including general observations/comments-Separate Sheet can be used):\n\n" + obs_text
    patched_buffer.seek(0)
    patched_buffer.truncate(0)
    doc.save(patched_buffer)
    patched_buffer.seek(0)
    return patched_buffer

def patch_docx_by_row_index(docx_buffer, audit_rows):
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    table = doc.tables[3]   # <-- THE CRITICAL FIX
    data_start_idx = 2      # still 2 header rows
    n_rows = min(len(audit_rows), len(table.rows) - data_start_idx)
    for i in range(n_rows):
        arow = audit_rows[i]
        trow = table.rows[i + data_start_idx]
        if len(trow.cells) < 4:
            print(f"Row {data_start_idx + i} has too few cells!")
            continue
        trow.cells[0].text = str(arow.get("Cl. NO", ""))
        trow.cells[1].text = str(arow.get("Description", ""))
        trow.cells[2].text = str(arow.get("C/NC/O", ""))
        trow.cells[3].text = str(arow.get("Document Verification detail with statement of Conformity", ""))
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer



def mistral_response_to_updated_rows(raw):
    import json, re
    if isinstance(raw, str):
        raw = json.loads(raw)
    codeblock = raw.get("response", "")
    json_str = re.sub(r'^`{3}json\s*|\s*`{3}$', '', codeblock.strip())
    return json.loads(json_str)


def patch_docx_buffer_with_na(docx_buffer, clause_status_dict):
    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    for table in doc.tables:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if "Clause & Description" in first_row:
            cl_no_idx = 0
            desc_idx = 1
            status_idx = 3  # Update if needed!
            for row in table.rows[2:]:
                cl_no = row.cells[cl_no_idx].text.strip()
                desc = row.cells[desc_idx].text.strip()
                clause_key = f"{cl_no} - {desc}" if cl_no and desc else cl_no or desc
                if clause_key in clause_status_dict:
                    row.cells[status_idx].text = clause_status_dict[clause_key]
            break
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)


def mark_na_clauses(extracted_data, na_clauses):
    """
    Sets exact-matched subclauses as NA, based on na_clauses being a list of strings
    like '4.1 - Understanding the organization and its context'
    """
    if not na_clauses:
        return extracted_data
    normalized = set(c.strip() for c in na_clauses)
    for key in extracted_data:
        if key.strip() in normalized:
            extracted_data[key] = "\t\t\tNA"
    return extracted_data



def extract_audit_table_iso9001_stage1(docx_path_or_stream):
    doc = Document(docx_path_or_stream)
    extracted = {}

    for table in doc.tables:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        if (
            "Clause & Description" in first_row
            and "Document Verification detail with statement of Conformity" in first_row
        ):
            second_row = [cell.text.strip() for cell in table.rows[1].cells]
            try:
                cl_no_idx = second_row.index("Cl. NO")
                desc_idx = second_row.index("Description")
                doc_ver_idx = next(
                    i for i, h in enumerate(first_row)
                    if "Document Verification detail with statement of Conformity" in h
                )
            except ValueError:
                continue

            for row in table.rows[2:]:
                cells = row.cells
                if len(cells) <= max(cl_no_idx, desc_idx, doc_ver_idx):
                    continue
                cl_no = cells[cl_no_idx].text.strip()
                desc = cells[desc_idx].text.strip()
                doc_ver = cells[doc_ver_idx].text.strip()
                key = f"{cl_no} - {desc}" if cl_no and desc else cl_no or desc
                extracted[key] = doc_ver
            break
    return extracted

def extract_audit_table_iso9001_stage1_ordered(docx_path_or_stream):
    doc = Document(docx_path_or_stream)
    data = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header_row = [cell.text.strip() for cell in table.rows[1].cells]
        colnames_row = table.rows[1]
        first_row_texts = [cell.text.strip() for cell in table.rows[0].cells]
        if (
            "Cl. NO" in header_row
            and "Description" in header_row
            and "Document Verification detail with statement of Conformity" in first_row_texts
        ):
            try:
                cl_no_idx = header_row.index("Cl. NO")
                desc_idx = header_row.index("Description")
                status_colnames = ["C/NC/O", "Status", "C / NC / O", "Conformity"]
                status_idx = next(
                    (i for i, val in enumerate(header_row) if val in status_colnames), 2
                )
                doc_ver_idx = next(
                    (ix for ix, txt in enumerate(first_row_texts)
                     if "Document Verification detail with statement of Conformity" in txt),
                    None
                )
                if doc_ver_idx is None:
                    continue
            except ValueError:
                continue

            for row in table.rows[2:]:
                cells = row.cells
                if len(cells) <= max(cl_no_idx, desc_idx, status_idx, doc_ver_idx):
                    continue
                data.append({
                    "Cl. NO": cells[cl_no_idx].text.strip(),
                    "Description": cells[desc_idx].text.strip(),
                    "C/NC/O": cells[status_idx].text.strip(),
                    "Document Verification detail with statement of Conformity": cells[doc_ver_idx].text.strip(),
                })
            break
    return data

def update_cnc_placeholders(rows):
    indices_with_placeholder = [
        idx for idx, row in enumerate(rows)
        if row["C/NC/O"] == "{{clause}}" and row["Document Verification detail with statement of Conformity"] != "NA"
    ]
    total = len(indices_with_placeholder)
    if total == 0:
        return rows

    nc_count = min(2, max(1, math.floor(0.1 * total)))
    o_count  = max(1, math.ceil(0.1 * total))
    c_count  = total - nc_count - o_count

    replacements = (["     C"] * c_count) + (["     O"] * o_count) + (["     NC"] * nc_count)
    random.shuffle(replacements)

    for i, idx in enumerate(indices_with_placeholder):
        rows[idx]["C/NC/O"] = replacements[i]

    for row in rows:
        if row["Document Verification detail with statement of Conformity"] == "NA":
            row["C/NC/O"] = ""
    return rows


def patch_docx_from_rows(docx_buffer, rows):
    # Robust key access for any LLM-drifted keys
    def keyval(row, *candidates):
        for k in candidates:
            if k in row:
                return row[k]
        # Debug print for missing keys:
        print("Row missing expected keys! Actual keys present:", list(row.keys()))
        print("Row content:", row)
        raise KeyError(f"Row missing keys {candidates}")

    key_to_cnc = {
        (
            keyval(row, "Cl. NO", "Cl. NO.", "Clause No", "Clause No."),
            keyval(row, "Description", "description")
        ): keyval(row, "C/NC/O", "CNC/O", "Status", "status")
        for row in rows
    }

    docx_buffer.seek(0)
    doc = Document(docx_buffer)
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header_row_1 = [cell.text.strip() for cell in table.rows[0].cells]
        header_row_2 = [cell.text.strip() for cell in table.rows[1].cells]
        if "Clause & Description" in header_row_1:
            cl_no_idx = header_row_2.index("Cl. NO")
            desc_idx = header_row_2.index("Description")
            cnc_idx = None
            for idx, val in enumerate(header_row_2):
                if val in ["C/NC/O", "C / NC / O", "Status"]:
                    cnc_idx = idx
                    break
            if cnc_idx is None: cnc_idx = 2  # Default/fallback
            for row in table.rows[2:]:
                cl_no = row.cells[cl_no_idx].text.strip()
                desc = row.cells[desc_idx].text.strip()
                key = (cl_no, desc)
                if key in key_to_cnc:
                    row.cells[cnc_idx].text = key_to_cnc[key]
            break
    docx_buffer.seek(0)
    docx_buffer.truncate(0)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


# =========================== ISO:9001 STAGE 1 FUNCTIONS END HERE ===================================================================


app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- MODELS ----------

class ISO9001Stage1Audit(BaseModel):
    organizationName: str
    address: str
    siteAddress: Optional[str] = ""
    numberOfEmployees: int
    emailId: EmailStr
    contactPerson: str
    telephoneFax: str
    scope: str
    riskCategory: str
    iafCode: str
    auditTeam: List[str]
    auditManDays: str
    startDateOfAudit: str
    endDateOfAudit: str
    auditMode: str
    quotedManDaysAdequate: str
    changeInEmployeeDetail: str
    changeInScope: str
    additionalInformation: str
    internalAuditFrequency: str
    dateOfLastInternalAudit: str
    managementReviewFrequency: str
    dateOfLastManagementReview: str
    recommendationForStage2: str
    reviewedBy: str
    na_clauses: Optional[List[str]] = []
    attendanceSheet: List[str]
    dateOfReview: str
    clientName: str
    designation: str
    auditorName: str

class ISO9001Stage2Audit(BaseModel):
    organizationName: str
    address: str
    siteAddress: Optional[str] = ""
    numberOfEmployees: int
    emailId: EmailStr
    contactPerson: str
    telephoneFax: str
    scope: str
    riskCategory: str
    iafCode: str
    auditTeam: List[str]
    auditManDays: str
    startDateOfAudit: str
    endDateOfAudit: str
    auditMode: str
    quotedManDaysAdequate: str
    changeInEmployeeDetail: str
    changeInScope: str
    additionalInformation: str
    internalAuditFrequency: str
    dateOfLastInternalAudit: str
    managementReviewFrequency: str
    dateOfLastManagementReview: str
    recommendationForStage2: str
    reviewedBy: str
    na_clauses: Optional[List[str]] = []
    attendanceSheet: List[str]
    dateOfReview: str
    clientName: str
    designation: str
    auditorName: str



# ---------- ROUTE ----------
@app.get("/test")
async def test():
    return {"message": "hello"}

@app.post("/iso9001/stage1/submit")
async def submit_iso9001_stage1(audit: ISO9001Stage1Audit):

    doc = DocxTemplate("templates/iso9001_stage1.docx")
    context = {
        "organizationName": audit.organizationName,
        "address": audit.address,
        "siteAddress": audit.siteAddress,
        "numberOfEmployees": audit.numberOfEmployees,
        "emailId": audit.emailId,
        "contactPerson": audit.contactPerson,
        "telephoneFax": audit.telephoneFax,
        "scope": audit.scope,
        "riskCategory": audit.riskCategory,
        "iafCode": audit.iafCode,
        "auditTeam": "\n".join(audit.auditTeam),
        "auditManDays": audit.auditManDays,
        "startDateOfAudit": audit.startDateOfAudit,
        "endDateOfAudit": audit.endDateOfAudit,
        "auditMode": audit.auditMode,
        "quotedManDaysAdequate": audit.quotedManDaysAdequate,
        "changeInEmployeeDetail": audit.changeInEmployeeDetail,
        "changeInScope": audit.changeInScope,
        "additionalInformation": audit.additionalInformation,
        "internalAuditFrequency": audit.internalAuditFrequency,
        "dateOfLastInternalAudit": audit.dateOfLastInternalAudit,
        "managementReviewFrequency": audit.managementReviewFrequency,
        "dateOfLastManagementReview": audit.dateOfLastManagementReview,
        "recommendationForStage2": audit.recommendationForStage2,
        "reviewedBy": audit.reviewedBy,
        "na_clauses": audit.na_clauses,
        "dateOfReview": audit.dateOfReview,
        "clientName": audit.clientName,
        "designation": audit.designation,
        "auditorName": audit.auditorName,
        "clause": "{{clause}}",
    }

    doc.render(context)
    extract_buffer = io.BytesIO()
    doc.save(extract_buffer)
    extract_buffer.seek(0)

    extracted_data = extract_audit_table_iso9001_stage1(extract_buffer)
    extracted_data = mark_na_clauses(extracted_data, audit.na_clauses)

    patch_docx_buffer_with_na(extract_buffer, extracted_data)
    rows = extract_audit_table_iso9001_stage1_ordered(extract_buffer)
    updated_rows = update_cnc_placeholders(rows)
    patched_buffer = patch_docx_from_rows(extract_buffer, updated_rows)

    pattern_name, pattern_desc, clause_map, prompt_table = choose_document_pattern()
    print(f"Pattern chosen: {pattern_desc}\n\n{prompt_table}")

    prompt = f"""
    You are an ISO 9001 audit reporting assistant.

    For this audit, use the following document numbering pattern: {pattern_desc}
    When mentioning any document as evidence for a clause, use the document name and document number from the table below. 
    If a document number uses a prefix like "XXX" or "BLPL", replace it with the initials of the organization's name 
    (for example, "Inzinc Consulting LLC" → "ICL-IMS-F-01"). If initials are not clear, use the first letter of each word.

    {prompt_table}

    Instructions for rephrasing:
    - For each row in the provided list (representing one audit clause), follow these rules:
        - If the "C/NC/O" value is "C", rephrase the "Document Verification detail with statement of Conformity" as a professional positive confirmation. 
          Clearly state requirements are confirmed, and use the correct document evidence for that clause from the table above.
        - If the "C/NC/O" value is "NC", rephrase as a documented nonconformity, stating what is nonconforming and referencing the relevant document(s) for that clause.
        - If the "C/NC/O" value is "O", rephrase as a neutral observation, mentioning the relevant document(s) for that clause from the table.
        - Leave other keys and their values (especially status) unchanged.
        - Do NOT add, remove, or restructure fields or rows.
        - Keep all keys in every dictionary exactly as in the input (including 'Cl. NO', 'Description', and 'C/NC/O', with case and punctuation).
        - Use the organization’s full name (“{audit.organizationName}”) only the first time it is mentioned in findings. Later, use "the organization", "the company", or "it".
        - Replace any dates in document references so that they fall within 6–10 months prior to the audit date. Update dates as needed.
        - Adjust any sample names, designations, or client/supplier data based on the attendance sheet and scope.

    Organization Name: {audit.organizationName}
    Scope: {audit.scope}
    Start Date of Audit: {audit.startDateOfAudit}
    End Date of Audit: {audit.endDateOfAudit}
    Attendance Sheet: {json.dumps(audit.attendanceSheet, ensure_ascii=False)}

    Audit Content to rephrase and adapt:
    {json.dumps(updated_rows, indent=2, ensure_ascii=False)}

    Output ONLY the resulting list of dictionaries, nothing else (no explanation, no code block, no extra whitespace).
    """

    mistral_api_url = "https://mistral-api-5icm.onrender.com/api/mistral"
    payload = {
        "prompt": prompt,
    }
    headers = {
        "Content-Type": "application/json"
    }

    try:
        async with httpx.AsyncClient(timeout=300.0) as client:
            response = await client.post(mistral_api_url, json=payload, headers=headers)
            response.raise_for_status()

            if response.headers.get("content-type") == "application/json":
                mistral_result = response.json()
                rephrased_text = mistral_result.get("response", "")
            else:
                rephrased_text = response.text

        try:
            audit_rows = mistral_response_to_updated_rows(
                rephrased_text)
            print("✅ Updated old rows to LLM output success")
            updated_rows = audit_rows

        except Exception as e:
            return JSONResponse(
                status_code=502,
                content={"error": f"Failed to get response from Mistral API: {str(e)}"}
            )

    except Exception as e:
        print(f"Error calling Mistral API: {e.__class__.__name__}: {e!r}")
        traceback.print_exc()
        rephrased_text = json.dumps(extracted_data, indent=2)

    patched_buffer = patch_docx_by_row_index(extract_buffer, updated_rows)
    nc_text = build_nc_section(updated_rows)
    obs_text = build_observation_section(updated_rows)
    patched_buffer = insert_audit_sections(patched_buffer, nc_text, obs_text)
    patched_buffer = await add_org_brief_to_docx(
        patched_buffer,
        audit.organizationName,
        audit.scope
    )
    final_doc_bytes = patched_buffer.getvalue()

    headers = {
        "Content-Disposition": f"attachment; filename={audit.organizationName}_iso9001_stage1_report.docx"
    }

    return Response(
        content=final_doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )




@app.post("/iso9001/stage2/submit")
async def submit_iso9001_stage2(audit : ISO9001Stage2Audit):
    doc = DocxTemplate("templates/iso9001_stage2.docx")
    context = {
        "organizationName": audit.organizationName,
        "address": audit.address,
        "siteAddress": audit.siteAddress,
        "numberOfEmployees": audit.numberOfEmployees,
        "emailId": audit.emailId,
        "contactPerson": audit.contactPerson,
        "telephoneFax": audit.telephoneFax,
        "scope": audit.scope,
        "iafCode": audit.iafCode,
        "auditTeam": "\n".join(audit.auditTeam),
        "auditManDays": audit.auditManDays,
        "startDateOfAudit": audit.startDateOfAudit,
        "clientName": audit.clientName,
        "designation": audit.designation,
        "auditorName": audit.auditorName,
        "endDateOfAudit": audit.endDateOfAudit
    }

    doc.render(context)
    extract_buffer = io.BytesIO()
    doc.save(extract_buffer)
    extract_buffer.seek(0)

    final_doc_bytes = extract_buffer.getvalue()
    headers = {
        "Content-Disposition": f"attachment; filename={audit.organizationName}_iso9001_stage2_report.docx"
    }
    return Response(
        content=final_doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )

